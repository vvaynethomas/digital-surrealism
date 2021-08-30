#! multidocumentmultitypewriter.py (MadMat) is a package of functions for scraping and transforming texts
## with the name of a poet, you can scrape all the works available on poetryfoundation.org and find their common props
## via the create_altered_texts function, one can shuffle the lines in a poem, construct a text from multiple texts
## by the same writer (or with the same keyword) or by multiple writers (or according to multiple keywords)
## you can also trade proper nouns or any other part of speech (eventually) between individual texts or writers/keywords

import docx
import os
import re
from random import sample, randint, shuffle
from nltk import tokenize, pos_tag, FreqDist, corpus
import PyPDF2
import string
import shutil
import matplotlib.pyplot as plt
# from poetry_foundation_scraper import scrape_all_poems
english_vocab = set(w.lower() for w in corpus.words.words())
def extract_from_pdf(filepath,type='page',subtype='string'):
    # filename = os.path.basename(filepath)
    with open(filepath, 'rb') as pdf_file_obj:  # open current pdf
        pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)  # store pdf as PyPDF2 reader object
        if pdf_reader and not pdf_reader.isEncrypted:
            # print("I'm looking for a line in " + filename)
            number_of_pages = pdf_reader.getNumPages()  # count number of pages in pdf
            if number_of_pages > 0:
                if type == 'all' :
                    text_in_lines = []
                    text_in_pages = []
                    for i in range(0,number_of_pages - 1):
                            page_obj = pdf_reader.getPage(i)
                            page_text = page_obj.extractText()
                            # page_as_list = page_obj.extractText().splitlines()
                            text_in_pages.append(page_text)
                            # for line in page_as_list:
                            #     text_in_lines.append(line)
                    return text_in_pages, text_in_lines
                elif type == 'page':
                    page_number = randint(0,number_of_pages-1)
                    page_obj = pdf_reader.getPage(page_number)
                    if subtype == 'lines':
                        page_in_lines = page_obj.extractText().splitlines()
                        # text_in_pages = [text_in_lines]
                        return page_in_lines, page_number
                    else:
                        page_as_string = page_obj.extractText()
                        return page_as_string, page_number

                elif type == 'line':
                    fail_count = 0
                    line = ""
                    page_num = 0
                    while not line and fail_count < number_of_pages:
                        page_num = randint(1, number_of_pages - 1)  # generate random index to choose page
                        page_obj = pdf_reader.getPage(page_num)  # reads and stores random page as page object
                        try:
                            page_in_lines = page_obj.extractText().splitlines()  # stores extracted text as multiline string
                            num_lines = len(page_in_lines)  # count number of lines detected
                            if num_lines > 0:
                                print("we've got lines here")
                                line = sample(page_in_lines,1)[0]
                                return line, page_num
                            else:
                                print("no lines detect on page " + page_num)
                                line = ""
                                fail_count +=1
                                continue

                        except:
                            print('failed to extract a line from pdf page ' + page_num)
                            fail_count +=1
                            continue
                    return line, page_num

                elif type == 'sentence':
                    fail_count = 0
                    sentence = ""
                    page_num = 0
                    while not sentence and fail_count < number_of_pages:
                        page_num = randint(0, number_of_pages-1)
                        page_obj = pdf_reader.getPage(page_num)
                        try:
                            page_as_string = page_obj.extractText()
                            page_in_sentences = tokenize.sent_tokenize(page_as_string)
                            sentence = sample(page_in_sentences,1)[0]
                        except:
                            print('failed to extract a sentence from pdf page ' + str(page_num))
                            fail_count +=1
                            continue
                    return sentence, page_num

def get_filepaths(directory):
    """
    This function will generate the file names in a directory
    tree by walking the tree either top-down or bottom-up. For each
    directory in the tree rooted at directory top (including top itself),
    it yields a 3-tuple (dirpath, dirnames, filenames).
    """
    file_paths = []  # List which will store all of the full filepaths.

    # Walk the tree.
    for root, directories, files in os.walk(directory):
        for filename in files:
            # Join the two strings in order to form the full filepath.
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)  # Add it to the list.

    return file_paths  # Self-explanatory.

def get_document_paths(file_paths):
    paths_for_potentials = []
    for path in file_paths:
        word = path
        if word.endswith('.docx') or word.endswith('.pdf') or word.endswith('.doc'):
            paths_for_potentials.append(path)
    return paths_for_potentials

def get_docx_paths(file_paths):
    paths_for_docs = []
    for path in file_paths:
        word = path
        regexp = re.compile(r'doc')
        if regexp.search(word):
            paths_for_docs.append(path)
    return paths_for_docs

def convert_pdf_to_doc(pdf_path,output_directory):
    os.chdir(output_directory)
    new_filename = os.path.basename(pdf_path).replace('.pdf','.docx')
    f = docx.Document()
    text_in_pages, text_in_lines = extract_from_pdf('all', pdf_path)
    for page in text_in_pages:
        for line in page:
            p = f.add_paragraph()
            p.add_run(line)
            p.add_run().add_break()
        f.add_page_break()
    f.save(new_filename)
    return text_in_lines

def store_source(sources,filename,section_num,sections_and_sources,section):
    if filename.endswith('.pdf'):
        if filename not in sources:
            sources[filename] = {'page or section number': section_num}  # store source of current line being generated
    elif filename.endswith('.doc') or filename.endswith('.docx'):
        if not filename in sources:
            sources[filename] = {'page or section number': section_num}  # store source of current line being generated
    if str(section) not in sections_and_sources:
        sections_and_sources[str(section)] = [filename,section_num]
    else:
        sections_and_sources[str(section)].append([filename,section_num])

def extract_from_document(file_path,type='line'):
    filename = os.path.basename(file_path)
    if filename:
        line = ""
        section_num = 0

        if filename.endswith('.pdf'):
            # print(filename)
            try:
                if type == 'all':
                    text_in_pages, text_in_lines = extract_from_pdf(file_path, type='all')
                    return text_in_pages, text_in_lines
                line, section_num = extract_from_pdf(file_path,type)
                return line, section_num
            except:
                print("unreadable pdf")
                number_extracted = 0
                return line, section_num
        elif filename.endswith('.doc') or filename.endswith('.docx'):
            # print("scraping " + filename + " at " + file_path + " for a line")
            # try:
            all_text = extract_all_text_from_docx(file_path)
            print('got all text from doc')
            if len(all_text) >= 1:
                if type == "line":
                    line_number = randint(0,len(all_text)-1)
                    line = all_text[line_number]
                    print('scraped a line')
                    return line, line_number
                elif type == "sentence":
                    all_sentences = tokenize.sent_tokenize(" ".join(all_text))
                    sentence_no = randint(0,len(all_sentences)-1)
                    sentence = all_sentences[sentence_no]
                    print('scraped a sentence')
                    return sentence, sentence_no
            else:
                print('no text here')
                line = ''
                line_number = 0
                return line, line_number
            return all_text

            # except:
            #     print("not docx friendly")
            #     number_extracted = 0
            #     return line, section_num, number_extracted

def grab_words(path, max):
    fail_count = 0
    word_list = []
    if path.endswith('.pdf'):
        page_text = ""
        while not page_text or len(page_text) < 2:
            page_text = extract_from_pdf(path)
            fail_count += 1
            if fail_count > 50:
                word_list = []
                return word_list

        page_as_sentences = split_text([page_text])
    elif path.endswith('.docx'):
        page_text = ""
        while not page_text or len(page_text) < 2:
            page_text = extract_all_text_from_docx(path)
            if not isinstance(page_text,str):
                try:
                    page_text = " ".join(page_text)
                except:
                    lines = ""
                    for run in page_text:
                        lines = lines + run.text
                    page_text = lines

        number_of_words = randint(2,max)
        page_in_words = page_text.split()
        start = randint(0,len(page_in_words)- number_of_words - 1)
        interval = range(start,start+number_of_words)
        for i in interval:
            word_list.append(page_in_words[i])

def extract_all_text_from_docx(filepath,paragraph_or_runs='runs'):
    try:
        doc = docx.Document(filepath)
        if doc:
            # print("attempting to scrape " + os.path.basename(filepath))
            file_fail_count = 0
            all_paragraphs = []
            all_runs = []
            length = len(doc.paragraphs)
            # print(str(length))
            if length >= 1:
                for paragraph in doc.paragraphs:
                    all_paragraphs.append(paragraph)
                    for run in paragraph.runs:
                        all_runs.append(run)
                if paragraph_or_runs == 'paragraphs':
                    return all_paragraphs
                else:
                    print("got all runs")
                    return all_runs
    except:
        print("couldn't scrape it")
        all_runs = []
        return all_runs

def split_text(string_list):
    print("let's see if we need to split this up into sentences")
    list_of_strings = []
    for run in string_list:
        list_of_strings.append(run)

    if '.' in [run for run in list_of_strings]:
            text_list = tokenize.sent_tokenize(" ".join(list_of_strings))
    else:
        text_list = list_of_strings

    return text_list

def scrape_numbering_tags(string):
    # print("let's test for numbering/sectional tags so they don't get lumped in as lines")
    s = string
    if not isinstance(string,str):
        s = string.text
    if re.search('[iv]+\.',s):
        lst = re.findall('[iv]+\.', s)
        for tag in lst:
            s.replace(tag,"")
    elif re.search('\d+\.',s):
        lst = re.findall('\d+\.', s)
        for tag in lst:
            s.replace(tag,"")
        # s.replace([i for i in lst],"")
    else:
        lst = []
    return s, lst

def get_number_of_lines(path):
    all_text_list = extract_all_text_from_docx(path)
    if all_text_list:
        all_numbers = []
        strings_no_numbers = []
        for string in all_text_list:
            new_string, tag_list = scrape_numbering_tags(string)
            strings_no_numbers.append(new_string)
            for tag in tag_list:
                all_numbers.append(tag)
        for line in strings_no_numbers:
            if line.isspace():
                strings_no_numbers.remove(line)
        return len(strings_no_numbers)
    else:
        return 0

def local_shuffle(region_name,specific_poem_or_random="",poem_name = ""): #shuffles line in a poem or each poem in a region
    print("hmm...")
    input_directory = 'C:/Users/Vvayne/Documents/FA 2017/Major Authors/Material/' # initialize global input folder
    output_directory = 'C:/Users/Vvayne/Documents/FA 2017/Major Authors/Outputs/local shuffles/' + region_name + '/' # initialize where to store new constructions
    if not specific_poem_or_random: #test for user specified text or single random text
        print("let's shuffle them all")
        path = input_directory + region_name + '/' #set path to whole region if user has not requested specified or random text
    else:
        if specific_poem_or_random == "specific": #test if user wants specific poem shuffled
            print("oh! you already know which one you want to shuffle")
            poem_name = poem_name.replace(" ", '-') #replace spaces in poem name with dashes
            path = input_directory + region_name + '/' + poem_name + '.docx' #stores specified poem path as current path
        else:
            print("let's just grab one")
            path = sample(get_docx_paths(get_filepaths(input_directory+region_name+'/')),1)[0] # randomly selects a document path in the region
    if path: # test taht path exists and is non-empty
        print(path)
        all_doc_paths = get_docx_paths(get_filepaths(path)) #filter paths in directory or singular path for documents only
        if len(all_doc_paths) == 0: #if path defined has no sub-paths
            all_doc_paths = [path] # establish path as singular list
        print(len(all_doc_paths))
        os.chdir(output_directory) # change current directory to output directory
        print(output_directory)
        for doc in all_doc_paths: # loop over all document paths selected
            f = docx.Document() # start a new word document
            all_text_list = extract_all_text_from_docx(doc) #store all text extracted from current document
            print(str(len(all_text_list)))
            all_numbers = [] #initialize list to store number tags
            strings_no_numbers = [] #initialize list to store string minus number tags
            for string in all_text_list: #loop over each line in extracted text
                new_string, tag_list = scrape_numbering_tags(string) #calls function to filter numbering tags
                strings_no_numbers.append(new_string) #add filtered line to new list
                for tag in tag_list: #loop over numbering tags
                    all_numbers.append(tag) #add numbering tag to list

            print(len(strings_no_numbers))
            # for line in strings_no_numbers: #loop over each line in text
            #     if line.isspace(): # test that line is not blank
            #         strings_no_numbers.remove(line) #remove blank line
            shuffled_list = sample(strings_no_numbers, len(strings_no_numbers)) # generate a new list of lines from selecting each at random from the filtered list
            for item in strings_no_numbers:
                if item not in shuffled_list:
                    shuffled_list.append(item)
            print(len(shuffled_list))
            numbered_lines = sample(range(0,len(shuffled_list)-1),len(all_numbers)) # determine randomly which lines will be pre-fixed with numbering tags
            if any(line.endswith('!') for line in shuffled_list): # tests for any line that ends in a period
                f.add_paragraph(" ".join(shuffled_list)) # write text as block of prose
            else:
                p = f.add_paragraph() #start new paragraph
                all_added = []
                for i in range(0,len(shuffled_list)): # loop over lines by number to match numbering
                    if i == any(l for l in numbered_lines) and not shuffled_list[i]: #test if current number line matches the list
                        p.add_run(str(numbered_lines[i]) + " " + shuffled_list[i]) # write number tag and line to page
                    else:
                        p.add_run(shuffled_list[i])  # add line to page
                    all_added.append(shuffled_list[i])
                    p.add_run().add_break() # add space
                for line in shuffled_list:
                    if not line in all_added:
                        p.add_run(line)
                        all_added.append(line)
                print(len(all_added))
            name = os.path.basename(doc).replace('.docx','') # initialize name for new text
            f.save(name + str(randint(0,300)) + '.docx') # save text with random integer and type suffix
            print("saved shuffled version of " + name)
    else:
        print("something wrong with these paths")

def intrashuffle(input_paths,output_directory): # generate or construct text from any lines in any text in the same region/directory
    region_name = os.path.dirname(sample(input_paths,1)[0]).split('/')[-1] #extract current region name
    print(region_name)
    os.chdir(output_directory) # change current working directory to output directory
    copy_output_directory = output_directory
    print(output_directory)
    f = docx.Document() #start a new word document to store new text
    line_count = 0 # initialize line counter
    number_of_lines_for_each = [] #initialize list of line numbers for texts in the region
    for path in input_paths: # loop through each document in the region
        number_of_lines_for_each.append(get_number_of_lines(path)) # add number of lines in document to list


    number_of_lines = sample(number_of_lines_for_each,1)[0] # get number of lines for current piece by randomly selecting a length from the new list
    print("let's get " + str(number_of_lines) + ' lines')
    sources = [] # initialize list of sources
    while line_count < number_of_lines: # loop until number of lines defined has been reached
        material_path = sample(input_paths, 1)[0] # randomly select document in region
        line = sample(extract_all_text_from_docx(material_path), 1)[0] # randomly select line from the document
        # print(output_directory)
        if line: # test that line has been acquired and is not empty
            sources.append(os.path.basename(material_path).replace(".docx", "")) #add source of line to list of sources
            p = f.add_paragraph() # start new block in word document started earlier
            p.add_run(line.text) # write line to new block
            line_count += 1 # add 1 to line counter
            if number_of_lines % line_count: # arbitrary test for block breaking
                p.add_run().add_break() # add break before next block
    print(copy_output_directory)
    f.add_page_break() # start new page in document
    f.add_heading('Sources', 1) # write 'Sources' at the top of the page
    last_p = f.add_paragraph() # start final block
    set_of_sources = [] # initialize list of different sources (without duplicates)
    for source in sources: # loop through each source that's been collected
        if not source in set_of_sources: # test if current source is already in the list of unique sources
            set_of_sources.append(source) # store source in in list of different sources if it isn't already in there
    for source in set_of_sources: # loop through unique sources
        last_p.add_run(source) # write source to page
        last_p.add_run().add_break() # add space before next source
    name = region_name + str(len(set_of_sources)) + '_' + str(randint(0,900)) + '.docx' # define name of new document based on how many different texts contributed and a random integer
    print(name)
    output_directory = 'C:/Users/Vvayne/Documents/FA 2017/Major Authors/Outputs/regional shuffles/' + region_name + '/'
    os.chdir(output_directory)
    f.save(name) # save the new document
    print("Wrote a regionally shuffled poem called: " + name)
    print(os.path.join(output_directory, name))
    print("It's got " + str(line_count) + " lines from " + str(len(set_of_sources)) + " different poems")
    print("You can find it in " + output_directory)

def construct_intrashuffled_poems(region_name, number_of_new_texts=1): #batch intrashuffled text creator
    print("let's shuffle up " + str(number_of_new_texts) + " new text(s)")
    input_directory = 'C:/Users/Vvayne/Documents/FA 2017/Major Authors/Material/' + region_name + '/' #initialize folder for inputs
    output_directory = 'C:/Users/Vvayne/Documents/FA 2017/Major Authors/Outputs/regional shuffles/' + region_name + '/' #initialize location to store new texts
    all_input_paths = get_docx_paths(get_filepaths(input_directory)) # list all paths in input directory
    for i in range(0, number_of_new_texts): #loop construct number of poems desired
        print("let's shuffle up " + str(number_of_new_texts-i) + " new text(s)")
        intrashuffle(all_input_paths, output_directory) # call function that deals with single intraregional construction
def common_elements(list1, list2):
    return [element for element in list1 if element in list2]
def refrain_or_no(all_lines_so_far): # function to determine if the next line should be a refrain and if so what should it be
    if len(all_lines_so_far) > 2: #test for at least two lines before we start getting crazy with repetition
        source_for_lines = all_lines_so_far #store a copy of the lines
        roll_of_die = randint(1,len(source_for_lines)) #generate random integer between 1 and the number of available lines for refrain
        if roll_of_die >= randint(7,9)*.1*len(source_for_lines): #probabilistic test to confirm refrain
            line = sample(all_lines_so_far, 1)[0] #store refrain as line
            return 'yes', line  #return string and line
        else: #if refrain 'failed'
            line = "" #store nothing as line
            return 'no', line #return string and empty line
    else: # if 2 or fewer lines
        line = "" #store nothing as line
        return 'no', line #return no and empty line

def construct_global_shuffle(input_directory,output_directory,mode="verse",number_of_lines= randint(5,40),english=False): #function to produce a text formed by lines from a sample of all available paths in a directory
    if mode == "verse":
        type = "line"
    else:
        type = "sentence"

    all_lines_so_far = [] # initialize list of lines
    input_paths = get_document_paths(get_filepaths(input_directory)) #store all available paths in a list
    print('got all the paths')
    os.chdir(output_directory) #print change current directory to output directory
    print(output_directory)
    f = docx.Document() #initialize new document to write and save construction
    line_count = 0 #initialize line counter

    # number_of_lines = max(5,min(40,get_number_of_lines(sample(input_paths,1)[0]))) #define number of lines in new construction between 5 and 40 and ideally equal to the length of a random text in the materials folder
    print("let's get " + str(number_of_lines) + ' ' + type +'s')
    sources = [] #initialize list to store the origin of each line
    while line_count < number_of_lines: # loop until construction has desired number of lines
        # print(str(line_count) + ' out of ' + str(number_of_lines))
        line = "" #initialize new line variable
        refrain_boolean, refrain = refrain_or_no(all_lines_so_far)  # random integer to decide if a line should be repeated and returns it if yes
        # print(refrain_boolean)
        if refrain_boolean == 'yes':  # test for refrain_affirmation
            line = refrain  # stores old line as new line for refrain
            print('refrained')
        while not line or not isinstance(line,str) or not line[0].isalpha(): # loop until a line is acquired
            material_path = sample(input_paths, 1)[0] # get a random document path
            # fail_count = 0
            print('working on getting a line')
            line, location = extract_from_document(material_path,type) # calls function and stores as a list
            print('something like a line extracted')
            if not isinstance(line,str):
                line = line.text

            # print(current_text_as_list[0])
            # if material_path.endswith('.pdf') and not isinstance(current_text_as_list[0],str):
            #     try:
            #         while not isinstance(current_text_as_list[0], str):
            #             total_pages = len(current_text_as_list)
            #             page_number = randint(0, total_pages - 1)
            #             current_text_as_list = current_text_as_list[page_number]
            #             print(current_text_as_list[0])
            #     except:
            #         continue

            # if refrain_boolean == 'no' and len(current_text_as_list) >= 1: # tests against if refrain is False and also that the text is more than just one string
            #     print('getting new line from ' + material_path)
            #     # print('not a refrain')
            #     line = sample(current_text_as_list, 1)[0] # store a line from the text
            # fail_count +=1
            # if fail_count > 50:
            #     continue

        if english == True:
            line = line.lower()
            line_in_words = tokenize.word_tokenize(line)
            line = ' '.join(common_elements(line_in_words,english_vocab))
            # for word in line_in_words:
            #     if not word in english_vocab:
            #         line_in_words.remove(word)
            # line = ' '.join(line_in_words)
        if line: # tests that we surely have a line and it's not empty

            if len(tokenize.sent_tokenize(line)) > 1:
                line = sample(tokenize.sent_tokenize(line),1)[0]
            if len(tokenize.word_tokenize(line)) > 10:
                line = ' '.join(tokenize.word_tokenize(line)[:10])
            print('got a line')
            print(line)
            sources.append(os.path.basename(os.path.dirname(material_path).split('/')[-1]) + " " + os.path.basename(material_path).replace(".docx", " p." + str(location))) #adds the name of the author and text to the list of sources
            p = f.add_paragraph() #starts a new paragraph/block in the new poem document
            if isinstance(line,str): # tests that the line variable is the right type, namely a string
                p.add_run(line) #adds scraped line to the new paragraph
                all_lines_so_far.append(line) #adds line to list of all lines
                line_count +=1 #adds 1 to line counter
                if number_of_lines % line_count: #tests to see if a break should occur by an arbitrary relationship
                    p.add_run().add_break() #adds break
            else: #deal with line being a docx.run or something else with a text property but isn't a string
                p.add_run(line.text) # write string version of line to text
                all_lines_so_far.append(line.text) #add string version of line to list of all lines
                line_count += 1 #add 1 to line counter
                if number_of_lines % line_count: #arbitrary test for break
                    p.add_run().add_break() #add paragraph break

    f.add_page_break() #starts a new page in document being constructed
    f.add_heading('Sources', 1) #write header called Sources on new document
    last_p = f.add_paragraph() #start final block
    set_of_sources = [] #initialize unique set of sources
    for source in sources: #loop through list of all sources for lines in the new document
        if not source in set_of_sources: #check for uniqueness
            set_of_sources.append(source) #add if unique
    for source in set_of_sources: #loop through unique sources
        last_p.add_run(source) #write source
        last_p.add_run().add_break() #give space before the next one
    name = 'global_shuffle_' + str(len(set_of_sources)) + '_' + str(randint(0, 900)) + '.docx' #generate a name using the number of different poems used and a random integer
    f.save(name) # save new document
    print("Wrote a globally shuffled poem called: " + name)
    print("It's got " + str(line_count) + " lines from " + str(len(set_of_sources)) + " different poems")
    print("You can find it in " + output_directory)

def proper_scraper(this_string): # function to scrape proper nouns from a string
    s = this_string
    exclude = ['“',"”","‘","—",'*','_',"’"] #list of characters that start some lines
    s_as_list = s.split() #split string into words
    for i in range (1,3): #pass filter over string twice in case double punctuation, etc.
        # print('passing filter through potential props; pass #' + str(i))
        # filter for words only deemed proper because they start with punctuation
        for ch in exclude: #loop over undesirables
            for word in s_as_list: #loop over words in string
                if word.startswith(ch): #test if word starts with undesirable character
                    word.replace(ch,'') #remove character for proper name scraping
        # clear up punctuation to match duplicates
        for word in s_as_list: #loop through words
            for c in word: #loop over character in word
                if not c.isalnum(): #make sure character is not punctuation
                    # print(c)
                    word.replace(c,'') #erase punctuation
                    # print(word)
            if "’" in word: #check for weird apostrophe
                word.replace("’","'") #replace weird apostrophe with regular apostrophe
            if '*' in word or word == '*': # check for weird asterisk thing
                word.replace('*','') #remove weird asterisk thing from being proper noun
            if "’" in word: #check for other weird apostrophe
                word.replace("’","'") #replace other weird apostrophe with regular apostrophe

    romanNumeralPattern = r'^M?M?M?(CM|CD|D?C?C?C?)(XC|XL|L?X?X?X?)(IX|IV|V?I?I?I?)$' #regex pattern to detect roman numerals
    tagged_string = pos_tag(s_as_list) #produce a list of tuples (word, POS) for each word in tag
    propnouns_list = [word for word,pos in tagged_string if pos == 'NNP' and word.isalnum()] #store words tagged as 'NNP' as list
    for word in propnouns_list: # loop through words in new list
        if re.search(romanNumeralPattern,word): # check for roman numerals
            propnouns_list.remove(word) # remove roman numeral from propnouns
    return propnouns_list

def get_all_props(pathlist,region):
    print("why don't we get the proper nouns from each text in " + region)
    all_props = []
    text_in_lines = []
    text_in_strings = []
    for path in pathlist:
        filename = os.path.basename(path)
        if filename.endswith('.pdf'):
            text_in_pages, text_in_lines = extract_from_pdf('all',path)
        elif filename.endswith('.docx'):
            text_in_lines = extract_all_text_from_docx(path)

        if text_in_lines:
            # print('got all text from ' + filename + " in the pursuit of named references")
            for line in text_in_lines:
                if not isinstance(line,str):
                    text_in_strings.append(line.text)
                else:
                    text_in_strings.append(line)
            propnouns_in_list = proper_scraper(" ".join(text_in_strings))
            for prop in propnouns_in_list:
                all_props.append(prop)

    print('got all specific signifiers for ' + region)
    return all_props

def get_most_common_props(props_w_repetition,number_to_get=10):
    fdist1 = FreqDist(props_w_repetition)
    most_common = fdist1.most_common(number_to_get)
    common_props = []
    for (a,b) in most_common:
        common_props.append(a)
    return common_props, fdist1

def get_regional_props_and_paths(region_directory):
    all_paths = []
    region = os.path.dirname(region_directory).split('/')[-1]
    intraregional_text_paths = get_document_paths(get_filepaths(region_directory))
    for itp in intraregional_text_paths:
        if not itp in all_paths:
            all_paths.append(itp)
    region_props = []
    region_props_w_repetition = get_all_props(intraregional_text_paths,region)
    number_to_get = 40
    region_props, fdist1 = get_most_common_props(region_props_w_repetition,number_to_get)
    return region_props, all_paths, fdist1

def get_regional_props_dict(region_directories_as_list,plot=False):
    regional_prop_dict = {}
    while not len(regional_prop_dict) == len(region_directories_as_list):
        for region_path in region_directories_as_list:
            print(region_path)
            region_name = os.path.dirname(region_path).split('/')[-1]
            if region_name:
                print(region_name)
            else:
                break

            region_props, intraregional_paths, fdist1 = get_regional_props_and_paths(region_path)
            regional_prop_dict[region_name] = region_props
            print(regional_prop_dict.get(region_name))
            if not plot == False:
                plot = fdist1.plot(len(region_props), cumulative=False)
    return regional_prop_dict
            # plt.savefig(os.path.join(region_path, "_".join([region_name, 'freqdist_plot.png'])))

def trade_props(region_directories_as_list,regional_prop_dict={}):
    print(region_directories_as_list)
    if not regional_prop_dict:
        regional_prop_dict = get_regional_props_dict(region_directories_as_list)
    # print(region_directories_as_list)
    shuffle(region_directories_as_list)
    copy_of_directories = region_directories_as_list
    for i in range(0,len(region_directories_as_list)-1):
        region_path = region_directories_as_list[i]
        input_directory = os.path.dirname(region_path)
        region_name = str(input_directory.split('/')[-1])
        region_props = regional_prop_dict.get(region_name)
        print(region_props)
        alt_region_paths = region_directories_as_list
        # alt_region_paths.remove(region_path)
        print(alt_region_paths)
        poem_paths = get_document_paths(get_filepaths(region_path))
        shuffle(poem_paths)
        for document_path in poem_paths:
            regions_traded = [] #initialize list of trade partners for naming/citation
            trade_dict = {}   #initialize substitution dictionary
            region_to_trade = "" #initialize variable to store current trade partner
            while not region_to_trade or region_to_trade == region_name: #check that trade partner isn't self
                alt_region_path = sample(alt_region_paths, 1)[0] #define path of alt region
                region_to_trade = os.path.dirname(alt_region_path).split('/')[-1] #exract partner name from path
            print("Current trading partner: " + region_to_trade)
            if region_to_trade not in regions_traded: #store trade partners for citation
                regions_traded.append(region_to_trade)  #store current trade partner
            alt_region_props = regional_prop_dict.get(region_to_trade) #store list of proper nouns to trade for
            # print(alt_region_props)
            os.chdir(input_directory) #change current working directory to wherever the initial materials come from
            filename = os.path.basename(document_path) #store name of file for parsing
            # print(filename)
            if filename.endswith('.pdf') and not os.path.exists(os.path.join(filename.replace('.pdf','.docx'))): #checks whether the file is a word doc or pdf
                text_in_lines = convert_pdf_to_doc(document_path) # converts pdf to word docx and stores all pdf text as a variable
            else:
                text_in_lines = extract_all_text_from_docx(document_path) #extracts all text from a document as a list of strings where each string is a line in the text
            # print(len(text_in_lines))
            new_text_list = [] #initialize text list with replacements to be populated
            # print("going line by line")
            for line in text_in_lines: #go line by line and perform these actions for each
                new_line = [] # initialize variable to store new line with replacements
                if not isinstance(line,str): #checks to see if the line is a run object rather than a string
                    line = line.text  #converts run object to string
                for word in line.split(): #loop over each word in the line
                    # print(word)
                    # print("let's try and substitute some signifiers")
                    if word in region_props: #check if that word is one of this region's proper nouns
                        # print("let's find a substitute for " + word)
                        # print(new_dict.keys())
                        if word not in trade_dict.keys(): #check to see if this prop has been traded before
                            new_word = '' #initialize variable to store prop to trade
                            fail_count = 0 #initialize counter in case we run out of props from one trading partner
                            while not new_word or new_word in list(trade_dict.values()): # loop to collect new word to substitute regional prop, checks that the word's already been used
                                new_word = sample(alt_region_props,1)[0] # grab a prop from the props up for trade
                                # print(new_word)
                                fail_count += 1 # add to fail counter since if we were successful, we wouldn't have made it this far in the loop
                                if fail_count >= 3*len(alt_region_props): # give a breaking point for failing to get a new word for trade
                                    print('not enough proper nouns in this new region to trade') # tell the user that we're gonna have to try something else
                                    alt_region_path = sample(alt_region_paths, 1)[0] # grab a new trading partner
                                    region_to_trade = os.path.dirname(alt_region_path).split('/')[-1] # store new trading partner
                                    while region_to_trade == region_name: # loop to make sure the trading partner isn't the home region, keeps going till it hits something else
                                        alt_region_path = sample(alt_region_paths, 1)[0] # grab a trading partner
                                        region_to_trade = os.path.dirname(alt_region_path).split('/')[-1]  # store region name
                                    alt_region_props = regional_prop_dict.get(region_to_trade) # stores props to trade for from the dictionary of regional props
                                    fail_count = 0 # re-initialize fail_count
                                    # print(alt_region_props)
                                    new_word = sample(alt_region_props,1)[0] #grab a new word to trade from the new trading partner
                                    break

                            trade_dict[word] = new_word
                        else:
                            new_word = trade_dict[word]
                            # print(new_word)
                        new_line.append(new_word)
                        # print('replaced ' + word + ' with ' + new_word)
                    else:
                        new_word = word
                        new_line.append(new_word)
                new_text_list.append(" ".join(new_line))
            trade_partners = "+".join([region_name,"+".join(regions_traded)])
            if len(trade_dict.keys()) < 3:
                print("No trades, let's move on")
                continue
            print(trade_dict)
            output_directory = 'C:/Users/Vvayne/Documents/FA 2017/Major Authors/Outputs/prop trade ' + trade_partners
            if not os.path.exists(output_directory):
                os.makedirs(output_directory)
            clean_and_write_to_doc(new_text_list,output_directory,trade_partners+ '_in_'+ filename.replace('.docx','') + '_' + str(len(trade_dict.keys())) + '_')

def clean_and_write_to_doc(text_list,output_directory,name):

        os.chdir(output_directory)
        f = docx.Document()
        strings_no_numbers = []
        for string in text_list:
            new_string, tag_list = scrape_numbering_tags(string)
            strings_no_numbers.append(new_string)

        if any(line.endswith('.') for line in strings_no_numbers):
            f.add_paragraph(" ".join(text_list))
        else:
            p = f.add_paragraph()
            for i in range(0, len(text_list) - 1):
                p.add_run(text_list[i])
                p.add_run().add_break()
        f.save(name + str(randint(0, 300)) + '.docx')
        print("saved " + name + " at " + output_directory)

def initialize_prop_trade(region_names):
    region_directories = []
    output_directory = 'C:/Users/Vvayne/Documents/FA 2017/Major Authors/Outputs/prop trade ' + '+'.join(region_names)
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    for region_name in region_names:
        input_directory = 'C:/Users/Vvayne/Documents/FA 2017/Major Authors/Material/' + region_name + '/'
        region_directories.append(input_directory)
    return region_directories, output_directory

def create_altered_texts(region_names):
    print("What should we do today?")
    type = str(
        input("type 'l' for local shuffle, 'r' for intraregional, 'g' for global, 's' for signifier trading:    "))
    number = input("also, how many would you like (you can say nothing if you really don't care)? ")
    if not type:
        type = sample(['l', 'r', 'g', 's'], 1)[0]
    if type == 'l':
        region_name = ""
        print("Any region/writer in particular?")
        specific_check = str(input("name or 'no'?"))
        if specific_check.islower() == 'no' or not specific_check:
            print("Fine, I'll choose")
            region_name = sample(region_names, 1)[0]
            print("I choose " + region_name)
        else:
            if specific_check in region_names:
                region_name = specific_check
                print("Okay, let's shuffle " + region_name)
        if not number:
            local_shuffle(region_name)
        elif isinstance(number, int):
            for i in range(0, number):
                local_shuffle(region_name)
        print("All done with your local shuffle")

    elif type == 'r':
        region_name = ""
        print("Any region/writer in particular?")
        specific_check = str(input("name or just hit enter"))
        if not specific_check:
            print("Fine, I'll choose")
            region_name = sample(region_names, 1)[0]
            print("I choose " + region_name)
        if not number:
            construct_intrashuffled_poems(region_name)
        elif isinstance(number, int):
            construct_intrashuffled_poems(region_name, number)
        print("All done with your intraregional mixing")

    elif type == 'g':
        print("Ooh yeah, let's see what it looks like when we assemble a collection of lines from all " + str(
            len(region_names)) + "!")
        input_directory = 'C:/Users/Vvayne/Documents/FA 2017/Major Authors/Material/'
        output_directory = 'C:/Users/Vvayne/Documents/FA 2017/General/Outputs/'
        if isinstance(number, int):
            number_of_poems = number
        else:
            number_of_poems = 1
        for i in range(0, number_of_poems):
            print("let's shuffle up " + str(number_of_poems - i) + " new text(s)")
            construct_global_shuffle(input_directory, output_directory)

    else:
        print("Here we go, let's perform some motif/privilege exchanges")
        region_paths, output_directory = initialize_prop_trade(region_names)
        trade_props(region_paths)

def scrape_and_pull_signifiers(poet,keep=True):
    # scrape_all_poems(poet)
    regional_directories, output_directory = initialize_prop_trade([poet.split()[-1]])
    regional_props_dict = get_regional_props_dict(regional_directories, True)
    # if keep==False:
    #     shutil.rmtree(output_directory)
    return regional_props_dict

def get_unique_paths(list_of_paths):
    paths_in_list = []
    for path in list_of_paths:
        all_paths = get_filepaths(path)
        all_doc_paths = get_document_paths(all_paths)
        for subpath in all_doc_paths:
            if not os.path.basename(subpath) in paths_in_list:
                paths_in_list.append(subpath)
    return paths_in_list

def write_vignettes(paths_to_use, output_path, type='sentence', number_of_vignettes=randint(3, 70)):
    all_sentences = []
    vignette_numbers = sample(range(1, number_of_vignettes + 1), number_of_vignettes)
    sources = {}
    sections_and_sources = {}
    vignettes = {}
    f = docx.Document()
    f.add_heading(str(number_of_vignettes) + " Collage/Assemblage Vignettes")
    print((str(number_of_vignettes) + " Collage/Assemblage Vignettes"))
    for i in vignette_numbers:
        key = str(i)
        print(str(i))
        number_of_sentences = i
        current_sentences = []
        total_extracted = 0
        remaining_sentences = number_of_sentences - total_extracted
        while True:
            if len(current_sentences) >= number_of_sentences:
                print("we've got enough lines for " + str(number_of_sentences))
                break
            sentence = ""
            while not sentence or not any(c.isalpha() for c in sentence):
                try:
                    random_path = sample(paths_to_use, 1)[0]
                    print(random_path)
                    print(remaining_sentences)
                    # try:
                    sentence, page, number_extracted = extract_from_document(random_path, type, remaining_sentences)
                    sentence = sentence + '.'
                except:
                    print("let's try another file")
                    continue

            if sentence:
                # try:
                # print("We're adding the lines to the list for " + key)
                if isinstance(sentence, list):
                    for item in sentence:
                        item = re.sub("\d+", "", item)
                        if '.' not in item:
                            item = item + '.'
                        current_sentences.append(item)
                elif isinstance(sentence, str):
                    list_of_sentences = sentence.split(".")
                    if len(list_of_sentences) > 1:
                        for item in list_of_sentences:
                            item = re.sub("\d+", "", item)
                            if '.' not in item:
                                item = item + '.'
                            current_sentences.append(item)
                    else:
                        sentence = re.sub("\d+", "", sentence)
                        current_sentences.append(sentence)
                # else:
                #     print("We're adding this line")
                #     current_sentences.append(sentence)
                store_source(sources, os.path.basename(random_path), page, sections_and_sources,
                                                  key)
                # print("we've stored the sentence(s) and source(s)")
                # print("total sentence(s) extracted at this stage = " + str(total_extracted))
                remaining_sentences -= total_extracted
                shuffle(current_sentences)
                copy_of_current_sentences = current_sentences
                for sentence in copy_of_current_sentences:
                    if not any(c.isalpha() for c in sentence):
                        # print('already got this line or this line is empty')
                        current_sentences.remove(sentence)

                        # except:
                        #     print("something went wrong with the storage of the lines")

        if not len(current_sentences) >= number_of_sentences:
            continue
        else:
            print(str(len(current_sentences)) + " out of " + str(number_of_sentences) + " sentences being added")
            if key not in vignettes:
                vignettes[key] = [" ".join(current_sentences)]
            else:
                vignettes[key].append([" ".join(current_sentences)])

    for key in vignettes:
        f.add_heading(key)
        f.add_paragraph(vignettes[key])
        p = f.add_paragraph()
        p.add_run().add_break()

    f.add_page_break()
    f.add_heading('Sources', 1)
    # last_p = f.add_paragraph()
    for x in sections_and_sources.keys():
        f.add_heading(x, 2)
        f.add_paragraph(str(sections_and_sources[x]))

        # for x in sources.keys():
        #     f.add_paragraph(str(x) + " " + str(sources[x]))
    # f.write(sources)
    assemblage_name = 'prose_poem' + str(len(vignette_numbers)) + '.docx'
    os.chdir(output_path)
    if os.path.isfile(assemblage_name):
        assemblage_name = 'prose_poem' + str(len(vignette_numbers)) + "_" + str(randint(1, 500)) + '.docx'

    f.save(assemblage_name)

    print("Complete! Check out: " + assemblage_name)


# region_names = ['Adnan','Zurita','Notley']
input_directory = 'C:/Users/Vvayne/Documents/FA 2017/Major Authors/Final Sortie Inputs/'
output_directory = 'C:/Users/Vvayne/Documents/FA 2017/Major Authors/Outputs/Final Sortie/'
if not os.path.exists(output_directory):
    os.mkdir(output_directory)
os.chdir(output_directory)


# input_directory = 'C:/Users/Vvayne/Documents/FA 2017/Major Authors/Material/'
#
# output_directory = 'C:/Users/Vvayne/Documents/FA 2017/Major Authors/Outputs/global shuffles/'
#
# # for region in region_names:
# #     # local_shuffle(region)
# #     construct_intrashuffled_poems(region,5)

# construct_intrashuffled_poems(input_directory,5)
construct_global_shuffle(input_directory,output_directory, english=True)

# poet = ""
# if poet:
#     scrape_and_pull_signifiers(poet)

# create_altered_texts(region_names)

# for region in regional_props_dict.keys():
#     print(region + ' props are:')
#     print(regional_props_dict[region])
# for i in range(0,3):
#     trade_props(regional_directories, regional_props_dict)

