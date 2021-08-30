
import docx
import os
import re
from random import sample, randint, shuffle
from nltk import tokenize, pos_tag, FreqDist
import nltk
import PyPDF2
import string
import shutil
import matplotlib.pyplot as plt
# from poetry_foundation_scraper import scrape_all_poems
english_vocab = set(w.lower() for w in nltk.corpus.words.words())

def primes(n):
    primfac = []
    d = 2
    while d*d <= n:
        while (n % d) == 0:
            primfac.append(d)  # supposing you want multiple factors repeated
            n //= d
        d += 1
    if n > 1:
       primfac.append(n)
    return primfac

def extract_from_pdf(filepath,type='page',subtype='string'):
    # filename = os.path.basename(filepath)
    with open(filepath, 'rb') as pdf_file_obj:  # open current pdf
        try:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)  # store pdf as PyPDF2 reader object
            if pdf_reader and not pdf_reader.isEncrypted:
                # print("I'm looking for a line in " + filename)
                number_of_pages = pdf_reader.getNumPages()  # count number of pages in pdf
                if number_of_pages > 0:
                    if type == 'all' :
                        text_in_lines = []
                        text_in_pages = []
                        for i in range(0,number_of_pages - 1):
                                page_obj = pdf_reader.getPage(i)
                                page_as_list = page_obj.extractText().splitlines()
                                text_in_pages.append(page_as_list)
                                for line in page_as_list:
                                    text_in_lines.append(line)
                        return text_in_pages, text_in_lines
                    elif type == 'page':
                        page_number = randint(0,number_of_pages-1)
                        page_obj = pdf_reader.getPage(page_number)
                        if subtype == 'lines':
                            page_in_lines = page_obj.extractText().splitlines()
                            # text_in_pages = [text_in_lines]
                            return page_in_lines, page_number
                        else:
                            page_as_string = page_obj.extractText()
                            return page_as_string, page_number

                    elif type == 'line':
                        fail_count = 0
                        line = ""
                        page_num = 0
                        while not line and fail_count < number_of_pages:
                            page_num = randint(1, number_of_pages - 1)  # generate random index to choose page
                            page_obj = pdf_reader.getPage(page_num)  # reads and stores random page as page object
                            try:
                                page_in_lines = page_obj.extractText().splitlines()  # stores extracted text as multiline string
                                num_lines = len(page_in_lines)  # count number of lines detected
                                if num_lines > 0:
                                    print("we've got lines here")
                                    line = sample(page_in_lines,1)[0]
                                    return line, page_num
                                else:
                                    print("no lines detect on page " + page_num)
                                    line = ""
                                    fail_count +=1
                                    continue

                            except:
                                print('failed to extract a line from pdf page ' + page_num)
                                fail_count +=1
                                continue
                        return line, page_num

                    elif type == 'sentence':
                        fail_count = 0
                        sentence = ""
                        page_num = 0
                        while not sentence and fail_count < number_of_pages:
                            page_num = randint(0, number_of_pages-1)
                            page_obj = pdf_reader.getPage(page_num)
                            try:
                                page_as_string = page_obj.extractText()
                                page_in_sentences = tokenize.sent_tokenize(page_as_string)
                                sentence = sample(page_in_sentences,1)[0]
                            except:
                                print('failed to extract a sentence from pdf page ' + str(page_num))
                                fail_count +=1
                                continue
                        return sentence, page_num
        except:
            print('unable to parse pdf')

def get_filepaths(directory):
    """
    This function will generate the file names in a directory
    tree by walking the tree either top-down or bottom-up. For each
    directory in the tree rooted at directory top (including top itself),
    it yields a 3-tuple (dirpath, dirnames, filenames).
    """
    file_paths = []  # List which will store all of the full filepaths.

    # Walk the tree.
    for root, directories, files in os.walk(directory):
        for filename in files:
            # Join the two strings in order to form the full filepath.
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)  # Add it to the list.

    return file_paths  # Self-explanatory.

def get_document_paths(file_paths):
    paths_for_potentials = []
    for path in file_paths:
        word = path
        if word.endswith('.docx') or word.endswith('.pdf') or word.endswith('.doc'):
            if not os.path.basename(path).startswith('~$'):
                paths_for_potentials.append(path)
    return paths_for_potentials

def grab_words(path, max_words):
    fail_count = 0
    word_list = []
    print('parsing ' + os.path.basename(path) + ' for words')
    page_text = ""
    if path.endswith('.pdf'):
        while not page_text or len(page_text.split(' ')) < 2:
            # print("scraping a page of text (after " + str(fail_count) + " tries)")
            if fail_count > 20:
                print("couldn't get any text")
                return word_list
            try:
                page_text,page_number = extract_from_pdf(path)

            except:
                fail_count += 1
                # print('failed to get a page (after failing ' + str(fail_count) + ' times')
            fail_count +=1


        # page_as_sentences = split_text([page_text])
    elif path.endswith('.docx'):
        fail_count = 0
        page_text = ""
        print("trying to extract all text from this doc")
        page_text = extract_all_text_from_docx(path)
        if page_text:
            # print(page_text)
            # fail_count +=1
            # if fail_count > 20:
            #     print("can't parse this doc")
            #     return word_list
            if not isinstance(page_text[0],str):
                print('converting runs to strings')
                all_lines_as_list = []
                for run in page_text:
                    all_lines_as_list.append(run.text)
                page_text = " ".join(all_lines_as_list)

            if not isinstance(page_text, str):
                print('this page text is in the wrong form')
                try:
                    page_text = " ".join(page_text)
                except:
                    print('no luck fixing it')
                    return word_list
        else:
            print('no text here')
            return word_list
    number_of_words_available = len(tokenize.word_tokenize(page_text))
    if not number_of_words_available:
        return word_list
    print("found " + str(number_of_words_available) + " words in this document")
    if number_of_words_available < 5:
        word_list = sample(tokenize.word_tokenize(page_text),len(tokenize.word_tokenize(page_text)))
        return word_list

    number_of_words = randint(2, max(2,min(max_words,number_of_words_available)))
    print("let's collect " + str(number_of_words) + " from this doc")
    page_in_words = tokenize.word_tokenize(page_text)
    number_of_available_words = len(page_in_words)
    start = randint(0, number_of_available_words - number_of_words)
    interval = range(start, start + number_of_words)
    for i in interval:

        word_list.append(page_in_words[i])
    # print(word_list)
    print('added ' + str(len(interval)) + ' words to list')
    return word_list

def extract_all_text_from_docx(filepath, paragraph_or_runs='runs'):
    try:
        doc = docx.Document(filepath)
    except:
        print("couldn't scrape it")
        all_runs = []
        return all_runs
    if doc:
        print("attempting to scrape " + os.path.basename(filepath))
        file_fail_count = 0
        all_paragraphs = []
        all_runs = []
        length = len(doc.paragraphs)
        # print(str(length))
        if length >= 1:
            for paragraph in doc.paragraphs:
                all_paragraphs.append(paragraph)
                for run in paragraph.runs:
                    all_runs.append(run.text)
            if paragraph_or_runs == 'paragraphs':
                return all_paragraphs
            else:
                print("got all " + str(length) + " run(s)")
                # print("sample: " + all_runs[0])
                return " ".join(tokenize.word_tokenize(" ".join(all_runs)))

def split_text(string_list):
    print("let's see if we need to split this up into sentences")
    list_of_strings = []
    for run in string_list:
        list_of_strings.append(run)

    if '.' in [run for run in list_of_strings]:
        text_list = tokenize.sent_tokenize(" ".join(list_of_strings))
    else:
        text_list = list_of_strings

    return text_list

def refrain_or_no(all_lines_so_far): # function to determine if the next line should be a refrain and if so what should it be
    if len(all_lines_so_far) > 2: #test for at least two lines before we start getting crazy with repetition
        source_for_lines = all_lines_so_far #store a copy of the lines
        roll_of_die = randint(1,len(source_for_lines)) #generate random integer between 1 and the number of available lines for refrain
        if roll_of_die >= randint(7,9)*.1*len(source_for_lines): #probabilistic test to confirm refrain
            line = sample(all_lines_so_far, 1)[0] #store refrain as line
            return 'yes', line  #return string and line
        else: #if refrain 'failed'
            line = "" #store nothing as line
            return 'no', line #return string and empty line
    else: # if 2 or fewer lines
        line = "" #store nothing as line
        return 'no', line #return no and empty line

def scrape_numbering_tags(string):
    # print("let's test for numbering/sectional tags so they don't get lumped in as lines")
    s = string
    if not isinstance(string,str):
        s = string.text
    if re.search('[iv]+\.',s):
        lst = re.findall('[iv]+\.', s)
        for tag in lst:
            s.replace(tag,"")
    elif re.search('\d+\.',s):
        lst = re.findall('\d+\.', s)
        for tag in lst:
            s.replace(tag,"")
        # s.replace([i for i in lst],"")
    else:
        lst = []
    return s, lst

def get_number_of_lines(path):
    all_text_list = extract_all_text_from_docx(path)
    if all_text_list:
        all_numbers = []
        strings_no_numbers = []
        for string in all_text_list:
            new_string, tag_list = scrape_numbering_tags(string)
            strings_no_numbers.append(new_string)
            for tag in tag_list:
                all_numbers.append(tag)
        for line in strings_no_numbers:
            if line.isspace():
                strings_no_numbers.remove(line)
        return len(strings_no_numbers)
    else:
        return 0

def get_regional_paths(region_directory):
    all_paths = []
    region = os.path.dirname(region_directory).split('/')[-1]
    intraregional_text_paths = get_document_paths(get_filepaths(region_directory))
    for itp in intraregional_text_paths:
        if not itp in all_paths:
            all_paths.append(itp)

    return all_paths

    def line_from_doc(filepath):
        doc = docx.Document(filepath)
        if doc:
            print("attempting to scrape " + os.path.basename(filepath))
            file_fail_count = 0
            line = ""
            length = len(doc.paragraphs)
            if length >= 1:
                while not line:
                    if file_fail_count >= 20:
                        print("Out of 20 tries, could not get a line")
                        break

                    paragraph_no = randint(0, length - 1)
                    # number_of_lines = len(doc.paragraphs[paragraph_no].text)
                    # paragraph_no = randint(0,number_of_lines-1)
                    potential_line = doc.paragraphs[paragraph_no].text
                    if potential_line:
                        potential_line_as_string = potential_line  # stores selected line for filtering
                        print("Paragraph Scraped")
                        # print(potential_line_as_string)
                        # filter punctuation and digits
                        for c in potential_line_as_string:  # loop through each character in line
                            if not c.isalpha() and not c.isspace() or c in "ÒÓÔÕßˆˇ":  # test for unsuitable characters
                                potential_line_as_string = potential_line_as_string.replace(c,
                                                                                            "")  # replace unsuitable character with ""
                        words_in_line = tokenize.word_tokenize(potential_line_as_string)  # separate line into recognizable words
                        # print("Paragraph filtered")
                        words_in_paragraph_count = len(words_in_line)
                        if words_in_line and words_in_paragraph_count >= 2:
                            # print("There are " + str(words_in_paragraph_count) + " words left in the filtered section.")
                            if words_in_paragraph_count >= 20:
                                number_of_words = min(randint(2, words_in_paragraph_count), 15)
                                # print("Let's cut that down to " + str(number_of_words) + " words")
                                beginning_of_line = randint(0, words_in_paragraph_count - number_of_words)
                                end_of_line = beginning_of_line + number_of_words - 1
                                selected_words = words_in_line[beginning_of_line:end_of_line]
                                line = " ".join(selected_words)
                                return line, paragraph_no
                            else:
                                line = " ".join(tokenize.word_tokenize(potential_line_as_string))
                                return line, paragraph_no
                            print(line)
                        else:
                            file_fail_count += 1
                    else:
                        file_fail_count += 1
                        # except:
                        #     line = ""
                        #     line_no = ""
                        #     print("let's try a different document")
                        #     return line, line_no
                        # try:
                        #     doc = open(filepath, 'r')
                        #     page = [line.split(" ") for line in doc.readlines()]
                        #     length = len(page)
                        #     line_no = randint(2,length)
                        #     # number_of_words_in_line = randint(2, length)
                        #     # beginning_of_line = randint(0, length - number_of_words_in_line - 1)
                        #     # end_of_line = beginning_of_line + number_of_words_in_line
                        #     # line_words = page[beginning_of_line:end_of_line]
                        #     line_words = page[line_no]
                        #     line = " ".join(line_words)
                        #     return line, line_no
                        # except:
                        #         line = ""
                        #         line_no = ""
                        #         print("nothing doing with this file")
                        #         print("nothing doing with this file")
                        #         shutil.move(filepath, destination)
                        #         return line, line_no


def line_from_pdf(filepath):
    filename = os.path.basename(filepath)
    with open(filepath, 'rb') as pdf_file_obj:  # open current pdf
        pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)  # store pdf as PyPDF2 reader object
        if pdf_reader and not pdf_reader.isEncrypted:
            print("I'm looking for a line in " + filename)
            number_of_pages = pdf_reader.getNumPages()  # count number of pages in pdf
            if number_of_pages > 2:
                file_fail_count = 0
                line = ""
                page_num = ""
                while not line:
                    if file_fail_count >= min(20, number_of_pages + 4):
                        print("Out of " + str(file_fail_count) + " pages, could not get a line. :-/")
                        break
                    page_num = randint(1, number_of_pages - 1)  # generate random index to choose page
                    page_obj = pdf_reader.getPage(page_num)  # reads and stores random page as page object
                    try:
                        page_in_lines = page_obj.extractText().splitlines()  # stores extracted text as multiline string
                        num_lines = len(page_in_lines)  # count number of lines detected
                        if num_lines > 3:  # ensure page without blank lines has been chosen
                            line_num = randint(0, num_lines - 1)  # store random number to select line from page
                        else:
                            file_fail_count += 1
                            continue  # if page has less than 3 non-blank lines, continues to next page

                        potential_line = page_in_lines[line_num]  # stores selected line for filtering
                        # filter punctuation and digits
                        for c in potential_line:  # loop through each character in line
                            if not c.isalpha() and not c.isspace() or c in "ÒÓÔÕßˆˇ":  # test for unsuitable characters
                                potential_line = potential_line.replace(c, "")  # replace unsuitable character with ""
                        words_in_line = tokenize.word_tokenize(potential_line)  # separate line into recognizable words

                        if words_in_line and len(words_in_line) > 1:
                            print(words_in_line)
                            # filter word abberrations
                            line_words = words_in_line
                            line = " ".join(line_words)  # join words back into line with blank space between
                            return line, page_num
                        else:
                            file_fail_count += 1
                    except:
                        file_fail_count += 1
                        continue


def line_from_document(filepath):
    if os.path.basename(filepath).endswith('.docx'):
        line, section_num = line_from_doc(filepath)
        return line, section_num
    elif os.path.basename(filepath).endswith('.pdf'):
        line, section_num = line_from_pdf(filepath)
        return line, section_num

def text_by_keywords(keywords, number_of_words= randint(60,500),input_directory="C:/Users/Vvayne/"):
    region_directories, keywords_list, output_directory = initialize_keyword_directories(input_directory,keywords)
    input_paths = []
    for directory in region_directories:
        reg_paths = get_regional_paths(directory)
        for path in reg_paths:
            if not path in input_paths:
                input_paths.append(path)
    print('got all the paths')
    os.chdir(output_directory) #print change current directory to output directory
    print(output_directory)
    all_lines_so_far = assemble_phrases(input_paths,output_directory,number_of_words)
    return all_lines_so_far

def assemble_phrases(input_paths,output_directory,number_of_words=randint(60,500)):
    all_words_so_far = []
    all_lines_so_far = []
    f = docx.Document() # initialize new document to write and save construction
    line_count = 0 # initialize line counter
    word_count = 0
    number_of_lines = randint(10,2*max(primes(number_of_words))+10) #define number of lines in new construction between 5 and 40 and ideally equal to the length of a random text in the materials folder
    print("let's get " + str(number_of_words) + " words into " + str(number_of_lines) + ' lines')
    sources = [] #initialize list to store the origin of each line
    while len(all_words_so_far) < number_of_words:
        remaining_words = number_of_words - len(all_words_so_far)
        print('still need ' + str(remaining_words) + ' words')
        if remaining_words <= 0:
            break
        current_path = sample(input_paths,1)[0]
        max_number_of_new_words = 0
        while max_number_of_new_words > .1*number_of_words or max_number_of_new_words == 0:
            max_number_of_new_words = min([randint(2,round(.1*number_of_words)),remaining_words])
        new_words_list = grab_words(current_path, max_number_of_new_words)
        if new_words_list:
            for word in new_words_list:
                if word and word[0].isalpha() and word in english_vocab:
                    all_words_so_far.append(word)
            print('added new word list to big word list')
            sources.append(os.path.basename(current_path))
            print('stored the source of the words too')
    print('[COLLECTED ' + str(len(all_words_so_far)) + ' WORDS]')
    print("let's make lines out of them")
    for line_number in range(1,number_of_lines):
        new_line_string = ""
        # refrain_boolean, refrain = refrain_or_no(
        #     all_lines_so_far)  # random integer to decide if a line should be repeated and returns it if yes
        # # print(refrain_boolean)
        # if refrain_boolean == 'yes':  # test for refrain_affirmation
        #     new_line = refrain  # stores old line as new line for refrain
        #     print('refrained')
        line_length = max(14, sample(primes(number_of_lines), 1)[0])
        words_in_chunks = partition_text_in_words(all_words_so_far,randint(3,7))
        while not new_line_string or len(new_line) < line_length:
            # start = randint(0,len(all_words_so_far) - line_length - 1)
            # end = start + line_length
            # new_line = " ".join(all_words_so_far[start:end])
            new_line = []
            while len(new_line) < line_length:
                chunk = sample(words_in_chunks,1)[0]
                while not isinstance(chunk,str):
                    chunk = " ".join(chunk)
                new_line.append(chunk + " ")
            new_line_string = " ".join(new_line)
        all_lines_so_far.append(new_line)
    shuffle(all_lines_so_far)
    for line in all_lines_so_far:
        f.add_paragraph(line)
        # print(line)

    f.add_page_break() #starts a new page in document being constructed
    f.add_heading('Sources', 1) #write header called Sources on new document
    last_p = f.add_paragraph() #start final block
    set_of_sources = [] #initialize unique set of sources
    for source in sources: #loop through list of all sources for lines in the new document
        if not source in set_of_sources: #check for uniqueness
            set_of_sources.append(source) #add if unique
    for source in set_of_sources: #loop through unique sources
        last_p.add_run(source) #write source
        last_p.add_run().add_break() #give space before the next one
    name = 'poem_by_phrases_' + str(len(set_of_sources)) + '_' + str(number_of_words) + '.docx' #generate a name using the number of different poems used and a random integer
    f.save(name) # save new document
    print("Wrote a thing called: " + name)
    print("It's got " + str(len(all_lines_so_far)) + " lines from " + str(len(set_of_sources)) + " different sources")
    print("You can find it in " + output_directory)
    return all_lines_so_far

def partition_text_in_words(text_in_words,chunk_length=randint(2,6)):
    l = text_in_words
    n = chunk_length
    """Yield successive n-sized chunks from l."""
    chunk_list = []
    for i in range(0, len(l), n):
        chunk_list.append(l[i:i + n])
    return chunk_list

# def phrase_by_subdirectory(region_dir_path):

def get_all_subdirectories(paths):
    dir_list = []
    for path in paths:
        sub_dir = os.path.dirname(path)
        dir_list.append(sub_dir)
            # print(path)
            # print(sub_dir)

    return dir_list

def initialize_keyword_directories(input_directory,keywords_list=None):
    print("let's get the directories we need")
    all_possible_paths = get_document_paths(get_filepaths(input_directory))
    print("found " + str(len(all_possible_paths)) + " total document paths")
    dir_list = get_all_subdirectories(all_possible_paths)
    print("including " + str(len(dir_list)) + " directories")
    if not keywords_list:
        print("let's choose some key directories")
        number_of_domains = randint(1,7)
        keyword_directories = sample(dir_list,number_of_domains)
        keywords_list = [re.split(r'[/\\]',sub_dir)[-1] for sub_dir in keyword_directories]
        print(keywords_list)

    region_directories = []
    relevant_paths = []
    for name in keywords_list:
        print("let's find any files related to " + name)
        for path in all_possible_paths:
            if re.search(name,path):
                sub_dir = os.path.dirname(path)
                if sub_dir not in region_directories:
                    print('got a sub-directory')
                    region_directories.append(sub_dir)
                if path not in relevant_paths:
                    relevant_paths.append(path)

    output_directory = 'C:/Users/Vvayne/Documents/FA 2017/Outputs/assemblages from ' + '+'.join(keywords_list)
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    # unique_relevant_paths = []
    # for path in relevant_paths:
    #     if not path in unique_relevant_paths:
    #         unique_relevant_paths.append(path)
    return region_directories, keywords_list,output_directory

def preformed_or_scrambled():
    coin_flip = randint(0,1)
    if coin_flip == 0:
        mode = 'preformed'
    else:
        mode = 'scrambled'
    return mode

def add_line_of_verse(f,max_lines,first_line,line_counter,stanza_break,line):
    try:
        p = f.add_paragraph()
        if line_counter > 0 and line_counter % stanza_break == 0:
            run = p.add_run()
            run.add_break()
            run.add_break()
            max_lines += 1
            line_counter += 1  # increases line counter
            print("Stanza break on line " + str(line_counter))
        if line_counter == 0 and not first_line:
            first_line = line
            print("First line selected!")
            line_counter+=1
            p.add_run(line)  # appends line to poem file and goes to next line
            return first_line,line_counter
        line_counter += 1
        # print("Line number " + str(line_counter) + " selected!")
        p.add_run(line)  # appends line to poem file and goes to next line
        return first_line,line_counter
    except:
        print("Failed to add line as verse")

def add_scramble_line(all_lines_segmented,first_line,line_counter,line,f):
    if line_counter == 0 and not first_line:
        first_line = line
        print("First line selected!")
        line_counter += 1
        p = f.add_paragraph()
        p.add_run(line)  # appends line to poem file and goes to next line
        return first_line,line_counter

    # print("Line " + str(line_counter) + " will be: " + line)  # shows recently generated line
    # print(segment(line))
    all_lines_segmented.append(tokenize.word_tokenize(line))
    line_counter += 1
    return first_line, line_counter
    # print(all_lines_segmented)

def store_source(sources,filename,section_num):
    if filename.endswith('.pdf'):
        if not filename in sources:
            sources[filename] = {'page or section number': section_num}  # store source of current line being generated
    elif filename.endswith('.doc') or filename.endswith('.docx'):
        if not filename in sources:
            sources[filename] = {'page or section number': section_num}  # store source of current line being generated

def poem_writer(input_directory):
    file_list = get_document_paths(input_directory)
    line_counter = 0
    fail_count = 0
    sources = {}
    first_line = ""
    mode = ''
    fl = ""

    # number_of_poems_to_write = 5

    f = docx.Document()
    if not mode:
        mode = preformed_or_scrambled()

    if mode == 'preformed':
        stanza_count = 0
        max_lines = randint(10, 30)
        stanza_break = randint(2, max_lines)
        # loop to continue running until poem length reaches max_lines
        print(stanza_break)
        print("Let's write a poem of " + str(max_lines) + " lines with a stanza break every " + str(
            stanza_break) + " lines")
    else:
        all_lines_segmented = []
        max_lines = randint(20, 60)
        print("Let's write a poem from a block with " + str(max_lines) + " lines from various sources!")

    # loop to continue running until poem length reaches max_lines
    while line_counter < max_lines:
        remaining_lines = max_lines - line_counter  # calculate remaining number of lines to grab
        for docpath in file_list:
            if fl and mode == 'preformed':
                line = refrain_or_no('preformed', f)
                add_line_of_verse(f, max_lines, first_line, line_counter, stanza_break, line)
            elif fl and mode == 'scrambled':
                while not line:
                    line = refrain_or_no('scrambled', all_lines_segmented)
                add_scramble_line(all_lines_segmented, first_line, line_counter, line)
            file_path = docpath
            filename = os.path.basename(file_path)
            if filename:
                line = ""
                section_num = 0
                # try:
                if filename.endswith('.pdf'):
                    # print(filename)
                    try:
                        line, section_num = line_from_pdf(file_path)
                    except:
                        print("unreadable pdf")
                        # try:
                        destination = 'C:/Users/wayne_000/Desktop/problem_files/problem_pdfs/'
                        shutil.move(file_path, destination)
                        # except:
                        #     print("already moved or couldn't move it")
                elif filename.endswith('.doc') or filename.endswith('.docx'):
                    # print("scraping " + filename + " at " + file_path + " for a line")
                    try:
                        line, section_num = line_from_doc(file_path)
                    except:
                        print("not docx friendly")
                        try:
                            destination = 'C:/Users/wayne_000/Desktop/problem_files/problem_docs/'
                            shutil.move(file_path, destination)
                        except:
                            print("already moved or couldn't move it")
                if line:
                    print("let's try writing this line")
                    if line_counter == 0:
                        first_line = line
                    store_source(sources,filename, section_num)
                    if mode == 'preformed':
                        fl, line_counter = add_line_of_verse(f, max_lines, first_line, line_counter, stanza_break, line)
                        if not first_line or first_line.isspace():
                            first_line = fl
                        continue
                    else:
                        fl, line_counter = add_scramble_line(all_lines_segmented, first_line, line_counter, line, f)
                        if not first_line or first_line.isspace():
                            first_line = fl
                        continue

                        # except:
                        #      print("Could not get a line from " + filename)
                        #      continue
        else:
            continue
    title_words = tokenize.word_tokenize(fl)
    shuffle(title_words)
    title = " ".join(title_words)

    print("Got the lines, now let's save 'em!")
    if mode == "preformed":
        f.add_page_break()
        f.add_heading('Sources', 1)
        last_p = f.add_paragraph()
        for x in sources.keys():
            f.add_paragraph(str(x) + " " + str(sources[x]))
        poem_name = ""
        try:
            poem_name = " ".join(title_words[0:min(len(title_words), 8)]) + '.docx'
            f.save(poem_name)
        except:
            poem_name = 'newpoem' + str(randint(0, 100000)) + '.docx'
            f.save(poem_name)

        print("Poem complete! Check out: " + poem_name)
    else:
        all_words = tokenize.word_tokenize(" ".join(str(x) for x in all_lines_segmented))
        f.add_paragraph(" ".join(y for y in all_words))
        # shuffle(all_words)
        # f.add_page_break()
        # f.add_paragraph(" ".join(y for y in all_words))

        f.add_page_break()
        f.add_heading('Sources', 1)
        last_p = f.add_paragraph()
        for x in sources.keys():
            f.add_paragraph(str(x) + " " + str(sources[x]))
        # f.write(sources)
        poem_name = ""
        try:
            poem_name = first_line[0:min(len(first_line), 30)] + '.docx'
            f.save(poem_name)
        except:
            poem_name = 'newpoem' + str(randint(0, 100000)) + '.docx'
            f.save(poem_name)

        print("Poem complete! Check out: " + poem_name)


# input_directory = 'C:/Users/Vvayne/Downloads/Downloads/'
# input_directory = 'C:/Users/Vvayne/Documents/FA 2017/General/Inputs/'
# output_directory = "C:/Users/Vvayne/Documents/FA 2017/General/Outputs/"
fiction_directory = "C:/Users/Vvayne/Documents/FA 2017/Inputs/other/literature/fiction"
input_directory = 'C:/Users/Vvayne/Documents/FA 2018/'

output_directory = 'C:/Users/Vvayne/Documents/FA 2018/Outputs/' + input_directory.split('/')[-2]
if not os.path.exists(output_directory):
    os.mkdir(output_directory)
os.chdir(output_directory)
# print(poem_as_list)
# print(poem_as_list)
# keywords = ['mechanics']
# relevant_directories,keywords,output_directory = initialize_keyword_directories(input_directory,keywords)
# print(keywords)
# print(relevant_directories)
#
# relevant_paths = []
# for directory in relevant_directories:
#     paths = get_document_paths(get_filepaths(directory))
#     for path in paths:
#         if not os.path.basename(path) in relevant_paths:
#             relevant_paths.append(path)
#
# poem_as_list = text_by_keywords(keywords)

poem_as_list = assemble_phrases(get_filepaths(input_directory),output_directory)
# phrases_by_specific_paths(relevant_paths, output_directory)




def line_from_doc(filepath):
        doc = docx.Document(filepath)
        if doc:
            print("attempting to scrape " + os.path.basename(filepath))
            file_fail_count = 0
            line = ""
            length = len(doc.paragraphs)
            if length >= 1:
                while not line:
                    if file_fail_count >= 20:
                        print("Out of 20 tries, could not get a line")
                        break

                    paragraph_no = randint(0, length - 1)
                    # number_of_lines = len(doc.paragraphs[paragraph_no].text)
                    # paragraph_no = randint(0,number_of_lines-1)
                    potential_line = doc.paragraphs[paragraph_no].text
                    if potential_line:
                        potential_line_as_string = potential_line  # stores selected line for filtering
                        print("Paragraph Scraped")
                        # print(potential_line_as_string)
                        # filter punctuation and digits
                        for c in potential_line_as_string:  # loop through each character in line
                            if not c.isalpha() and not c.isspace() or c in "ÒÓÔÕßˆˇ":  # test for unsuitable characters
                                potential_line_as_string = potential_line_as_string.replace(c,
                                                                                            "")  # replace unsuitable character with ""
                        words_in_line = tokenize.word_tokenize(potential_line_as_string)  # separate line into recognizable words
                        # print("Paragraph filtered")
                        words_in_paragraph_count = len(words_in_line)
                        if words_in_line and words_in_paragraph_count >= 2:
                            # print("There are " + str(words_in_paragraph_count) + " words left in the filtered section.")
                            if words_in_paragraph_count >= 20:
                                number_of_words = min(randint(2, words_in_paragraph_count),15)
                                # print("Let's cut that down to " + str(number_of_words) + " words")
                                beginning_of_line = randint(0, words_in_paragraph_count - number_of_words)
                                end_of_line = beginning_of_line + number_of_words - 1
                                selected_words = words_in_line[beginning_of_line:end_of_line]
                                line = " ".join(selected_words)
                                return line, paragraph_no
                            else:
                                line = " ".join(tokenize.word_tokenize(potential_line_as_string))
                                return line, paragraph_no
                        else:
                            file_fail_count +=1
                    else:
                        file_fail_count += 1

