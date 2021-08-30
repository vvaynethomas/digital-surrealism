import re
import os
from random import randint
from random import sample
from random import shuffle
import nltk
import docx
from PIL import Image
# import pyocr
# import pyocr.builders
import PyPDF2

def pdf_to_image(pdf_path):
    input1 = PyPDF2.PdfFileReader(open(pdf_path, "rb"))
    page0 = input1.getPage(0)
    xObject = page0['/Resources']['/XObject'].getObject()

    for obj in xObject:
        if xObject[obj]['/Subtype'] == '/Image':
            size = (xObject[obj]['/Width'], xObject[obj]['/Height'])
            data = xObject[obj].getData()
            if xObject[obj]['/ColorSpace'] == '/DeviceRGB':
                mode = "RGB"
            else:
                mode = "P"

            if xObject[obj]['/Filter'] == '/FlateDecode':
                img = Image.frombytes(mode, size, data)
                img.save(obj[1:] + ".png")
            elif xObject[obj]['/Filter'] == '/DCTDecode':
                img = open(obj[1:] + ".jpg", "wb")
                img.write(data)
                img.close()
            elif xObject[obj]['/Filter'] == '/JPXDecode':
                img = open(obj[1:] + ".jp2", "wb")
                img.write(data)
                img.close()
        return img

# def text_from_image(image_path):
#     tools = pyocr.get_available_tools()[0]
#     line_and_word_boxes = tools.image_to_string(
#         Image.open(image_path), lang="eng",
#         builder=pyocr.builders.TextBuilder())
#     return line_and_word_boxes

def get_filepaths(directory):
    """
    This function will generate the file names in a directory
    tree by walking the tree either top-down or bottom-up. For each
    directory in the tree rooted at directory top (including top itself),
    it yields a 3-tuple (dirpath, dirnames, filenames).
    """
    file_paths = []  # List which will store all of the full filepaths.

    # Walk the tree.
    for root, directories, files in os.walk(directory):
        for filename in files:
            # Join the two strings in order to form the full filepath.
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)  # Add it to the list.

    return file_paths  # Self-explanatory.

def grab_a_file_to_scrape(selected_path):
    number_of_files = len(selected_path)
    grabbed_file = ""
    while not grabbed_file:
        proposed_file_no = randint(0,number_of_files-1)
        proposed_file_path = selected_path[proposed_file_no]
        proposed_filename = os.path.basename(proposed_file_path)
        if proposed_filename.endswith('pdf') or proposed_filename.endswith('.doc' or 'docx'):
            grabbed_file = proposed_filename
            return grabbed_file, proposed_file_path

def get_document_paths(file_paths):
    paths_for_potentials = []
    for path in file_paths:
        word = path
        if word.endswith('.docx') or word.endswith('.pdf') or word.endswith('.doc'):
            paths_for_potentials.append(path)
    return paths_for_potentials

def get_pdf_paths(file_paths):
    paths_for_pdfs = []
    for path in file_paths:
        word = path
        regexp = re.compile(r'pdf')
        if regexp.search(word):
            paths_for_pdfs.append(path)
    return paths_for_pdfs

def get_docx_paths(file_paths):
    paths_for_docs = []
    for path in file_paths:
        word = path
        regexp = re.compile(r'doc')
        if regexp.search(word):
            paths_for_docs.append(path)
    return paths_for_docs

def extract_from_doc(filepath,type,max_lines):
        try:
            doc = docx.Document(filepath)
        except:
            print("couldn't make this one work")
            line = ""
            page_num = 0
            number_extracted = 0
            return line, page_num, number_extracted
        if doc:
            # print("attempting to scrape " + os.path.basename(filepath))
            file_fail_count = 0
            line = ""
            length = len(doc.paragraphs)
            if length >= 1:
                while not line:
                    if file_fail_count >= 20:
                        # print("Out of 20 tries, could not get a line")
                        break

                    paragraph_no = randint(0, length - 1)
                    # number_of_lines = len(doc.paragraphs[paragraph_no].text)
                    # paragraph_no = randint(0,number_of_lines-1)
                    potential_line = doc.paragraphs[paragraph_no].text
                    if potential_line:
                        if type == "line":
                            potential_line_as_string = potential_line  # stores selected line for filtering
                            # print("Paragraph Scraped")
                            # print(potential_line_as_string)
                            # filter punctuation and digits
                            for c in potential_line_as_string:  # loop through each character in line
                                if not c.isalpha() and not c.isspace() or c in "ÒÓÔÕßˆˇ":  # test for unsuitable characters
                                    potential_line_as_string = potential_line_as_string.replace(c,
                                                                                                "")  # replace unsuitable character with ""
                            words_in_line = nltk.tokenize.word_tokenize(potential_line_as_string)  # separate line into recognizable words
                            # print("Paragraph filtered")
                            words_in_paragraph_count = len(words_in_line)
                            if words_in_line and words_in_paragraph_count >= 2:
                                # print("There are " + str(words_in_paragraph_count) + " words left in the filtered section.")
                                if words_in_paragraph_count >= 20:
                                    number_of_words = min(randint(2, words_in_paragraph_count),15)
                                    # print("Let's cut that down to " + str(number_of_words) + " words")
                                    beginning_of_line = randint(0, words_in_paragraph_count - number_of_words)
                                    end_of_line = beginning_of_line + number_of_words - 1
                                    selected_words = words_in_line[beginning_of_line:end_of_line]
                                    line = " ".join(selected_words)
                                    number_extracted = 1
                                    return line, paragraph_no, number_extracted
                                else:
                                    line = " ".join(nltk.tokenize.word_tokenize(potential_line_as_string))
                                    number_extracted = 1
                                    return line, paragraph_no, number_extracted

                            else:
                                file_fail_count +=1
                        elif type == "sentence":
                            # try:
                                paragraph_as_string = potential_line
                                paragraph_as_list = paragraph_as_string.split(".")
                                number_of_sentences = len(paragraph_as_list)
                                # print("Total sentences in paragraph = " + str(number_of_sentences))
                                try:
                                    sentences = sample(paragraph_as_list, (min(max_lines,randint(1,3))))
                                    copy_of_sentences = sentences
                                    for sentence in copy_of_sentences:
                                        if sentence.isspace() or not any(c.isalpha() for c in sentence):
                                            sentences.remove(sentence)
                                    number_keeping = len(sentences)
                                    shuffle(sentences)
                                    joined_sentences = ". ".join(sentences)
                                    # print("Let's keep " + str(number_keeping) + " of them")
                                    return joined_sentences, paragraph_no, number_keeping
                                except:
                                    try:
                                        sentence = paragraph_as_list[randint(0,number_of_sentences)]
                                        number_keeping = 1
                                        if not "." in sentence:
                                            sentence = sentence + ". "
                                        return sentence, paragraph_no, number_keeping
                                    except:
                                        print("still not getting there")
                                # if number_of_sentences > 1:
                                #     number_keeping = min(randint(1,number_of_sentences),max_lines)
                                #     print("Let's keep " + str(number_keeping) + " of them")
                                #     starting_number = randint(0,number_of_sentences - number_keeping - 1)
                                #     ending_number = min(starting_number + number_keeping - 1,starting_number + i -1)
                                #     if ending_number > starting_number:
                                #         print("How about the " + str(ending_number - starting_number) + " between sentences " + str(starting_number) + " and " + str(ending_number) + ".")
                                #         sentences = ". ".join(paragraph_as_list[starting_number:ending_number])
                                #         return sentences, paragraph_no, number_keeping
                                #     elif ending_number == starting_number:
                                #         sentences = paragraph_as_list[starting_number]
                                #         return sentences, paragraph_no, number_keeping
                                # else:
                                #     sentences = paragraph_as_list[0]
                                #     number_keeping = 1
                                #     return sentences, paragraph_no, number_keeping
                            # except:
                            #     print("couldn't close the deal")
                            #     file_fail_count+=1

                    else:
                        file_fail_count += 1
        else:
            line = ""
            section_num = 0
            number_extracted = 0
            return line, section_num, number_extracted

def extract_from_pdf(filepath,type,max_lines):
    filename = os.path.basename(filepath)
    with open(filepath, 'rb') as pdf_file_obj:   # open current pdf
        pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)  # store pdf as PyPDF2 reader object
        if pdf_reader and not pdf_reader.isEncrypted:
            # print("I'm looking for a line in " + filename)
            number_of_pages = pdf_reader.getNumPages()  # count number of pages in pdf
            if number_of_pages > 0:
                file_fail_count = 0
                line = ""
                page_num = ""
                while not line:
                    if file_fail_count >= min(20, number_of_pages+4):
                        print("Out of " + str(file_fail_count) + " pages, could not get a line. :-/")
                        break
                    page_num = randint(1, number_of_pages - 1)  # generate random index to choose page
                    page_obj = pdf_reader.getPage(page_num)  # reads and stores random page as page object
                    if type == "line":
                        try:
                            page_in_lines = page_obj.extractText().splitlines()  # stores extracted text as multiline string
                            num_lines = len(page_in_lines)  # count number of lines detected
                            if num_lines > 3:  # ensure page without blank lines has been chosen
                                line_num = randint(0, num_lines - 1)  # store random number to select line from page
                            else:
                                file_fail_count += 1
                                continue  # if page has less than 3 non-blank lines, continues to next page

                            potential_line = page_in_lines[line_num]  # stores selected line for filtering
                            # filter punctuation and digits
                            for c in potential_line:  # loop through each character in line
                                if not c.isalpha() and not c.isspace() or c in "ÒÓÔÕßˆˇ" :  # test for unsuitable characters
                                    potential_line = potential_line.replace(c, "")  # replace unsuitable character with ""
                            words_in_line = nltk.tokenize.word_tokenize(potential_line)  # separate line into recognizable words

                            if words_in_line and len(words_in_line) > 1:
                                # print(words_in_line)
                                # filter word abberrations
                                line_words = words_in_line
                                line = " ".join(line_words)  # join words back into line with blank space between
                                return line, page_num
                            else:
                                file_fail_count += 1
                        except:
                            file_fail_count+=1
                            continue
                    elif type == "sentence":
                        try:
                            page_as_string = page_obj.extractText()
                            page_as_sentences = nltk.tokenize.sent_tokenize(page_as_string)
                            num_sentences = len(page_as_sentences)
                            # print("We've detected " + str(num_sentences) + " sentences in this paragraph.")
                            if num_sentences > 1 and max_lines >= 1:
                                # try:
                                    potential_sentences = sample(page_as_sentences,min(randint(1,3),max_lines))
                                    number_extracted = len(potential_sentences)
                                    # print("Let's keep " + str(number_extracted) + " of them.")
                                    copy_of_potential_sentences = potential_sentences
                                    for sentence in copy_of_potential_sentences:
                                        if sentence.isspace() or not any(c.isalpha() for c in sentence):
                                            potential_sentences.remove(sentence)
                                    shuffle(potential_sentences)
                                    extracted_sentences = ".".join(potential_sentences)
                                    # for c in extracted_sentences:  # loop through each character in line
                                    #     if not c.isalpha() and not c.isspace() or c in "ÒÓÔÕßˆˇ" and not c in ".,!?:\"();\'":  # test for unsuitable characters
                                    #         extracted_sentences = extracted_sentences.replace(c, "")  # replace unsuitable character with ""
                                    return extracted_sentences, page_num, number_extracted
                                # except:
                            else:
                                    potential_sentence = page_as_sentences[randint(0,num_sentences-1)]
                                    # for c in potential_sentence:  # loop through each character in line
                                    #     if not c.isalpha() and not c.isspace() or c in "ÒÓÔÕßˆˇ" and not c in ".,!?\":;()\'":  # test for unsuitable characters
                                    #         potential_sentence = potential_sentence.replace(c, "")  # replace unsuitable character with ""
                                    if not "." in potential_sentence and not potential_sentence.isspace():
                                        potential_sentence = potential_sentence + ". "
                                        number_extracted = 1
                                        # print("Let's take one and be happy about it.")
                                        return potential_sentence, page_num, number_extracted
                                    else:
                                        potential_sentence = ""
                                        number_extracted = 0
                                        return potential_sentence, page_num, number_extracted

                        except:
                            # print("failed to get sentence(s) from " + filename)
                            file_fail_count+=1
                            continue


def extract_from_document(file_path,type,max_lines):
    filename = os.path.basename(file_path)
    if filename:
        line = ""
        section_num = 0
        # try:
        if filename.endswith('.pdf'):
            # print(filename)
            try:
                line, section_num, number_extracted = extract_from_pdf(file_path,type,max_lines)
                return line, section_num, number_extracted
            except:
                print("unreadable pdf")
                number_extracted = 0
                return line, section_num, number_extracted
        elif filename.endswith('.doc') or filename.endswith('.docx'):
            # print("scraping " + filename + " at " + file_path + " for a line")
            # try:
            line, section_num, number_extracted = extract_from_doc(file_path,type,max_lines)
            return line, section_num, number_extracted
            # except:
            #     print("not docx friendly")
            #     number_extracted = 0
            #     return line, section_num, number_extracted

def extract_page_text_as_list(file_path,page_num):
    filename = os.path.basename(file_path)
    if filename:
        if filename.endswith('.pdf'):
            try:
                with open(file_path, 'rb') as pdf_file_obj:  # open current pdf
                    pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)  # store pdf as PyPDF2 reader object
                    if pdf_reader and not pdf_reader.isEncrypted:
                        # print("I'm looking for a line in " + filename)
                        number_of_pages = pdf_reader.getNumPages()  # count number of pages in pdf
                        if number_of_pages > 0:
                            file_fail_count = 0
                            line = ""
                            while not line:
                                if file_fail_count >= min(20, number_of_pages + 4):
                                    print("Out of " + str(file_fail_count) + " pages, could not get a line. :-/")
                                    break
                                page_obj = pdf_reader.getPage(page_num)
                                if '.' in page_obj.extractText():
                                    page_in_sentences = re.split("(?<=[.!?])+",page_obj.extractText())
                                    return page_in_sentences
                                else:
                                    page_in_lines = page_obj.extractText().splitlines()
                                    return page_in_lines
            except:
                print("cannot extract page text as list")

def preformed_or_scrambled():
    coin_flip = randint(0,1)
    if coin_flip == 0:
        mode = 'preformed'
    else:
        mode = 'scrambled'
    return mode

def refrain_or_no(mode, source_for_lines):
    roll_of_die = randint(1,8)
    if roll_of_die >= 7:
        refrain = 'yes'
    else:
        refrain = 'no'
    if refrain == 'yes':
        if mode == 'preformed':
            line = ""
            length = len(source_for_lines.paragraphs)
            if length >= 1:
                while not line:
                    paragraph_no = randint(0, length - 1)
                    # number_of_lines = len(doc.paragraphs[paragraph_no].text)
                    # paragraph_no = randint(0,number_of_lines-1)
                    potential_line = source_for_lines.paragraphs[paragraph_no].text
                    if potential_line:
                        potential_line_as_string = potential_line  # stores selected line for filtering
                        # print("Paragraph Scraped")
                        # print(potential_line_as_string)
                        # filter punctuation and digits
                        for c in potential_line_as_string:  # loop through each character in line
                            if not c.isalpha() and not c.isspace() or c in "ÒÓÔÕßˆˇ":  # test for unsuitable characters
                                potential_line_as_string = potential_line_as_string.replace(c,
                                                                                            "")  # replace unsuitable character with ""
                        words_in_line = nltk.tokenize.word_tokenize(potential_line_as_string)  # separate line into recognizable words
                        # print("Paragraph filtered")
                        words_in_paragraph_count = len(words_in_line)
                        if words_in_line and words_in_paragraph_count >= 2:
                            # print("There are " + str(words_in_paragraph_count) + " words left in the filtered section.")
                            if words_in_paragraph_count >= 20:
                                number_of_words = min(randint(4, words_in_paragraph_count), 10)
                                # print("Let's cut that down to " + str(number_of_words) + " words")
                                beginning_of_line = randint(0, words_in_paragraph_count - number_of_words)
                                end_of_line = beginning_of_line + number_of_words - 1
                                selected_words = words_in_line[beginning_of_line:end_of_line]
                                line = " ".join(selected_words)
                                return line
                            else:
                                line = " ".join(nltk.tokenize.word_tokenize(potential_line_as_string))
                                return line
        if mode == 'scrambled':
            line = " ".join(source_for_lines[randint(0,len(source_for_lines)-1)])
            return line

def select_a_path(path1,path2,path3, line_counter):
    # coin_flip = randint(0, 1)
    # if coin_flip == 0:
    #     this_path = path1
    # else:
    #     this_path = path2
    number = randint(1,3)
    if number == 1:
        this_path = path1
    elif number == 2:
        this_path = path2
    else:
        this_path = path3
    return this_path

def store_source(sources,filename,section_num,sections_and_sources,section):
    if filename.endswith('.pdf'):
        if filename not in sources:
            sources[filename] = {'page or section number': section_num}  # store source of current line being generated
    elif filename.endswith('.doc') or filename.endswith('.docx'):
        if not filename in sources:
            sources[filename] = {'page or section number': section_num}  # store source of current line being generated
    if str(section) not in sections_and_sources:
        sections_and_sources[str(section)] = [filename,section_num]
    else:
        sections_and_sources[str(section)].append([filename,section_num])
def add_line_of_verse(f,max_lines,first_line,line_counter,stanza_break,line):
    try:
        p = f.add_paragraph()
        if line_counter > 0 and line_counter % stanza_break == 0:
            run = p.add_run()
            run.add_break()
            run.add_break()
            max_lines += 1
            line_counter += 1  # increases line counter
            print("Stanza break on line " + str(line_counter))
        if line_counter == 0 and not first_line:
            first_line = line
            print("First line selected!")
            line_counter+=1
            p.add_run(line)  # appends line to poem file and goes to next line
            return first_line,line_counter
        line_counter += 1
        # print("Line number " + str(line_counter) + " selected!")
        p.add_run(line)  # appends line to poem file and goes to next line
        return first_line,line_counter
    except:
        print("Failed to add line as verse")

def add_scramble_line(f, all_lines_segmented,first_line,line_counter,line):
    if line_counter == 0 and not first_line:
        first_line = line
        print("First line selected!")
        line_counter += 1
        p = f.add_paragraph()
        p.add_run(line)  # appends line to poem file and goes to next line
        return first_line,line_counter

    # print("Line " + str(line_counter) + " will be: " + line)  # shows recently generated line
    # print(segment(line))
    all_lines_segmented.append(nltk.tokenize.word_tokenize(line))
    line_counter += 1
    return first_line, line_counter

def make_vignette(number_of_sentences,alldocpaths,sources,sections_and_sources):
    key = str(number_of_sentences)
    current_sentences = []
    total_extracted = 0
    remaining_sentences = number_of_sentences - total_extracted
    while True:
        if len(current_sentences) >= number_of_sentences:
            print("we've got enough lines for " + str(number_of_sentences))
            break
        sentence = ""
        page = 0
        random_path = ""
        number_extracted = 0
        while not sentence:
            random_path = sample(alldocpaths, 1)[0]
            sentence, page, number_extracted = extract_from_document(random_path, "sentence", remaining_sentences)
                # print("we got " + str(number_extracted) + " sentences from personal sources")
        if sentence:
            # try:
            # print("We're adding the lines to the list for " + key)
            if isinstance(sentence, list):
                for item in sentence:
                    item = re.sub("\d+", "", item)
                    if '.' not in item:
                        item = item + '.'
                    current_sentences.append(item)
            elif isinstance(sentence, str):
                list_of_sentences = sentence.split(".")
                if len(list_of_sentences) > 1:
                    for item in list_of_sentences:
                        item = re.sub("\d+", "", item)
                        if '.' not in item:
                            item = item + '.'
                        current_sentences.append(item)
                else:
                    sentence = re.sub("\d+", "", sentence)
                    current_sentences.append(sentence)
            # else:
            #     print("We're adding this line")
            #     current_sentences.append(sentence)
            store_source(sources, os.path.basename(random_path), page, sections_and_sources, number_of_sentences)
            # print("we've stored the sentence(s) and source(s)")
            # print("total sentence(s) extracted at this stage = " + str(total_extracted))
            remaining_sentences -= total_extracted
            shuffle(current_sentences)
            copy_of_current_sentences = current_sentences
            for sentence in copy_of_current_sentences:
                if not any(c.isalpha for c in sentence):
                    # print('already got this line or this line is empty')
                    current_sentences.remove(sentence)

                    # except:
                    #     print("something went wrong with the storage of the lines")

    print(str(len(current_sentences)) + " out of " + str(number_of_sentences) + " sentences being added")
    vignette = " ".join(current_sentences)
    return vignette, sources, sections_and_sources

def get_page_numbers_pdf(filepath):
    with open(filepath, 'rb') as pdf_file_obj:  # open current pdf
        pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)  # store pdf as PyPDF2 reader object
        if pdf_reader and not pdf_reader.isEncrypted:
            # print("I'm looking for a line in " + filename)
            number_of_pages = pdf_reader.getNumPages()  # count number of pages in pdf
            return number_of_pages

def shuffle_page(filepath, page_num):
    page_list = extract_page_text_as_list(filepath, page_num)
    new_page_list = shuffle(page_list)
    if new_page_list:
        if " ".join(new_page_list).count('(?<=[.!?])+') > 1:
            new_page_text = " ".join(new_page_list)
            return new_page_text
        else:
            new_page_text = "/n".join(new_page_list)
            return new_page_text
    else:
        print("couldn't get new page")
        new_page_text = ""
        return new_page_text

def shuffle_all_pages_in_pdf(filepath):
    x = docx.Document()
    page_numbers = get_page_numbers_pdf(filepath)
    for number in range(0, page_numbers - 1):
        try:
            new_page_text = shuffle_page(filepath, number)
            if new_page_text.count('(?<=[.!?])+') > 1:
                x.add_paragraph(new_page_text)

            elif new_page_text:
                y = x.add_paragraph()
                for line in new_page_text.split("/n"):
                    y.add_run(line)
                    if line.isspace():
                        y = x.add_paragraph()
        except:
            # page_text = text_from_image(pdf_to_image(filepath))
            # try:
            #     new_page_text = " ".join(shuffle(page_text))
            #     if new_page_text.count('(?<=[.!?])+') > 1:
            #         x.add_paragraph(new_page_text)
            #
            #     elif new_page_text:
            #         y = x.add_paragraph()
            #         for line in new_page_text.split("/n"):
            #             y.add_run(line)
            #             if line.isspace():
            #                 y = x.add_paragraph()
            # except:
            #     print("didn't work even with the ocr")

    filename = os.path.basename(filepath)
    x.save(filename.replace('.pdf',str(randint(0,5000) + '.docx')))



