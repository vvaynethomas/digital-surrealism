## for use in Anaconda 3.x
import datetime
import sys
import os
import re
from random import *
import numpy as np
from scipy import *
from PIL import Image, ImageEnhance, ImageFilter
import itertools
from itertools import product
from pdf2image import convert_from_path, convert_from_bytes
import requests
from bs4 import BeautifulSoup
import docx
import PyPDF2
import nltk
import pytesseract
pytesseract.pytesseract.tesseract_cmd = 'C:/Program Files (x86)/Tesseract-OCR/tesseract'
pytesseract.pytesseract.tessdata_prefix = 'C:/Program Files (x86)/Tsseract-OCR/tessdata'
english_vocab = set(w.lower() for w in nltk.corpus.words.words())


def merge_random_images(input_directory,output_directory,alpha=0,size=None):
    d1 = input_directory
    image1 = None
    image2 = None
    fail_count = 0
    while not image1 or not image2:
        try:
            image1 = sample(get_imlist(directory=d1), 1)[0]
            # print(image1)
            image2 = sample(get_imlist(directory=d1), 1)[0]
            # print(image2)
        except:
            print('keep looking for images')
            fail_count += 1
            if fail_count >= 20:
                break
    if image1 and image2:
        print('got images to blend')
        merge_images(image_path1=image1, image_path2=image2, size=size, save_directory=output_directory,alpha=alpha)

def get_filepaths(directory):

    """
    This function will generate the file names in a directory
    tree by walking the tree either top-down or bottom-up. For each
    directory in the tree rooted at directory top (including top itself),
    it yields a 3-tuple (dirpath, dirnames, filenames).
    """
    file_paths = []  # List which will store all of the full filepaths.

    # Walk the tree.
    if isinstance(directory,list):
        for folder in directory:
            for root, directories, files in os.walk(folder):
                for filename in files:
                    # Join the two strings in order to form the full filepath.
                    filepath = os.path.join(root, filename)
                    file_paths.append(filepath)  # Add it to the list.
    else:
        for root, directories, files in os.walk(directory):
            for filename in files:
                # Join the two strings in order to form the full filepath.
                filepath = os.path.join(root, filename)
                file_paths.append(filepath)  # Add it to the list.

    return file_paths  # Self-explanatory.

def get_pdf_paths(file_paths):
    paths_for_pdfs = []
    for path in file_paths:
        word = path
        regexp = re.compile(r'pdf')
        if regexp.search(word):
            paths_for_pdfs.append(path)
    return paths_for_pdfs


def get_imlist(directory):
    image_list = []
    if isinstance(directory,str):
        if directory.endswith('.jpg'):
            image_list.append(directory)
        elif directory.endswith('/'):
            for path in get_filepaths(directory):
                if path.endswith('.jpg'):
                    image_list.append(path)
    elif isinstance(directory,list):
        for item in directory:
            if isinstance(item,str):
                if item.endswith('.jpg'):
                    image_list.append(item)
                elif item.endswith('/'):
                    for path in get_filepaths(item):
                        if path.endswith('.jpg'):
                            image_list.append(path)
    return image_list

def initialize_images(imlist):
  for infile in imlist:
    outfile = os.path.splitext(infile)[0] + ".jpg"
    if infile != outfile:
      try:
        Image.open(infile).save(outfile)
      except IOError:
        print("cannot convert " + infile)

def powerset(L):
  pset = set()
  for n in range(len(L) + 1):
    for sset in itertools.combinations(L, n):
      pset.add(sset)
  return pset

def get_pairs(L):
    pset=set()
    for n in range(3):
        for sset in itertools.combinations(L,n):
            pset.add(sset)
    return pset

def primes(n):
    primfac = [1]
    d = 2
    while d*d <= n:
        while (n % d) == 0:
            primfac.append(d)  # supposing you want multiple factors repeated
            n //= d
        d += 1
    # if n > 1:
    primfac.append(n)
    # print(primfac)
    return primfac

def get_all_divisors(n):
    all_divisors = [1]
    prime_factors = primes(n)
    # print(set(prime_factors))
    # print(set(powerset(prime_factors)))
    prime_combinations = set(powerset(prime_factors))

    for tuple in prime_combinations:
        if len(tuple) > 1:
            divisor = 0

            # print(tuple)
            divisor = np.product(tuple)
            # print(str(divisor))
            if divisor not in all_divisors:
                all_divisors.append(divisor)
                # print('added ' + str(divisor) + ' to list of divisors')
    # print('got all divisors of ' + str(n))
    return all_divisors

def get_common_divisors(N,M):
    list_1 = get_all_divisors(N)
    # print(list_1)
    list_2 = get_all_divisors(M)
    # print(list_2)
    common_divisors = [1]
    all_divisors = list_1+list_2
    # divisor_set = list(itertools.chain(*all_divisors))
    divisor_set = set(all_divisors)
    for a in divisor_set:
        if N % a == 0 and M % a == 0 and a not in common_divisors:
            common_divisors.append(a)
    # print(common_divisors)
    return common_divisors

def gcd(N,M):
    return max(get_common_divisors(N,M))

def initialize_polygon(im, length=0, type='square',columns=0,rows=0,orientation=None):
    # print(length)
    # print(N)
    # print(M)
    # print(color_channels)
    N = im.size[0]
    M = im.size[1]
    # fail_count = 0
    if type == 'square':
        while not length or length <= 80:
            # print("let's find a good length for the cuboids")
            all_divisors = get_common_divisors(N,M)
            # print('options: ')
            # print(all_divisors)
            length = sample(all_divisors,1)[0]
            # print("how about " + str(length))
        if N-length-1 > 0 and M-length-1 > 0:
            left = randint(0,N-length-1)
            right = left + length
            upper = randint(0,M - length - 1)
            lower = upper + length
            box = (left,upper,right,lower)
            # print(str(right-left))
            # print(str(lower-upper))
            return length, box
    # elif type == 'triangle': #insert triangulation algorithm here
    elif type == 'strip':
        if orientation == 'horizontal':
            # print("Final image width = " + str(columns) + " & this image width = " + str(im.size[0]))
            N = min([columns, im.size[0]])
            alt_dimension = N
            # print(str(N))
            M = im.size[1]
            preferred_dimension = M

        else:
            # print("Final image height = " + str(rows) + " & this image height = " + str(im.size[1]))
            M = min([rows, im.size[1]])
            alt_dimension = M
            # print(str(M))
            N = im.size[0]
            preferred_dimension = N

        # while not length or length <= 50:
        #     print("let's find a good length for the cuboids")
        #     all_divisors = get_all_divisors(preferred_dimension)
        #     print('options: ')
        #     print(all_divisors)
        #     length = sample(all_divisors, 1)[0]
        #     print("how about " + str(length))



        if orientation == 'horizontal':

            left = 0
            right = N
            upper = randint(0,M-length-1)
            lower = upper + length
            box = (left,upper,right,lower)
        else:
            left = randint(0,N-length-1)
            right = left + length
            upper = 0
            lower = M
            box = (left,upper,right,lower)
        return length, box

def initialize_image(im, length=0, type='square',orientation=None):
    # print(length)
    # print(N)
    # print(M)
    # print(color_channels)
    N = im.size[0]
    M = im.size[1]
    fail_count = 0
    if type == 'strip':
        if not orientation:
            orientation = sample(['horizontal','vertical'],1)[0]
        # print(orientation)
        if orientation == 'horizontal':
            preferred_dimension = M
            alt_dimension = N
        else:
            preferred_dimension = N
            alt_dimension = M
        # else:
        #     preferred_dimension = sample([N,M],1)
        #     alt_dimension = [n for n in [N,M] if not n == preferred_dimension][0]

        pref_distance = length
        while not pref_distance or pref_distance <= 20:
            pref_distance = sample(get_all_divisors(preferred_dimension),1)[0]
        alt_distance = alt_dimension
        if preferred_dimension % pref_distance == 0 and pref_distance > 1:
            try:
                if orientation == 'horizontal':
                    left = 0
                    right = alt_distance-1
                    upper = randint(0,preferred_dimension-pref_distance-1)
                    lower = upper + pref_distance
                    box = (left,upper,right,lower)
                    return pref_distance,box,orientation
                else:
                    left = randint(0,preferred_dimension-pref_distance-1)
                    right = left + pref_distance
                    upper = 0
                    lower = alt_distance-1
                    box = (left,upper,right,lower)
                    return pref_distance,box,orientation
            except:
                print("had a problem, let's move on")
                pref_distance = 0
                box = None
                return pref_distance,box,orientation

    elif type == 'square':
        while not length or length <= 50:
            # print("let's find a good length for the cuboids")
            all_divisors = get_common_divisors(N,M)
            # print('options: ')
            # print(all_divisors)
            length = sample(all_divisors,1)[0]
            # print("how about " + str(length))
            fail_count +=1
            if fail_count >= len(all_divisors)*len(all_divisors):
                print('cuboids too small')
                break
        if N % length == 0 and length > 1:
            left = randint(0,N-length-1)
            right = left + length
            upper = randint(0,M - length - 1)
            lower = upper + length
            box = (left,upper,right,lower)
            # print(str(right-left))
            # print(str(lower-upper))
            return length, box

def extract_image_properties(image):
  pil_im = image
  im_array = np.array(pil_im)
  N = im_array.shape[1]
  M = im_array.shape[0]
  if len(im_array.shape) == 3:
    color_channels = im_array.shape[2]
  else:
      color_channels = 'None'
  return N, M, color_channels

def grid_image(im,length,type='square',orientation="horizontal"):
    columns,rows = im.size
    # print(columns,rows)
    all_boxes = []
    if type=='square':
        for i in range(0,int(columns/length)):
            left = length * i
            right = left + length
            for j in range(0,int(rows/length)):
                upper = length * j
                lower = upper + length
                box = (left,upper,right,lower)
                if not box in all_boxes:
                    all_boxes.append(box)
    elif type=='strip':
        if orientation == "horizontal":
            for i in range(0,int(rows/length)):
                left = 0
                right = columns
                upper = length*i
                lower = upper + length
                box = (left,upper,right,lower)
                if not box in all_boxes:
                    all_boxes.append(box)
        else:
            for i in range(0,int(columns/length)):
                left = length*i
                right = left + length
                upper = 0
                lower = rows
                box = (left,upper,right,lower)
                if not box in all_boxes:
                    all_boxes.append(box)
    return all_boxes

def tile_image(im):
    columns,rows = im.size
    full_area = columns*rows
    all_boxes = []
    lengths = get_common_divisors(columns,rows)
    area_added = 0
    while area_added < full_area:
        remaining_area = full_area - area_added
        length = min(columns, rows)
        number_to_fit = floor(remaining_area/length^2)
        for i in range(number_to_fit):
            left = 0
            upper = 0
            right = left + length
            lower = upper + length
            box = (left,upper,right,lower)
            all_boxes.append(box)
        area_added += length*length*number_to_fit

def get_random_size():
    length_1,length_2 = 0,0
    while length_1 < 600 or length_2 < 600:
        length_1 = randint(3,24)*100
        length_2 = int(.75*length_1)
        lengths = [length_1,length_2]
    return (i for i in sample(lengths,len(lengths)))

def create_cubomanic(input_directory,output_directory,columns=2700,rows=3600,length=0,rotate=False,conversions=False):
    # print(input_directory)
    image_list = get_imlist(input_directory)
    # print("Found " + str(len(image_list)) + " images")
    while not len(image_list) >= 1:
        for directory in input_directory:
            im_list = get_imlist(directory)
            for im_path in im_list:
                if not im_path in image_list:
                    image_list.append(im_path)
        image_list = input_directory
        break
    if length and length > rows/40:
        im = Image.new("RGB", (2700, 3600))
        length, box = initialize_image(im,length=length,type='square')

    while not length or length < rows/40:
        if not (rows % 100 == 0 and rows > 500) or not (columns % 100 == 0 and columns > 500):
            rows, columns = get_random_size()
        im = Image.new("RGB", (columns, rows))
        length, box = initialize_image(im)
        # print(length, box)
    # print(length)
    sources = []
    area = columns*rows
    number_of_cuboids = int(area / (length * length))
    l = length
    print(l)
    # print("now let's fill a page with them")
    all_boxes = grid_image(im, length)
    # print(all_boxes)
    # print("let's get " + str(number_of_cuboids) + " cuboids")
    for i in range(0,len(all_boxes)):
        # print("working on cuboid number " + str(i))
        while True:
            this_box = all_boxes[i]
            # print(this_box)
            this_cuboid = Image.new('RGB',(columns,length))
            # print(this_cuboid.size)
            current_path = sample(image_list, 1)[0]
            # try:
            try:
                current_image = Image.open(current_path)
            except:
                continue
            # print(current_image.size)
            current_box = None
            if current_image and length > 0 and current_image.size[1] > length:
                l = None
                while not l or not current_box:
                    try:
                        l, current_box = initialize_polygon(current_image, length)
                    except:
                        # print("couldn't make a big enough box")
                        break
                # print(l)
                # print(current_box)

            if l and current_box:
                try:
                    cuboid = current_image.crop(current_box)
                    # cuboid.show()
                    # print(cuboid.size)
                    if os.path.basename(current_path) not in sources:
                        sources.append(os.path.basename(current_path))
                    this_cuboid = cuboid
                    # this_cuboid.show()
                    # print(this_cuboid.size)
                    # try:
                    if rotate == True:
                        rotation_values = [0,90,180,270]
                        this_cuboid = this_cuboid.rotate(sample(rotation_values,1)[0])
                    im.paste(this_cuboid,this_box)
                    # print('pasted')
                    break
                except:
                    print("failed")
            # except:
            #     print("couldn't paste")
            #     continue
            # except:
            #     print('failed to paste cuboid')
            #     continue
    print("finally, let's save the new cubomanic")
    print(str(l))
    name = 'cubomanic' + '_' + str(len(sources)) + '_' + str(l)
    outfile = output_directory + name + ".jpg"
          # try:
    if conversions == True:
        success = False
        while not success:
            try:
                mode = str(randint(1,3))
                im.convert(mode)
                success = True
            except:
                print("bad conversion")
                continue
    im.save(outfile)

          # except IOError:
          #   print("cannot convert " + infile)
    # print(outfile)
    # print(sources)
    # im.show()

def create_stripomanic(input_directory,output_directory,columns=2700,rows=3600,pref_distance=0,orientation="",conversions=False):
    # print(input_directory)
    if isinstance(input_directory,str):
        image_list = get_imlist(input_directory)
    elif isinstance(input_directory,list):
        image_list = []
        for directory in input_directory:
            # print(directory)
            if isinstance(directory,str):
                if directory.endswith('.jpg'):
                    image_list.append(directory)
                elif directory.endswith('/'):
                    # print('directory')
                    for path in get_imlist(directory):
                        # print(path)
                        if not os.path.basename(path) in image_list:
                            image_list.append(path)
            elif isinstance(directory,list):
                for item in directory:
                    if not os.path.basename(item) in image_list:
                        image_list.append(item)

    print("Found " + str(len(image_list)) + " images")
    if image_list:
        # if pref_distance and pref_distance > 10:
        im = Image.new("RGB", (columns, rows))
        fail_count = 0
        while not pref_distance or pref_distance < 100:
            pref_distance, box,orientation = initialize_image(im,length=pref_distance,type='strip',orientation=orientation)
            fail_count += 1
            if fail_count > 50:
                break
        # while not pref_distance or pref_distance < 100:
        #     if not (rows % 100 == 0 and rows > 500) or not (columns % 100 == 0 and columns > 500):
        #         rows, columns = get_random_size()
            # print(rows,columns)
            # im = Image.new("RGB", (2700, 3600))
            # pref_distance, box,orientation = initialize_image(im,type='strip',orientation=orientation)
            # print(pref_distance, box)
        # print(pref_distance)
        if pref_distance == 0:
            print("big problem")
            return
        sources = []
        area = columns*rows
        if orientation == 'horizontal':
            preferred_dimension = rows
            alt_dimension = columns
        else:
            preferred_dimension = columns
            alt_dimension = rows
        number_of_cuboids = int(area / (alt_dimension * pref_distance))


        # print("now let's fill a page with them")
        all_boxes = grid_image(im, pref_distance,type='strip',orientation=orientation)
        # print(all_boxes)
        print("let's get " + str(number_of_cuboids) + " strips")
        for i in range(0,len(all_boxes)):
            # print("working on strip number " + str(i) + " out of " + str(len(all_boxes)))
            this_box = all_boxes[i]
            # print(this_box)
            while True:
                # print('grabbing an image')
                current_path = sample(image_list, 1)[0]
                # print(current_path)
                try:
                    # print('opening it')
                    current_image = Image.open(current_path)
                    # print('opened it')
                except:
                    # print("couldn't")
                    continue
                # print("This image's size: " + "(" + ",".join([str(i) for i in current_image.size]) + ")")

                if orientation == 'horizontal':
                    this_cuboid = Image.new('RGB',(columns,pref_distance))
                    # print(this_cuboid.size)
                else:
                    this_cuboid = Image.new('RGB',(pref_distance,rows))

                try:
                    # print("Let's get a strip from it")
                    l, current_box = initialize_polygon(current_image, pref_distance,type='strip',columns=columns,rows=rows,orientation=orientation)

                    cuboid = current_image.crop(current_box)
                    if not sum(np.array(cuboid)) > 10:
                        continue
                    else:
                        # print('got a strip')
                    # cuboid.show()
                    #     print(cuboid.size)
                        if os.path.basename(current_path) not in sources:
                            sources.append(os.path.basename(current_path))

                        if not this_cuboid.size == cuboid.size:
                            this_cuboid = cuboid.resize(this_cuboid.size)
                            # this_cuboid.show()
                        if not sum(np.array(this_cuboid)) > 10:
                            # print("it's all black")
                            continue
                        # elif not all(this_cuboid,axis=1) == 255:
                        #     print("")
                        else:
                            im.paste(this_cuboid,this_box)
                            # print('pasted')
                            break
                except:
                    # print("couldn't paste")
                    continue
                # except:
                #     print('failed to paste cuboid')
                #     continue
        print("finally, let's save the new stripomanic")
        name = 'stripomanic' + '_' + str(len(sources)) + '_' + str(randint(0,500))
        outfile = output_directory + name + ".jpg"
              # try:
        if conversions == True:
            mode = str(randint(1,3))
            im.convert(mode)
        im.save(outfile)
              # except IOError:
              #   print("cannot convert " + infile)
        # print(outfile)
        # print(sources)
        # im.show()

def create_tiling_cubomanic(input_directory,output_directory,columns=900,rows=1200):
    area_added = 0
    while area_added < columns*rows:
        length = min(columns,rows)

def pdf_to_jpg(pdf_path):
    # Extract jpg's from pdf's. Quick and dirty
    pil_ims = convert_from_path(pdf_path)
    directory = os.path.dirname(pdf_path)
    os.chdir(directory)
    name = os.path.basename(pdf_path).replace('.pdf','')
    regexp = re.compile(name)
    if not regexp.search(" ".join(get_imlist(os.path.dirname(pdf_path)))):
        pages = len(pil_ims)
        for i in range(pages):
                with pil_ims[i] as img:
                    outfile = directory + name + ".jpg"
                    # try:
                    img.save(outfile)

def merge_images(image_path1,image_path2,size=(2700,3600),alpha=0,mask=None,primary=None,save_directory=None):
    # print("resizing for merge")
    # try:
    # print(image_path1,image_path2)
    print('opening image')
    image1 = Image.open(image_path1)
    image2 = Image.open(image_path2)
    # if not size:
    #     size1 = image1.size
    #     size2 = image2.size
    #     # print(size1,size2)
    #     while not size1 == size2:
    #        if not primary:
    #            image1 = image1.resize(size2)
    #            size1 = image1.size
    #            size2 = image2.size
    #            image2 = image2.resize(size1)
    #        elif primary == True:
    #             image2 = image2.resize(size1)
    #             size2 = image2.size
    # else:
    print('resizing')
    image1 = image1.resize(size,Image.ANTIALIAS)
    image2 = image2.resize(size,Image.ANTIALIAS)
    # except:
    #     print("can't do it")
    #     blend_pil = None
    #     comp_pil = None
    #     alpha = None
    #     return blend_pil, comp_pil,alpha

    #     print(size1,size2)
    # print(size1,size2)
    # print(array(image1).shape,array(image2).shape)
    while not alpha > 0 or alpha >=1:
        alpha = randint(10,85)/100
    print('blending')
    blend_pil = Image.blend(image1,image2,alpha)
    image_list = []
    for image in [image1,image2]:
        if not image.mode == "RGBA":
            image = image.convert("RGBA")
            image_list.append(image)
        else:
            image_list.append(image)

    comp_pil = Image.alpha_composite(image_list[0],image_list[1])
    # blend_im = Image.new("RGB", (blend_pil.size[0], blend_pil.size[1]))
    if mask:
        comp_pil = Image.composite(image1,image2,mask)
    else:
        comp_arr = np.array(image1) + np.array(image2)
        # too_high_indices = comp_arr > 255
        # too_low_indices = comp_arr < 0
        # comp_arr[too_low_indices] = 0
        # comp_arr[too_high_indices] = 255
        # for j in comp_arr.shape[1]:
        #     for i in comp_arr.shape[0]:
        #         if comp_arr[i,j] > 255:
        #             comp_arr[i,j] = 255
        #         elif comp_arr[i,j] < 0:
        #             comp_arr[i,j] = 0

        comp_pil = Image.fromarray(comp_arr)
        # comp_im = Image.new("RGB", (comp_pil.size[0], comp_pil.size[1]))
    # comp_pil = None
    # comp_im = None
    if save_directory:
        print('trying to save blend')
        try:
            blend_pil.save(save_directory+'blend' +'_' + os.path.basename(image_path1).replace('.jpg','') + '_' + os.path.basename(image_path2).replace('.jpg','') + '_' + str(alpha) + '.jpg')
        except:

            blend_pil.save(save_directory + 'blend' + '_' + str(randint(10,1000)) + '.jpg')
    print('all done')
    return blend_pil, comp_pil, alpha
    # except:
    #     print("something went wrong with these images")
    #     blend_pil = None
    #     comp_pil = None
    #     alpha = None
    #     return blend_pil,comp_pil, alpha

def merge_combinations(input_directory,output_directory,max_to_make=50,no_of_pairs=None,alpha=0):
    # print("let's find all the images in the chosen directory")
    im_list = get_imlist(input_directory)
    # print("found: " + str(len(im_list)) + " images to blend.")
    # if len(im_list) > 100:
    #     number_to_get = randint(10,80)
    #     im_list = sample(im_list,number_to_get)
    #     print("Let's get " + str(number_to_get) + " of them for blending")
    image_count = 0
    while image_count < max_to_make:
        # print("let's find out how many pairs we can make from these images")
        # all_combos = get_pairs(im_list)
        # # print(all_combos)
        # print("looks like there are a possible " + str(len(all_combos)) + \
        # " different pairs of images that can be combined")
        # if not isinstance(no_of_pairs,int):
        #     no_of_pairs = len(all_combos)
        #     print("since no numbers of pairs were specified, let's make "  + str(max_to_make))
        #     print("in fact, we could make up to " + str(no_of_pairs))
        # for combo in sample(all_combos,no_of_pairs):
        #     # try:
            if image_count >= max_to_make:
                # print("we've got enough")
                break
            combo = sample(im_list,2)
            # print(combo)
            # print(combo)
            # if not len(combo) > 1:
            #     continue
            image_path1,image_path2 = combo[0],combo[1]
            sources = [os.path.basename(image_path1),os.path.basename(image_path2)]
            blend_im,comp_im,alpha = merge_images(image_path1,image_path2)
            # try:
            # if comp_im:
            #     # try:
            #     name = 'composite' + '_' + "_".join(sources).replace('.jpg','') + '_' + str(alpha)
            #     outfile = output_directory + name + ".jpg"
            #     # comp_im.show()
            #     comp_im.save(outfile)
            #     image_count+=1
                # except:
                #     print('failed to produce/save alpha composite')
            if blend_im:
                sources_string = '_'.join([sources[0],sources[1]]).replace('.jpg','')
                if len(sources_string) > 30:
                    source_1_last = sources[0][-15:].replace('.jpg','')
                    source_2_last = sources[1][-15:].replace('.jpg','')
                    sources_string = '_'.join([source_1_last,source_2_last])
                # sources_string = " ".join(sources).replace('.jpg','')
                # print(sources_string)
                name = 'blend' + '_' + sources_string + '_' + str(alpha) + '_' + str(randint(0, 500))
                outfile = output_directory + name + ".jpg"
                # print("finally, let's save the new blend as " + name + ".jpg")
                # blend_im.show()

                blend_im.save(outfile)

                image_count += 1
            # print(str(image_count))
            if image_count >= max_to_make:
                break
                # except:
                #     print("Something went wrong trying to save the merged versions")
            #     continue
                # try:
def text_on_image(text_directory,image_directory,output_directory=None,max_to_make=1,list_of_paths=None):
    # print('getting all text image paths')
    if not list_of_paths:
        pages = get_imlist(text_directory)
    else:
        pages = list_of_paths
    # print('getting all image paths to merge with')
    images = get_imlist(image_directory)
    if pages and images:
        # print("we've got both")
        for i in range(max_to_make):
            # print(i)
            text_path = sample(pages,1)[0]
            image_path = sample(images,1)[0]
            # try:
            # print("let's merge a page with an image")
            blend_im,comp_im,alpha = merge_images(text_path,image_path,primary=True)

            if blend_im:
                # print("got it")
                sources = [os.path.basename(text_path), os.path.basename(image_path)]
                sources_string = '_'.join([sources[0], sources[1]]).replace('.jpg', '')
                # if len(sources_string) > 30:
                #     source_1_last = sources[0][-15:].replace('.jpg', '')
                #     source_2_last = sources[1][-15:].replace('.jpg', '')
                #     sources_string = '_'.join([source_1_last, source_2_last])
                # sources_string = " ".join(sources).replace('.jpg','')
                # print(sources_string)
                try:
                    name = 'blend' + '_' + sources_string + '_' + str(alpha) + '_' + str(randint(0, 500))
                    outfile = output_directory + name + ".jpg"
                    print("finally, let's save the new blend as " + name + ".jpg")
                    # blend_im.show()

                    blend_im.save(outfile)
                except:
                    try:
                        name = 'blend' + '_' + sources_string[-20:] + '_' + str(alpha) + '_' + str(randint(0, 500))
                        outfile = output_directory + name + ".jpg"
                        print("finally, let's save the new blend as " + name + ".jpg")
                        # blend_im.show()

                        blend_im.save(outfile)
                    except:
                        print("something went wrong")
                        continue

            # except:
            #     print("failed to merge text and image")
            #     continue
def text_on_image2(text_directory,image_directory,max_to_make=1):
    pages = get_imlist(text_directory)
    if not pages:
        pages = text_directory

    images = get_imlist(image_directory)
    for i in range(max_to_make):
        blend_im = None
        text_path = sample(pages,1)[0]
        image_path = sample(images,1)[0]
        text_image = Image.open(text_path)
        page_image = Image.open(image_path)
        text_array = np.array(text_image)
        page_array = np.array(page_image)
        if not text_array.shape == page_array.shape:
            # print(text_array.shape)
            # print(page_array.shape)
            larger_width = max(text_array.shape[0],page_array.shape[0])
            # print(larger_width)
            taller_height = max(text_array.shape[1],page_array.shape[1])
            # print(taller_height)
            text_template = np.zeros((larger_width,taller_height,3),dtype=np.uint8)
            # print(text_template.shape)
            x_offset = randint(0,larger_width - text_array.shape[0]) # 0 would be what you wanted
            y_offset = randint(0,taller_height- text_array.shape[1])  # 0 in your case
            text_template[x_offset:text_array.shape[0] + x_offset, y_offset:text_array.shape[1] + y_offset] = text_array
            text_array = text_template
            # print(text_array.shape)
            page_template = np.zeros((larger_width, taller_height,3),dtype=np.uint8)
            # print(page_template.shape)
            x_offset = randint(0, larger_width - page_array.shape[0])  # 0 would be what you wanted
            y_offset = randint(0, taller_height - page_array.shape[1])  # 0 in your case
            page_template[x_offset:page_array.shape[0] + x_offset, y_offset:page_array.shape[1] + y_offset] = page_array
            page_array = page_template
            # print(page_array.shape)
            text_image = Image.fromarray(text_array,"RGB")
            page_image = Image.fromarray(page_array,"RGB")

        comp_arr = np.array(text_image) + np.array(page_image)
        # too_high_indices = comp_arr >= 255
        # too_low_indices = comp_arr <= 0
        # comp_arr[too_low_indices] = 0
        # comp_arr[too_high_indices] = 0
        comp_im = Image.fromarray(comp_arr)
        # if blend_im:
        #     sources = [os.path.basename(text_path), os.path.basename(image_path)]
        #     sources_string = '_'.join([sources[0], sources[1]]).replace('.jpg', '')
        #     # if len(sources_string) > 30:
        #     #     source_1_last = sources[0][-15:].replace('.jpg', '')
        #     #     source_2_last = sources[1][-15:].replace('.jpg', '')
        #     #     sources_string = '_'.join([source_1_last, source_2_last])
        #     # sources_string = " ".join(sources).replace('.jpg','')
        #     print(sources_string)
        #     name = 'blend' + '_' + sources_string + '_' + '_' + str(randint(0, 500))
        #     outfile = output_directory + name + ".jpg"
        #     print("finally, let's save the new blend as " + name + ".jpg")
        #     # blend_im.show()
        #     blend_im.save(outfile)
        if comp_im:
            sources = [os.path.basename(text_path), os.path.basename(image_path)]
            sources_string = '_'.join([sources[0], sources[1]]).replace('.jpg', '')
            # if len(sources_string) > 30:
            #     source_1_last = sources[0][-15:].replace('.jpg', '')
            #     source_2_last = sources[1][-15:].replace('.jpg', '')
            #     sources_string = '_'.join([source_1_last, source_2_last])
            # sources_string = " ".join(sources).replace('.jpg','')
            print(sources_string)
            # try:
            #     name = 'combination' + '_' + sources_string + '_' +'_' + str(randint(0, 500))
            #     outfile = output_directory + name + ".jpg"
            #     print("finally, let's save the new combination as " + name + ".jpg")
            #     # blend_im.show()
            #     comp_im.save(outfile)
            # except:
            #     try:
            #         name = 'combination'  + '_' + sources_string[-15:] + '_' + str(randint(0,1000))
            #         outfile = output_directory + name + ".jpg"
            #         print("finally, let's save the new combination as " + name + ".jpg")
            #         # blend_im.show()
            #         comp_im.save(outfile)
            #     except:
            #         print("couldn't save this one")
        # except:
        #     print("failed to merge text and image")
        #     continue

def scrape_images(site):
    r = requests.get(site)
    data = r.text
    soup = BeautifulSoup(data, "lxml")
    for link in soup.find_all('img'):
            image = link.get('src')
            if image.endswith('.jpg'):
                try:
                    image_name = os.path.split(image)[1]
                    r2 = requests.get(image)
                    with open(image_name, "wb") as f:
                        f.write(r2.content)
                    # print("got an image")
                except:
                    try:
                        if not image.startswith('http'):
                            image = 'http:/' + image
                        image_name = os.path.split(image)[1]
                        r2 = requests.get(image)
                        with open(image_name, "wb") as f:
                            f.write(r2.content)
                        # print("got an image")
                    except:
                        print("something went wrong with this one")
                        continue

def download_free_images(keywords,output_directory=None):
    site = 'https://pixabay.com/en/photos/?q=' + "+".join(keywords.split(" "))
    if not output_directory:
        output_directory = picture_path + keywords + '/'
        if not os.path.exists(output_directory):
            os.makedirs(output_directory)
    os.chdir(output_directory)
    scrape_images(site)
    print('images scraped for: ' + keywords)

def grab_pos(path,pos_tag='NN', max_words=500):
    fail_count = 0
    word_list = []
    print('parsing ' + os.path.basename(path) + ' for words')
    page_text = ""
    if path.endswith('.pdf'):
        while not page_text or len(page_text.split(' ')) < 2:
            # print("scraping a page of text (after " + str(fail_count) + " tries)")
            if fail_count > 20:
                print("couldn't get any text")
                return word_list
            try:
                page_text,page_number = extract_from_pdf(path)

            except:
                fail_count += 1
                # print('failed to get a page (after failing ' + str(fail_count) + ' times')
            fail_count +=1


        # page_as_sentences = split_text([page_text])
    elif path.endswith('.docx'):
        fail_count = 0
        page_text = ""
        print("trying to extract all text from this doc")
        page_text = extract_all_text_from_docx(path)
        if page_text:
            # print(page_text)
            # fail_count +=1
            # if fail_count > 20:
            #     print("can't parse this doc")
            #     return word_list
            if not isinstance(page_text[0],str):
                print('converting runs to strings')
                all_lines_as_list = []
                for run in page_text:
                    all_lines_as_list.append(run.text)
                page_text = " ".join(all_lines_as_list)

            if not isinstance(page_text, str):
                print('this page text is in the wrong form')
                try:
                    page_text = " ".join(page_text)
                except:
                    print('no luck fixing it')
                    return word_list
        else:
            print('no text here')
            return word_list
    number_of_words_available = len(nltk.tokenize.word_tokenize(page_text))
    if not number_of_words_available:
        return word_list
    print("found " + str(number_of_words_available) + " words in this document")
    if number_of_words_available < 5:
        word_list = sample(nltk.tokenize.word_tokenize(page_text),len(nltk.tokenize.word_tokenize(page_text)))
        return word_list

    page_in_words = nltk.tokenize.word_tokenize(page_text)
    page_tagged = nltk.pos_tag(page_in_words)
    page_in_pos = []
    for word,tag in page_tagged:
        # print(word,tag)
        if pos_tag in tag and not word in page_in_pos:
            page_in_pos.append(word)
    number_of_available_words = len(page_in_words)
    number_of_available_pos = len(page_in_pos)
    word_list = page_in_pos
    # print(word_list)
    print('added ' + str(number_of_available_pos) + ' words to list')
    return word_list

def extract_from_pdf(filepath,type='page',subtype='string'):
    # filename = os.path.basename(filepath)
    with open(filepath, 'rb') as pdf_file_obj:  # open current pdf
        try:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)  # store pdf as PyPDF2 reader object
            if pdf_reader and not pdf_reader.isEncrypted:
                # print("I'm looking for a line in " + filename)
                number_of_pages = pdf_reader.getNumPages()  # count number of pages in pdf
                if number_of_pages > 0:
                    if type == 'all' :
                        text_in_lines = []
                        text_in_pages = []
                        for i in range(0,number_of_pages - 1):
                                page_obj = pdf_reader.getPage(i)
                                page_as_list = page_obj.extractText().splitlines()
                                text_in_pages.append(page_as_list)
                                for line in page_as_list:
                                    text_in_lines.append(line)
                        return text_in_pages, text_in_lines
                    elif type == 'page':
                        page_number = randint(0,number_of_pages-1)
                        page_obj = pdf_reader.getPage(page_number)
                        if subtype == 'lines':
                            page_in_lines = page_obj.extractText().splitlines()
                            # text_in_pages = [text_in_lines]
                            return page_in_lines, page_number
                        else:
                            page_as_string = page_obj.extractText()
                            return page_as_string, page_number

                    elif type == 'line':
                        fail_count = 0
                        line = ""
                        page_num = 0
                        while not line and fail_count < number_of_pages:
                            page_num = randint(1, number_of_pages - 1)  # generate random index to choose page
                            page_obj = pdf_reader.getPage(page_num)  # reads and stores random page as page object
                            try:
                                page_in_lines = page_obj.extractText().splitlines()  # stores extracted text as multiline string
                                num_lines = len(page_in_lines)  # count number of lines detected
                                if num_lines > 0:
                                    # print("we've got lines here")
                                    line = sample(page_in_lines,1)[0]
                                    return line, page_num
                                else:
                                    # print("no lines detect on page " + page_num)
                                    line = ""
                                    fail_count +=1
                                    continue

                            except:
                                print('failed to extract a line from pdf page ' + page_num)
                                fail_count +=1
                                continue
                        return line, page_num

                    elif type == 'sentence':
                        fail_count = 0
                        sentence = ""
                        page_num = 0
                        while not sentence and fail_count < number_of_pages:
                            page_num = randint(0, number_of_pages-1)
                            page_obj = pdf_reader.getPage(page_num)
                            try:
                                page_as_string = page_obj.extractText()
                                page_in_sentences = nltk.tokenize.sent_tokenize(page_as_string)
                                sentence = sample(page_in_sentences,1)[0]
                            except:
                                print('failed to extract a sentence from pdf page ' + str(page_num))
                                fail_count +=1
                                continue
                        return sentence, page_num
        except:
            print('unable to parse pdf')
def extract_all_text_from_docx(filepath, paragraph_or_runs='runs'):
    try:
        doc = docx.Document(filepath)
    except:
        print("couldn't scrape it")
        all_runs = []
        return all_runs
    if doc:
        # print("attempting to scrape " + os.path.basename(filepath))
        file_fail_count = 0
        all_paragraphs = []
        all_runs = []
        length = len(doc.paragraphs)
        # print(str(length))
        if length >= 1:
            for paragraph in doc.paragraphs:
                all_paragraphs.append(paragraph)
                for run in paragraph.runs:
                    all_runs.append(run.text)
            if paragraph_or_runs == 'paragraphs':
                return all_paragraphs
            else:
                # print("got all " + str(length) + " run(s)")
                # print("sample: " + all_runs[0])
                return " ".join(nltk.tokenize.word_tokenize(" ".join(all_runs)))

def download_relevant_images(text_path,actually=True):
    text_name = os.path.basename(text_path)
    output_directory = picture_path + text_name + '/'
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    word_list=grab_pos(text_path)
    output_list = []
    for word in word_list:
        if actually == True:
            download_free_images(word,output_directory=output_directory)
        if not output_directory in output_list:
            output_list.append(output_directory)
    return output_list

def initialize_art_for_text(text_path,need_images=False):
    text_name = os.path.basename(text_path)
    text_name = text_name[:-4]
    # print(text_name)
    text_image_pages = []
    for path in get_imlist('C:/Users/Vvayne/Pictures/images_from_pdfs/'):
        if re.search(text_name, path):
            if not path in text_image_pages:
                text_image_pages.append(path)
                # print('got the image of a page')
                # print(path)
    # print(len(text_image_pages))
    if text_image_pages:
        output_directory = picture_path + text_name + '/'
        if not os.path.exists(output_directory):
            os.makedirs(output_directory)
        if need_images == True:
            word_list = grab_pos(text_path)
            # print('grabbed words')

            for word in word_list:
                if len(word) > 1 and word.lower() in english_vocab:
                    download_free_images(word,output_directory=output_directory)
                    print("downloaded free images for: " + word)
            print("got all the free images to play with")
        final_output_directory = 'C:/Users/Vvayne/Documents/FA 2017/General/Outputs/' + text_name + '/'
        if not os.path.exists(final_output_directory):
            os.makedirs(final_output_directory)
        os.chdir(final_output_directory)
        return output_directory, text_image_pages,final_output_directory
    else:
        print("couldn't find the image version of the document")

def convert_to_bw(image_path):
    image_file = Image.open(image_path) # open colour image
    image_file = image_file.convert('1') # convert image to black and white
    # image_file.save('result.png')
    return image_file

def get_text_from_image(image_path):
    im = Image.open(image_path)
    im = im.filter(ImageFilter.MedianFilter())
    # enhancer = ImageEnhance.Contrast(im)
    # im = enhancer.enhance(2)
    im = im.convert('1')
    # im.save('temp2.jpg')
    text = pytesseract.image_to_string(im)
    # print(text)
    return text, im

def random_ocr(input_directory='C:/Users/Vvayne/'):
    f = docx.Document()
    print('hunting for an image to ocr')
    all_images = get_imlist(input_directory)
    if len(all_images) < 1:
        all_images = input_directory
    image_path = sample(all_images,1)[0]
    text,im = get_text_from_image(image_path)

    text_in_words = nltk.tokenize.word_tokenize(text)

    intersection = set(text_in_words).intersection(english_vocab)
    text_in_sentences = nltk.tokenize.sent_tokenize(text=' '.join(intersection))
    while not len(text_in_sentences) > 2 or not len(text_in_words) > 4 or not len(intersection) > 5:
        image_path = sample(all_images, 1)[0]
        text,im = get_text_from_image(image_path)
        text_in_words = nltk.tokenize.word_tokenize(text)
        intersection = set(text_in_words).intersection(english_vocab)
        text_in_sentences = nltk.tokenize.sent_tokenize(' '.join(intersection))
    im.show()
    for sentence in text_in_sentences:
        # print('got an image with words we can read')
        f.add_paragraph(sentence)
    name = os.path.basename(image_path).replace('.jpg','') + '.docx'
    text_output_directory = 'C:/Users/Vvayne/Documents/FA 2017/General/Outputs/extracted_texts/'
    if not os.path.exists(text_output_directory):
        os.makedirs(text_output_directory)
    os.chdir(text_output_directory)
    f.save(os.path.join(text_output_directory,name))
    print('saved the extracted text to ' + text_output_directory)
    print("it's called " + name)
def get_document_paths(file_paths):
    paths_for_potentials = []
    for path in file_paths:
        word = path
        if word.endswith('.docx') or word.endswith('.pdf') or word.endswith('.doc'):
            paths_for_potentials.append(path)
    return paths_for_potentials
def high_res_merge(image_1,image_2,size=(3600,5400),alpha=None):
    image1 = Image.open(image_1)
    image2 = Image.open(image_2)
    image1.resize(size, Image.ANTIALIAS)
    image2.resize(size, Image.ANTIALIAS)
    while not alpha > 0 or alpha >=1:
        alpha = randint(10,85)/100
    blend_pil = Image.blend(image1,image2,alpha)
    return blend_pil

def convert_to_highres(image_path,size=(5400,3600)):
    print('opening original image:' + os.path.basename(image_path))
    image1 = Image.open(image_path)
    os.chdir(os.path.dirname(image_path))
    print('resizing')
    highres = image1.resize(size, Image.ANTIALIAS)
    print('saving')
    name = os.path.basename(image_path).replace('.jpg','_'+str(size)+'_high_res.jpg')
    highres.save(name,dpi=(600,600))

def create_altered_text_images(text_pdf_path,runs=10,need_images=False):
    filename = os.path.basename(text_pdf_path)
    output, text_image_pages, final_directory = initialize_art_for_text(text_pdf_path,need_images=need_images)
    strip_output = "C:/Users/Vvayne/Documents/FA 2017/General/Outputs/" + filename + '_stripomanics/'
    if not os.path.exists(strip_output):
        os.makedirs(strip_output)
    cubo_output = "C:/Users/Vvayne/Documents/FA 2017/General/Outputs/" + filename + '_cubomanics/'
    if not os.path.exists(cubo_output):
        os.makedirs(cubo_output)
    print(output)
    print(final_directory)
    images_to_merge = get_imlist(output)
    print(len(images_to_merge))
    for i in range(runs):
        if isinstance(output,str):
        # print(output)
            output = output.split(" C")
        print('making stripomanic')
        create_stripomanic([images_to_merge, cubo_output], output_directory = strip_output)
        print('making cubomanic')
        create_cubomanic([images_to_merge, strip_output], output_directory=cubo_output)

        d1 = get_imlist(sample([strip_output,cubo_output],1)[0])

        image1 = None
        image2 = None
        fail_count = 0
        while not image1 or not image2:
            try:
                image1 = sample(text_image_pages, 1)[0]
                # print(image1)
                image2 = sample(d1, 1)[0]
                # print(image2)
            except:
                print('keep looking for images')
                fail_count += 1
                if fail_count >= 20:
                    break
        if image1 and image2:
            print('got images to blend')
            merge_images(image_path1=image1, image_path2=image2, size=(2700, 3600),
                         save_directory=final_directory)

directory = 'C:/Users/Vvayne/'
download_directory = 'C:/Users/Vvayne/Downloads/facebook-wayneyandell/'
test_directory = 'C:/Users/Vvayne/Documents/FA 2017/'
memoir_directory = 'C:/Users/Vvayne/Documents/FA 2017/Memori AntiMemoir/'
cubo_directory = 'C:/Users/Vvayne/Documents/FA 2017/General/Outputs/Cubomanics/'
# print(len(image_list))
other_directory = 'C:/Users/Vvayne/Pictures/Saved Pictures/'
pdf_jpgs = 'C:/Users/Vvayne/Pictures/images_from_pdfs/'
book_path = 'C:/Users/Vvayne/Pictures/images_from_pdfs/FA 2017/General/Inputs/'
            # '[Maggie_Ann_Bowers]_Magic(al)_Realism_(The_New_Cri(BookZZ.org)/'
derma_path = 'C:/Users/Vvayne/Pictures/cubopics/Derma Dentata/'
derma_cubo_path = 'C:/Users/Vvayne/Documents/FA 2017/General/Outputs/Cubomanics/Derma Dentata/'
picture_path = 'C:/Users/Vvayne/Pictures/cubopics/'
phone_path = 'C:/Users/Vvayne/Pictures/Camera Roll/phone/'
all_outputs = 'C:/Users/Vvayne/Documents/FA 2017/General/Outputs/'
l = 0
columns = 2700
rows = 3600
max_to_make = 50
output_category = 'Descent of Alette'
output_directory = 'C:/Users/Vvayne/Documents/FA 2017/General/Outputs/' + output_category + '/'
if not os.path.exists(output_directory):
    os.makedirs(output_directory)
eric_directories = [output_directory,os.path.join(cubo_directory,output_category +'/')]
derma_directories = [derma_path,derma_cubo_path,output_directory]
site = 'https://pixabay.com'

text_path = 'C:/Users/Vvayne/Documents/FA 2017/General/Inputs/personal/Writing/most_unabridged_collection_from_2013_to_110417.pdf'
text_path2 = 'C:/Users/Vvayne/Documents/FA 2017/'
max_to_make = 10
# input_directory = 'C:/Users/Vvayne/Documents/FA 2017/Memori AntiMemoir/Second Antimemoir Outputs/'
paths_to_use = 'C:/Users/Vvayne/Documents/FA 2017/'
input_directory = 'C:/Users/Vvayne/'
runs = 30
for i in range(max_to_make):
    create_stripomanic(get_imlist(input_directory), input_directory, columns=900, rows=1200, orientation='horizontal')
    image1 = None
    image2 = None
    fail_count = 0
    while not image1 or not image2:
        try:
            image1 = sample(get_imlist(directory=input_directory),1)[0]
            # print(image1)
            image2 = sample(get_imlist(directory=input_directory),1)[0]
            # print(image2)
        except:
            print('keep looking for images')
            fail_count +=1
            if fail_count >= 20:
                "no images or bad images"
                break

    if image1 and image2:
        print('got images to blend')
        merge_images(image_path1=image1,image_path2=image2,size=(900,1200),save_directory=input_directory)

# specific_path = input_directory + 'Poetry & Prose for Mem 2.pdf'
# for specific_path in get_filepaths(input_directory):
#     if specific_path.endswith('.pdf'):
#         output, text_image_pages, final_directory = initialize_art_for_text(specific_path,need_images=False)
#
#         # print(output)
#         # print(final_directory)
#
#
#
#         print(output)
#         print(final_directory)
#         print(final_directory.type)
#         print("makin a cubomanic")
#
#         print("now let's make a stripomanic")
#
#         d1 = input_directory
#
#         image1 = None
#         image2 = None
#         fail_count = 0
#         while not image1 or not image2:
#             try:
#                 image1 = sample(get_imlist(directory=d1),1)[0]
#                 # print(image1)
#                 image2 = sample(get_imlist(directory=final_directory),1)[0]
#                 # print(image2)
#             except:
#                 print('keep looking for images')
#                 fail_count +=1
#                 if fail_count >= 20:
#
#                     break
#         if image1 and image2:
#             print('got images to blend')
#             create_stripomanic([text_image_pages], d1, columns=900, rows=1200, orientation='horizontal')
#             merge_images(image_path1=image1,image_path2=image2,size=(2700,3600),save_directory=final_directory)

# now = datetime.datetime.now()
# #
# output = 'C:/Users/Vvayne/Documents/SP 2018/General/Outputs/' + now.strftime("%Y-%m-%d")
# if not os.path.exists(output):
#     os.makedirs(path=output)
# os.chdir(output)
# #
# for i in range(runs):
#     if isinstance(output, str):
#         # print(output)
#         output = output.split(" C")
#
#     # create_cubomanic([output_directory,text_image_pages,final_directory],output_directory,columns=900,rows=1200,conversions=True)
#
#     d1 = input_directory
#
#     image1 = None
#     image2 = None
#     fail_count = 0
#     while not image1 or not image2:
#         try:
#             image1 = sample(get_imlist(directory=d1), 1)[0]
#             # print(image1)
#             image2 = sample(get_imlist(directory=final_directory), 1)[0]
#             # print(image2)
#         except:
#             print('keep looking for images')
#             fail_count += 1
#             if fail_count >= 20:
#                 break
#     if image1 and image2:
#         print('got images to blend')
#         create_stripomanic([text_image_pages], d1, columns=900, rows=1200, orientation='horizontal')
#         merge_images(image_path1=image1, image_path2=image2, size=(900, 1200), save_directory=final_directory)
#
#     if isinstance(output, str):
#         # print(output)
#         output = output.split(" C")
#
#     print(output)
#     print(final_directory)
#     print(final_directory.type)
#     print("makin a cubomanic")
#     create_cubomanic([output_directory,text_image_pages,final_directory],output_directory,columns=900,rows=1200,conversions=True)
#     print("now let's make a stripomanic")
#
#     d1 = input_directory
#
#     image1 = None
#     image2 = None
#     fail_count = 0
#     while not image1 or not image2:
#         try:
#             image1 = sample(get_imlist(directory=d1), 1)[0]
#             # print(image1)
#             image2 = sample(get_imlist(directory=final_directory), 1)[0]
#             # print(image2)
#         except:
#             print('keep looking for images')
#             fail_count += 1
#             if fail_count >= 20:
#                 break
#     if image1 and image2:
#         print('got images to blend')
#         create_stripomanic([text_image_pages], d1, columns=900, rows=1200, orientation='horizontal')
#         merge_images(image_path1=image1, image_path2=image2, size=(900, 1200), save_directory=final_directory)


# specific_path = 'C:/Users/Vvayne/Documents/FA 2017/Derma Dentata Alt.pdf'
# output_directory, text_pages, final_output = initialize_art_for_text(specific_path)
# create_altered_text_images(specific_path,runs,need_images=False)

# purgatory_directory = 'C:/Users/Vvayne/Pictures/images_from_pdfs/Purgatory - Raul Zurita/'
# dreams_directory = 'C:/Users/Vvayne/Pictures/images_from_pdfs/Dreams for Kurosawa _ Zurita _ Anna Deeny-rotated'
# descent_directory ='C:/Users/Vvayne/Pictures/images_from_pdfs/the descent of alette/'
# image_texts = get_imlist(descent_directory)
# print(str(len(image_texts)))
# os.chdir(output_directory)
# for image_path in image_texts:
#     print(image_path)
#     text, im = get_text_from_image(image_path)
#     if text:
#         print('got text')
#         f = docx.Document()
#         p = f.add_paragraph()
#         text_in_lines = text.split('/n')
#         for line in text_in_lines:
#             if line[0].isalpha():
#                 p.add_run(line)
#             else:
#                 p = f.add_paragraph()
#         print('saving transcription')
#         text_name = os.path.basename(image_path).replace('.jpg','') + '_text_extracted.docx'
#         f.save(text_name)

