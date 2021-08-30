##program that takes one or more documents, parses out
import time
import docx
from docx.shared import Inches
import os
import re
from random import sample, randint, shuffle
from nltk import tokenize, pos_tag, FreqDist
import string
import shutil
def get_filepaths(directory):
    """
    This function will generate the file names in a directory
    tree by walking the tree either top-down or bottom-up. For each
    directory in the tree rooted at directory top (including top itself),
    it yields a 3-tuple (dirpath, dirnames, filenames).
    """
    file_paths = []  # List which will store all of the full filepaths.

    # Walk the tree.
    for root, directories, files in os.walk(directory):
        for filename in files:
            # Join the two strings in order to form the full filepath.
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)  # Add it to the list.

    return file_paths  # Self-explanatory.
def get_document_paths(file_paths):
    paths_for_potentials = []
    for path in file_paths:
        word = path
        if word.endswith('.docx') or word.endswith('.pdf') or word.endswith('.doc'):
            if not os.path.basename(path).startswith('~$'):
                paths_for_potentials.append(path)
    return paths_for_potentials
def extract_all_text_from_docx(filepath, paragraph_or_runs='runs'):
    try:
        doc = docx.Document(filepath)
    except:
        print("couldn't scrape it")
        all_runs = []
        return all_runs
    if doc:
        # print("attempting to scrape " + os.path.basename(filepath))
        file_fail_count = 0
        all_paragraphs = []
        all_runs = []
        length = len(doc.paragraphs)
        # print(str(length))
        if length >= 1:
            for paragraph in doc.paragraphs:
                all_paragraphs.append(paragraph)
                for run in paragraph.runs:
                    all_runs.append(run.text)
            if paragraph_or_runs == 'paragraphs':
                return all_paragraphs
            else:
                # print("got all " + str(length) + " run(s)")
                # print("sample: " + all_runs[0])
                return all_runs
def extract_sentence(path):
    print('trying to extract a sentence from ' + path)

    fail_count = 0
    sentence = ''
    # print('parsing ' + os.path.basename(path) + ' for a line')
    page_text = ""
    if path.endswith('.docx'):
        page_text = ""
        # print("trying to extract all text from this doc")

        doc_in_lines = extract_all_text_from_docx(path)
        # doc_in_sentences = tokenize.sent_tokenize(' '.join(doc_in_lines))
        # line_to_try = sample(doc_in_lines, 1)[0]
        while not sentence or len(tokenize.word_tokenize(sentence)) > 20 or len(tokenize.word_tokenize(sentence))  < 5\
                or re.search(r'"([A-Za-z0-9_\./\\-]*)"', sentence) or '"' in sentence or '“' in sentence or '”' in sentence:
                # sentences_from_run = tokenize.sent_tokenize(line_to_try)
                # sentence = sample(sentences_from_run,1)[0]
                # line_to_try = sample(doc_in_lines, 1)[0]
                candidate = sample(doc_in_lines, 1)[0]
                if len(tokenize.word_tokenize(sentence)) > 20 and len(tokenize.sent_tokenize(candidate)) > 1:
                    candidate = tokenize.sent_tokenize(candidate)[randint(0,len(tokenize.sent_tokenize(candidate))-1)]
                while candidate and not candidate[0].isalpha():
                    candidate = candidate[1:]
                # while candidate and (candidate[-1] not in end_punctuation and not candidate[-1].isalpha()):
                #     candidate = candidate[:-1]
                if candidate:
                    if candidate[-1].isalpha() or candidate[-1] not in ['.','?','!']:
                        candidate += sample(['.', '?', '.', '.', '.', '.', '.','.','.'], 1)[0]
                    sentence = candidate.capitalize()
                fail_count += 1
                print(fail_count)
                if fail_count > 20:
                    print("something wrong here")
                    sentence = ''
                    return sentence
    if sentence:
        sentence_to_add = sentence

        # print('got a sentence, checking to see if it contains a quote')
        # # pattern = r'"([A-Za-z0-9_\./\\-]*)"'
        # # quote = re.search(pattern, sentence)
        # # if quote:
        # if '"' in sentence:
        #     print('we got a quote here: ' + quote.group())
        #     list_of_quotes.append(quote.group())
        #     sentence = ''
        # print("nope, it's a free sentence")
        return sentence_to_add

    def generate_essay(input_directory):
        list_of_lines = []
        list_of_sentences = []
        list_of_quotes = []
        end_punctuation = ['.', '?', ')', "’", "'", ']']
        list_of_cohesive_ties = ['also', 'moreover', 'and', 'in addition', 'besides', 'as well', 'furthermore',
                                 'not only',
                                 'additionally', 'whereas', 'besides', 'but', 'however', 'in contrast', 'instead',
                                 'conversely',
                                 'it may be the case that', 'certainly', 'also', 'likewise', 'naturally',
                                 'nevertheless',
                                 'of course', 'on the contrary', 'on the other hand', 'regardless', 'granted', 'like',
                                 'alternatively', 'still', 'whereas', 'while', 'yet', 'although', 'despite',
                                 'it is true that',
                                 'notwithstanding', 'it may appear', 'regardless', 'certainly', 'granted that',
                                 'naturally',
                                 'I admit that', 'it may be the case that', 'admittedly', 'all things considered',
                                 'as a general rule', 'as far as we know', 'astonishingly', 'broadly', 'by and large',
                                 'characteristically', 'clearly', 'coincidentally', 'conveniently', 'curiously',
                                 'disappointingly', 'equally', 'essentially', 'explicitly', 'even so', 'eventually',
                                 'fortunately', 'fundamentally', 'generally speaking', 'interestingly', 'ironically',
                                 'in essence', 'in general', 'in particular', 'in practice', 'in reality',
                                 'in retrospect', 'in hindsight', 'in theory', 'in view of this', 'more interestingly',
                                 'more seriously', 'more specifically', 'naturally', 'on balance', 'obviously',
                                 'on reflection',
                                 'overall', 'paradoxically', 'potentially', 'predictably', 'presumably', 'primarily',
                                 'probably',
                                 'remarkably', 'seemingly', 'significantly', 'surprisingly', 'theoretically',
                                 'to all intents and purposes', 'typically', 'ultimately', 'understandably',
                                 'undoubtedly',
                                 'unfortunately', 'with hindsight']

        number_of_paragraphs = randint(28, 33)
        f = docx.Document()

        # get filepaths in input directory
        doc_paths = get_document_paths(get_filepaths(input_directory))
        print('number of sourcetexts: ' + str(len(doc_paths)))
        lines_added = []
        title = ''
        while not title:
            title = extract_sentence(sample(doc_paths, 1)[0])
        print('the title for this essay will be ' + title)
        f.add_heading(title, 0)

        for i in range(number_of_paragraphs):
            print('populating paragraph ' + str(i))
            sent_per_para = randint(4, 9)
            p = f.add_paragraph()
            paragraph_sentences = []
            while len(paragraph_sentences) < sent_per_para:
                sentence_to_add = extract_sentence(sample(doc_paths, 1)[0])
                if sentence_to_add and sentence_to_add[:-1] not in lines_added:
                    lines_added.append(sentence_to_add)
                    cohesive_tie_index = randint(0, 4)
                    if cohesive_tie_index == 2:
                        print("let's add a cohesive tag")
                        cohesive_tie = sample(list_of_cohesive_ties, 1)[0].capitalize()
                        sentence_to_add = cohesive_tie + ', ' + sentence_to_add[0].lower() + ''.join(
                            sentence_to_add[1:])
                    else:
                        sentence_to_add = sentence_to_add
                    paragraph_sentences.append(sentence_to_add)
                    print('sentence ' + str(len(paragraph_sentences)) + ' added')
                else:
                    print('either wrote this one already or fake sentence')

            p.add_run(' '.join(paragraph_sentences))
            print('paragraph ' + str(i) + ' added')
        clean_title = re.sub(r'[^\w\s]', '', title)
        paper_name = clean_title + '.docx'

        output_directory = 'C:/Users/Vvayne/Documents/FA 2018/Outputs/' + os.path.basename(input_directory)
        if not os.path.exists(output_directory):
            os.makedirs(output_directory)
        os.chdir(output_directory)
        if os.path.isfile(paper_name):
            paper_name = title + "_" + str(randint(1, 500)) + '.docx'
        f.save(paper_name)
        print('all done! you can find ' + title + '.docx @ ' + output_directory)

#initialize variables
input_directory = 'C:/Users/Vvayne/Documents/FA 2018/Multigenre Workshop/Workshop/'
