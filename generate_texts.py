##program that takes one or more documents, parses out
import time
import docx
from docx.shared import Inches
import os
import re
from random import sample, randint, shuffle, choice
from nltk import tokenize, pos_tag, FreqDist
import nltk
import string
import shutil
english_vocab = set(w.lower() for w in nltk.corpus.words.words())

def get_filepaths(directory):
    """
    This function will generate the file names in a directory
    tree by walking the tree either top-down or bottom-up. For each
    directory in the tree rooted at directory top (including top itself),
    it yields a 3-tuple (dirpath, dirnames, filenames).
    """
    file_paths = []  # List which will store all of the full filepaths.

    # Walk the tree.
    for root, directories, files in os.walk(directory):
        for filename in files:
            # Join the two strings in order to form the full filepath.
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)  # Add it to the list.

    return file_paths  # Self-explanatory.
def get_document_paths(file_paths):
    paths_for_potentials = []
    for path in file_paths:
        word = path
        if word.endswith('.docx') or word.endswith('.pdf') or word.endswith('.doc'):
            if not os.path.basename(path).startswith('~$'):
                paths_for_potentials.append(path)
    return paths_for_potentials
def extract_all_text_from_docx(filepath, paragraph_or_runs='runs'):
    try:
        doc = docx.Document(filepath)
    except:
        return []
    #     print("couldn't scrape it")
    #     all_runs = []
    #     return all_runs
    if doc:
        # print("attempting to scrape " + os.path.basename(filepath))
        file_fail_count = 0
        all_paragraphs = []
        all_runs = []
        length = len(doc.paragraphs)
        # print(str(length))
        if length >= 1:
            for paragraph in doc.paragraphs:
                all_paragraphs.append(paragraph)
                for run in paragraph.runs:
                    all_runs.append(run.text)
            if paragraph_or_runs == 'paragraphs':
                return all_paragraphs
            else:
                # print("got all " + str(length) + " run(s)")
                # print("sample: " + all_runs[0])
                return all_runs

def remove_repeated_words(string):
    clean_string = []
    string_in_words = string.split()
    for i in range(string_in_words):
        if not string_in_words[i] == string_in_words[i+1]:
            clean_string.append(string_in_words[i])
    return clean_string

def extract_sentence(path):
    print('trying to extract a sentence from ' + path)
    end_punctuation = ['.', '?', ')', "'", "'", ']', '--', ';', ':']
    fail_count = 0
    sentence = ''
    print('parsing ' + os.path.basename(path) + ' for a line')
    page_text = ""
    if path.endswith('.docx'):
        page_text = ""
        print("trying to extract all text from this doc")
        try:
            doc_in_lines = extract_all_text_from_docx(path)
        except:
            print('this doc failed')
            return ''
        # doc_in_sentences = tokenize.sent_tokenize(' '.join(doc_in_lines))
        # line_to_try = sample(doc_in_lines, 1)[0]
        while not sentence or len(tokenize.word_tokenize(sentence)) > 20 or len(tokenize.word_tokenize(sentence))  < 5\
                or re.search(r'"([A-Za-z0-9_\./\\-]*)"', sentence) or '"' in sentence or '“' in sentence or '”' in sentence:
                # sentences_from_run = tokenize.sent_tokenize(line_to_try)
                # sentence = sample(sentences_from_run,1)[0]
                # line_to_try = sample(doc_in_lines, 1)[0]
                try:
                    candidate = sample(doc_in_lines, 1)[0]
                    if len(tokenize.word_tokenize(sentence)) > 20 and len(tokenize.sent_tokenize(candidate)) > 1:
                        candidate = tokenize.sent_tokenize(candidate)[randint(0,len(tokenize.sent_tokenize(candidate))-1)]
                    while candidate and not candidate[0].isalpha():
                        candidate = candidate[1:]
                    while candidate and (candidate[-1] not in end_punctuation and not candidate[-1].isalpha()):
                        candidate = candidate[:-1]
                    if candidate:
                        if candidate[-1].isalpha() or candidate[-1] not in ['.','?','!']:
                            candidate += choice(end_punctuation)
                        sentence = candidate.capitalize()
                except:
                    sentence = ''
                    return sentence
                fail_count += 1
                # print(fail_count)
                if fail_count > 20:
                    print("something wrong here")
                    sentence = ''
                    return sentence
    if sentence:
        sentence_to_add = sentence

        # print('got a sentence, checking to see if it contains a quote')
        # # pattern = r'"([A-Za-z0-9_\./\\-]*)"'
        # # quote = re.search(pattern, sentence)
        # # if quote:
        # if '"' in sentence:
        #     print('we got a quote here: ' + quote.group())
        #     list_of_quotes.append(quote.group())
        #     sentence = ''
        # print("nope, it's a free sentence")
        return sentence_to_add

def strip_numbering_tags(string):
    # print("let's test for numbering/sectional tags so they don't get lumped in as lines")
    s = string
    if not isinstance(string, str):
        s = string.text
    if re.search('[iv]+\.', s):
        lst = re.findall('[iv]+\.', s)
        for tag in lst:
            s.replace(tag, "")
    elif re.search('\d+\.', s):
        lst = re.findall('\d+\.', s)
        for tag in lst:
            s.replace(tag, "")
            # s.replace([i for i in lst],"")
    return s

def generate_paragraphs(input_directory,number_of_paragraphs=randint(5,10)):
    # get filepaths in input directory
    doc_paths = get_document_paths(get_filepaths(input_directory))
    print('number of strophes: ' + str(number_of_paragraphs) + ' |  number of sourcetexts: ' + str(len(doc_paths)))
    paragraphs = []
    lines_added = []
    for i in range(number_of_paragraphs):
        print('populating strophe ' + str(i))
        sent_per_para = randint(4,9)
        paragraph_sentences = []
        while len(paragraph_sentences) < sent_per_para:
            sentence_to_add = extract_sentence(choice(doc_paths))
            if sentence_to_add and sentence_to_add[:-1] not in lines_added:
                sent_no_numbers = []
                for string in sentence_to_add.split():
                    new_string = strip_numbering_tags(string)
                    sent_no_numbers.append(new_string)
                lines_added.append(' '.join(sent_no_numbers))
                ulist = []
                [ulist.append(x) for x in sentence_to_add.split() if x not in ulist]
                paragraph_sentences.append(' '.join(ulist))
                print('sentence ' + str(len(paragraph_sentences)) + ' of ' + str(sent_per_para) + ' added')
            else:
                print('either wrote this one already or fake sentence')

        
        paragraphs.append(paragraph_sentences)
        print('strophe ' + str(i) + ' added')
    print('finished strophe generation')
    return paragraphs

def write_list_of_text_units(text_list,output_directory,titles=False):
    print("now let's write these strophes to a document")
    f = docx.Document()
    print("first let's generate a title")
    title = generate_phrase()
    print("the title will be " + title)
    filename = title + '.docx'
    f.add_heading(title, 0)
    for i in range(len(text_list)):
        print('writing strophe ' + str(i) + ' of ' + str(len(text_list)))
        if titles:
            subtitle = text_list[0][0]
            f.add_heading(subtitle,2)
            del text_list[0][0]
        p = f.add_paragraph()
        p.add_run(' '.join(text_list[i]))
        p.add_run().add_break()
        print('strophe ' + str(i) + ' added')
    print('all strophes written to the document')
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    os.chdir(output_directory)
    if os.path.isfile(filename):
        filename = title + "_" + str(randint(1, 500)) + '.docx'
    print('saving document')
    f.save(filename)
    print('all done! you can find ' + filename + ' @ ' + output_directory)

def generate_line_from_seed(corpus, cfdist, word, num_of_words=15):
    line_list = []
    fail_count = 0
    while len(line_list) < num_of_words:
        while word in line_list or word not in english_vocab:
            word = choice(list(cfdist[word].most_common()))[0]
            fail_count += 1
            if fail_count >= len(list(cfdist[word])):
                # print(pos_tag([word]))
                # print(word)
                token = tokenize.word_tokenize(word)
                # print(token)
                tagged_token = pos_tag(token)
                # print(tagged_token)
                tag = ' '.join([x[1] for x in tagged_token])
                # print(tag)
                # print(tag)

                word = choice(list_of_pos(corpus,tag))
        line_list.append(word)
            # if len(list(cfdist[word])) > 2:
            #     word = choice(cfdist[word].most_common(3))
            # else:



    line = ' '.join(line_list)
    return line

def generate_bigrams_cfd(text):
    bigrams = nltk.bigrams(text)
    cfd = nltk.ConditionalFreqDist(bigrams)
    return bigrams,cfd

def generate_trigrams_cfd(text):
    text_in_words = tokenize.word_tokenize(text)
    trigrams = []
    for i in range(text_in_words):
        trigrams.append((text_in_words[i-1],text_in_words[i],text_in_words[i+1]))
    cfd = nltk.ConditionalFreqDist(trigrams)
    return trigrams, cfd

# def generate_pre_words(word):
    
def create_corpus(input_directory):
    f = docx.Document()
    list_of_fulldoc_strings = []
    doc_paths = get_document_paths(get_filepaths(input_directory))
    intermediary_paths = [path for path in doc_paths if path.endswith('.docx')]
    paths = []
    for address in intermediary_paths:
        if not os.path.split(address)[-1] in ' '.join(paths):
            paths.append(address)
    for path in paths:
        try:
            all_text = extract_all_text_from_docx(path)
            if all_text:
                list_of_fulldoc_strings.append(' '.join(all_text))
        except:
            continue
    corpus_string = ' '.join(list_of_fulldoc_strings)
    corpus_list = re.findall(r"[\w']+|[.,!?;]",corpus_string)
    return corpus_list

def list_of_pos(corpus,tag):
    word_list = []
    words_and_tags = nltk.pos_tag(corpus)
    for a,b in words_and_tags:
        if tag in b:
            word_list.append(a)
    return word_list

def generate_phrase(n = randint(4,9)):
    phrase = ''
    while not phrase:
        source_path = ''
        while not source_path or not source_path.endswith('.docx') or not source_path[0].isalnum():
            source_path = choice(get_document_paths(get_filepaths("C:/Users/Vvayne/Documents/")))
            # print(source_path)
        print('looking for words to make a title in ' + os.path.split(source_path)[-1])
        text_in_words = ''
        number_of_words = 0
        # try:
        all_text = extract_all_text_from_docx(source_path)
        if all_text:
            # print(type(all_text))
            # print(' '.join(all_text)[0:20])
            text_in_words = tokenize.word_tokenize(' '.join(all_text))
            # print(text_in_words[0:3])
            number_of_words = len(text_in_words)
            print(str(number_of_words))
        # except:
        #     print("couldn't get a phrase from " + source_path)
        #     continue

        if text_in_words and number_of_words > n:
            phrase = ' '.join(sample(text_in_words,n))

    clean_phrase = re.sub(r'[^\w\s]', '', phrase)
    clean_phrase = ' '.join(clean_phrase.split())
    print('phrase generated: ' + clean_phrase)
    return clean_phrase

def generate_lines_from_corpus(input_directory,number_of_lines = 12,seed_type='most common'):
    print("let's create a corpus from which to generate")
    current_corpus = create_corpus(input_directory)
    current_corpus_words = [x for x in tokenize.word_tokenize(' '.join(current_corpus)) if x in english_vocab]
    # stopwords = nltk.corpus.stopwords.words('english')
    # current_corpus_words = [word for word in current_corpus_words if word not in stopwords]
    if current_corpus_words:
        print(len(current_corpus_words))
        print("now let's produce a list of bigrams and conditional frequency distribution of them")
        bigrams, cfd = generate_bigrams_cfd(current_corpus_words)
        print("now let's initialize and find a seed word")
        used_seeds = []
        lines_generated = []
        for i in range(number_of_lines):
            seed_word = ''
            fail_counter = 0
            n = 10
            while not seed_word or seed_word in used_seeds:
                if seed_type == 'most common':
                    fdist1 = FreqDist(current_corpus_words)
                    top_words = [x for x, y in fdist1.most_common(n) if len(x) > 3]
                    seed_word = choice(top_words)
                    fail_counter += 1
                    if fail_counter > number_of_lines:
                        n += 5
                if seed_type == 'noun':
                    seed_word = choice(list_of_pos(current_corpus_words,'NN'))

            if seed_word:
                print("this seed word will be " + seed_word)
                used_seeds.append(seed_word)
                line_length = randint(3,20)
                print("now let's generate a line with it")
                new_line = generate_line_from_seed(current_corpus_words, cfd,seed_word, line_length)
                if new_line and not new_line in lines_generated:
                    print("it's a new line")
                    lines_generated.append(new_line)
        print("all lines generated")
        return lines_generated

def write_lines_in_strophes(lines,output_directory,number_of_strophes = 0,lines_per_strophe=0):
    f = docx.Document()
    title = choice(lines)
    filename = title + '.docx'
    f.add_heading(title,2)
    lines.remove(title)
    if number_of_strophes == 0:
        number_of_strophes = randint(1, len(lines))
    if lines_per_strophe == 0:
        lines_per_strophe = int(len(lines)/number_of_strophes)
    for i in range(number_of_strophes):
        p = f.add_paragraph()
        new_lines = sample(lines,lines_per_strophe)
        for line in new_lines:
            p.add_run(line)
            p.add_run().add_break()

        lines = [x for x in lines if not x in new_lines]
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    os.chdir(output_directory)
    if os.path.isfile(filename):
        filename = title + "_" + str(randint(1, 500)) + '.docx'
    print('saving document')
    f.save(filename)
    print('all done! you can find ' + filename + ' @ ' + output_directory)


prepositions = ['aboard','about','above','across','after','against','along','amid','among','anti','around','as','at',
                'before','behind','below','beneath','beside','besides','between','beyond','but','by','concerning',
                'considering','despite','down','during','except','excepting','excluding','following','for','from','in',
                'inside','into','like','minus','near','of','off','on','onto','opposite','outside','over','past','per',
                'plus','regarding','round','save','since','than','through','to','toward','towards','under','underneath',
                'unlike','until','up','upon','versus','via','with','within','without']

sub_cons = ['after','in order (that)','unless','although','insofar as','until','as','in that','when','as far as','lest',
            'whenever','as soon as','no matter how','where','as if','now that','wherever','as though','once','whether',
            'because','provided (that)','while','before','since','why','even if','so that','even though',
            'supposing (that)','how','than','if','that','inasmuch as','though','in case (that)','till']

con_adverbs = ['after all','in addition','next','also','incidentally','nonetheless','as a result','indeed',
               'on the contrary','besides','in fact','on the other hand','consequently','in other words','otherwise',
               'finally','instead','still','for example','likewise','then','furthermore','meanwhile','therefore',
               'hence','moreover','thus','however','nevertheless']

coord_cons = ['for', 'and', 'nor', 'but', 'or', 'yet', 'so']

articles = ['a','an','the']


#initialize directories
# input_directory = 'C:/Users/Vvayne/Documents/'
input_folder = 'C:/Users/Vvayne/Documents/FA 2018/Multigenre Workshop/Workshop/'
output_folder = 'C:/Users/Vvayne/Documents/FA 2018/Multigenre Workshop/Outputs/' + os.path.basename(input_folder)


number_o_lines = 10
gen_lines = generate_lines_from_corpus(input_folder,number_of_lines=randint(5,15),seed_type='noun')
if gen_lines:
    write_lines_in_strophes(gen_lines,output_folder)
#generate paragraphs
# paragraphs = generate_paragraphs(input_directory)
# #
# # #write paragraphs to word doc
# write_list_of_text_units(paragraphs,output_directory,True)

# phrase = generate_phrase()
# print(phrase)





