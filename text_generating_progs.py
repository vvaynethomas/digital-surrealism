import os
from random import randint,sample, shuffle,choice
import nltk
import time
import docx
import re
from docx.shared import Inches
from nltk import tokenize, pos_tag, FreqDist
import PyPDF2
from bs4 import BeautifulSoup, SoupStrainer

english_vocab = set(w.lower() for w in nltk.corpus.words.words())
prepositions = ['aboard','about','above','across','after','against','along','amid','among','anti','around','as','at',
                'before','behind','below','beneath','beside','besides','between','beyond','but','by','concerning',
                'considering','despite','down','during','except','excepting','excluding','following','for','from','in',
                'inside','into','like','minus','near','of','off','on','onto','opposite','outside','over','past','per',
                'plus','regarding','round','save','since','than','through','to','toward','towards','under','underneath',
                'unlike','until','up','upon','versus','via','with','within','without']
sub_cons = ['after','in order (that)','unless','although','insofar as','until','as','in that','when','as far as','lest',
            'whenever','as soon as','no matter how','where','as if','now that','wherever','as though','once','whether',
            'because','provided (that)','while','before','since','why','even if','so that','even though',
            'supposing (that)','how','than','if','that','inasmuch as','though','in case (that)','till']
con_adverbs = ['after all','in addition','next','also','incidentally','nonetheless','as a result','indeed',
               'on the contrary','besides','in fact','on the other hand','consequently','in other words','otherwise',
               'finally','instead','still','for example','likewise','then','furthermore','meanwhile','therefore',
               'hence','moreover','thus','however','nevertheless']
coord_cons = ['for', 'and', 'nor', 'but', 'or', 'yet', 'so']
articles = ['a','an','the']
end_punct = ['.', '?',')', "’", "'", ']','"',':']

def primes(n):
    primfac = []
    d = 2
    while d*d <= n:
        while (n % d) == 0:
            primfac.append(d)  # supposing you want multiple factors repeated
            n //= d
        d += 1
    if n > 1:
       primfac.append(n)
    return primfac

def partition_text_in_words(text_in_words,chunk_length=randint(2,6)):
    l = text_in_words
    n = chunk_length
    """Yield successive n-sized chunks from l."""
    chunk_list = []
    for i in range(0, len(l), n):
        chunk_list.append(l[i:i + n])
    print(chunk_list[0])
    return chunk_list

def extract_from_pdf(filepath,type='page',subtype='string'):
    # filename = os.path.basename(filepath)
    with open(filepath, 'rb') as pdf_file_obj:  # open current pdf
        try:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)  # store pdf as PyPDF2 reader object
            if pdf_reader and not pdf_reader.isEncrypted:
                # print("I'm looking for a line in " + filename)
                number_of_pages = pdf_reader.getNumPages()  # count number of pages in pdf
                if number_of_pages > 0:
                    if type == 'all' :
                        text_in_lines = []
                        text_in_pages = []
                        for i in range(0,number_of_pages - 1):
                                page_obj = pdf_reader.getPage(i)
                                page_as_list = page_obj.extractText().splitlines()
                                text_in_pages.append(page_as_list)
                                for line in page_as_list:
                                    text_in_lines.append(line)
                        return text_in_pages, text_in_lines
                    elif type == 'page':
                        page_number = randint(0,number_of_pages-1)
                        page_obj = pdf_reader.getPage(page_number)
                        if subtype == 'lines':
                            page_in_lines = page_obj.extractText().splitlines()
                            # text_in_pages = [text_in_lines]
                            return page_in_lines, page_number
                        else:
                            page_as_string = page_obj.extractText()
                            return page_as_string, page_number

                    elif type == 'line':
                        fail_count = 0
                        line = ""
                        page_num = 0
                        while not line and fail_count < number_of_pages:
                            page_num = randint(1, number_of_pages - 1)  # generate random index to choose page
                            page_obj = pdf_reader.getPage(page_num)  # reads and stores random page as page object
                            try:
                                page_in_lines = page_obj.extractText().splitlines()  # stores extracted text as multiline string
                                num_lines = len(page_in_lines)  # count number of lines detected
                                if num_lines > 0:
                                    print("we've got lines here")
                                    line = sample(page_in_lines,1)[0]
                                    return line, page_num
                                else:
                                    print("no lines detected on page " + page_num)
                                    line = ""
                                    fail_count +=1
                                    continue

                            except:
                                print('failed to extract a line from pdf page ' + page_num)
                                fail_count +=1
                                continue
                        return line, page_num

                    elif type == 'sentence':
                        fail_count = 0
                        sentence = ""
                        page_num = 0
                        while not sentence and fail_count < number_of_pages:
                            page_num = randint(0, number_of_pages-1)
                            page_obj = pdf_reader.getPage(page_num)
                            try:
                                page_as_string = page_obj.extractText()
                                page_in_sentences = tokenize.sent_tokenize(page_as_string)
                                sentence = sample(page_in_sentences,1)[0]
                            except:
                                print('failed to extract a sentence from pdf page ' + str(page_num))
                                fail_count +=1
                                continue
                        return sentence, page_num
        except:
            print('unable to parse pdf')

def get_filepaths(directory):
    """
    This function will generate the file names in a directory
    tree by walking the tree either top-down or bottom-up. For each
    directory in the tree rooted at directory top (including top itself),
    it yields a 3-tuple (dirpath, dirnames, filenames).
    """
    file_paths = []  # List which will store all of the full filepaths.

    # Walk the tree.
    for root, directories, files in os.walk(directory):
        for filename in files:
            # Join the two strings in order to form the full filepath.
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)  # Add it to the list.

    return file_paths  # Self-explanatory.

def get_document_paths(directory):
    file_paths = []  # List which will store all of the full filepaths.

    # Walk the tree.
    for root, directories, files in os.walk(directory):
        for filename in files:
            # Join the two strings in order to form the full filepath.
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)  # Add it to the list.

    paths_for_potentials = []
    for path in file_paths:
        if path.endswith('.docx') or path.endswith('.pdf') or path.endswith('.doc'):
            if not os.path.basename(path).startswith('~$'):
                paths_for_potentials.append(path)
    return paths_for_potentials

def extract_sentence(path):
    fail_count = 0
    sentence = ''
    page_number=0
    # print('parsing ' + os.path.basename(path) + ' for a line')
    page_text = ""
    if path.endswith('.pdf'):
        while not sentence:
            # print("scraping a page of text (after " + str(fail_count) + " tries)")
            if fail_count > 20:
                print("couldn't get a sentence")
                return sentence
            # try:
            page_text,page_number = extract_from_pdf(path,subtype='lines')
            page_in_sentences = tokenize.sent_tokenize(' '.join(page_text))
            if len(page_in_sentences) > 3:
                sentence = sample(page_in_sentences,1)[0]
            else:
                try:
                    sentence = sample(page_text,1)[0]
                except:
                    fail_count+=1

            fail_count +=1
    elif path.endswith('.docx'):
        fail_count = 0
        page_text = ""
        # print("trying to extract all text from this doc")
        while not sentence:
            doc_in_lines = extract_all_text_from_docx(path)
            if not doc_in_lines:
                print('scraped nothing from' + path)
                return
            doc_in_sentences = tokenize.sent_tokenize(' '.join(doc_in_lines))
            sentence = sample(doc_in_sentences,1)[0]
            if len(tokenize.word_tokenize(sentence)) > 20 or len(tokenize.word_tokenize(sentence)) < 2:
                sentence = sample(doc_in_lines,1)[0]
            fail_count +=1
            if fail_count > 20:
                print("something wrong here")
                return sentence,page_number

    if sentence:
        # print("got a sentence")
        return sentence,page_number

def extract_all_text_from_docx(filepath, paragraph_or_runs='runs'):
    try:
        doc = docx.Document(filepath)
    except:
        print("couldn't scrape it")
        all_runs = []
        return all_runs
    if doc:
        # print("attempting to scrape " + os.path.basename(filepath))
        file_fail_count = 0
        all_paragraphs = []
        all_runs = []
        length = len(doc.paragraphs)
        # print(str(length))
        if length >= 1:
            for paragraph in doc.paragraphs:
                all_paragraphs.append(paragraph)
                for run in paragraph.runs:
                    all_runs.append(run.text)
            if paragraph_or_runs == 'paragraphs':
                return all_paragraphs
            else:
                # print("got all " + str(length) + " run(s)")
                # print("sample: " + all_runs[0])
                return all_runs

def refrain_or_no(all_lines_so_far): # function to determine if the next line should be a refrain and if so what should it be
    if len(all_lines_so_far) > 2: #test for at least two lines before we start getting crazy with repetition
        source_for_lines = all_lines_so_far #store a copy of the lines
        roll_of_die = randint(1,len(source_for_lines)) #generate random integer between 1 and the number of available lines for refrain
        if roll_of_die >= randint(7,9)*.1*len(source_for_lines): #probabilistic test to confirm refrain
            line = sample(all_lines_so_far, 1)[0] #store refrain as line
            return 'yes', line  #return string and line
        else: #if refrain 'failed'
            line = "" #store nothing as line
            return 'no', line #return string and empty line
    else: # if 2 or fewer lines
        line = "" #store nothing as line
        return 'no', line #return no and empty line

def get_imlist(directory):
    all_paths = []
    try:
        for path in get_filepaths(directory):
            if path.endswith('.jpg') and not os.path.basename(path) in all_paths:
                all_paths.append(path)
    except:
        for folder in directory:
            for path in get_filepaths(folder):
                if not os.path.basename(path) in all_paths and path.endswith('.jpg'):
                    all_paths.append(path)
    return all_paths

def get_all_subdirectories(paths):
    dir_list = []
    for path in paths:
        sub_dir = os.path.dirname(path)
        dir_list.append(sub_dir)
            # print(path)
            # print(sub_dir)

    return dir_list

def write_vignettes(paths_to_use, output_path, set=True, number_of_vignettes=randint(5,20),image_directory=None,end_punct=end_punct):
    print("let's try to write " + str(number_of_vignettes) + " prose vignettes with these texts")
    if not set:
        vignette_numbers = sample(range(3,number_of_vignettes+3), number_of_vignettes)
    else:
        vignette_numbers = [randint(3,15)]*number_of_vignettes
    sections_and_sources = {}
    vignettes = {}
    f = docx.Document()
    f.add_heading(str(number_of_vignettes) + " Collage/Assemblage Vignettes")
    print((str(number_of_vignettes) + " Collage/Assemblage Vignettes"))
    for i in range(number_of_vignettes):
        key = str(i)
        print(str(i))
        number_of_sentences = sample(vignette_numbers,1)[0]
        current_sentences = []
        total_extracted = 0
        remaining_sentences = number_of_sentences - total_extracted
        while len(current_sentences) < number_of_sentences:
            sentence = ""
            while not sentence or not any(c.isalpha() for c in sentence):
                # try:
                if len(paths_to_use) > 1:
                    random_path = sample(paths_to_use, 1)[0]
                else:
                    random_path = paths_to_use[0]
                # print(random_path)
                # print(remaining_sentences)
                # try:
                sentence,page_number = extract_sentence(random_path)
                # except:
                # print(os.path.basename(random_path) + ' failed')
                continue

            if sentence and len(sentence) > 1:
                # print(sentence)
                words_in_sentence = tokenize.word_tokenize(sentence)
                if len(words_in_sentence) > 25 or len(words_in_sentence) < 2:
                    sentence = ' '.join(sample(partition_text_in_words(words_in_sentence,randint(7,20)),1)[0])

                end_punct = ['!','?','.',';','--',':','/']

                week_days = ['Sun ','Mon ','Tue ','Wed ','Thu ','Fri ']
                months = ['Jan ','Feb ','Mar ','Apr ','Jun ','Jul','Aug ','Sep ','Oct ','Nov ','Dec ']
                try:
                    while not sentence[-1].isalpha():
                        sentence = sentence[:len(sentence)-1]
                except:
                    continue
                # while not sentence[-1].isalpha():
                #     sentence = sentence[:-1]
                # if not sentence[-1] in end_punct:
                #     # sentence += sample(end_punct,1)[0]
                #     sentence += '.'
                outer_punct = [['"','" '],['[','] '],['{','} '],['(',') '],['...','...']]
                chosen_punct = sample(outer_punct,1)[0]
                pattern = r'\[.*?\]'
                sentence = re.sub(pattern,'',sentence)
                for c in [str(i) for i in range(0,100)]+week_days+months:
                    sentence = sentence.replace(c,'')
                try:
                    sentence = time.strptime(sentence, '%H:%M:%S')
                except:
                    # print('no time to strip')
                    this = False
                try:
                    while sentence[0] and not sentence[0].isalpha():
                        sentence = sentence[1:]

                    sentence = chosen_punct[0] + sentence + chosen_punct[1]
                    current_sentences.append(sentence.capitalize())
                    if not random_path in sections_and_sources.keys():
                        sections_and_sources[(os.path.basename(random_path)[:len(random_path)-4])] = page_number
                # print("We're adding this line")
                #     current_sentences.append(sentence)
                # store_source(sources, os.path.basename(random_path), page, sections_and_sources,
                #                                   key)
                # print("we've stored the sentence(s) and source(s)")
                # print("total sentence(s) extracted at this stage = " + str(total_extracted))
                    remaining_sentences -= 1
                except:
                    print('bad sentence')
                shuffle(current_sentences)
                # if not random_path + ', p.' + page_number in sources:

        if not len(current_sentences) >= number_of_sentences:
            continue
        else:
            print(str(len(current_sentences)) + " out of " + str(number_of_sentences) + " sentences being added")
            if key not in vignettes:
                vignettes[key] = [" ".join(current_sentences)]
            else:
                vignettes[key].append([" ".join(current_sentences)])
    print('compiling vignette')
    section_break_symbol = sample(end_punct,1)[0]*randint(4,30)
    im_list = None
    if image_directory:
        im_list = get_imlist(image_directory)
    for key in vignettes:
        print(key)
        f.add_heading(section_break_symbol)
        if im_list:
            image_added = False
            while not image_added:
                try:
                    picture_path = sample(im_list,1)[0]
                    if isinstance(picture_path,str):
                        f.add_picture(picture_path,width=Inches(4))
                        image_added = True
                except:
                    print("failed to insert " + picture_path )

        f.add_paragraph(vignettes[key])
        p = f.add_paragraph()
        p.add_run().add_break()

    f.add_page_break()
    f.add_heading('Sources', 1)
    # last_p = f.add_paragraph()
    for x in sections_and_sources.keys():
        p = f.add_paragraph(x)
        if not sections_and_sources[x] == 0:
            p.add_run(str(sections_and_sources[x]))

        # for x in sources.keys():
        #     f.add_paragraph(str(x) + " " + str(sources[x]))
    # f.write(sources)
    assemblage_name = 'prose_poem' + str(len(vignette_numbers)) + '.docx'
    os.chdir(output_path)
    if os.path.isfile(assemblage_name):
        assemblage_name = 'prose_poem' + str(len(vignette_numbers)) + "_" + str(randint(1, 500)) + '.docx'

    f.save(assemblage_name)

    print("Complete! Check out: " + assemblage_name)

def write_paragraph_shuffle(doc_path, output_path):
    print(doc_path)
    doc = docx.Document(doc_path)
    full_text = ''
    text_in_para = []
    for para in doc.paragraphs:
        if para.text:
            text_in_para.append(para.text)
    # full_text = '\n'.join(text_in_para)
    # print(full_text)
    shuffle(text_in_para)
    f = docx.Document()
    for para in text_in_para:
        f.add_paragraph(para)
    source_name = os.path.splitext(os.path.basename(doc_path))[0]
    shuffle_name = source_name + str(randint(1000,7000))
    os.chdir(output_path)
    while os.path.isfile(shuffle_name +'.docx'):
        shuffle_name = source_name + str(randint(1000,7000))
    f.save(shuffle_name + '.docx')

def write_shuffle(input_directory, output_directory, number_of_paragraphs = 0, repetition=False):
    doc_list = get_document_paths(input_directory)
    corpus_in_docpara = []
    corpus_in_docstring = []
    corpus_in_para = []
    corpus_in_sent = []
    corpus_in_parasent = []
    corpus_in_string = ''
    for doc_path in doc_list:
        print('extracting text from ' + os.path.basename(doc_path))
        doc = docx.Document(doc_path)
        doc_in_para = []
        for paragraph in doc.paragraphs:
            doc_in_para.append(paragraph.text)
            corpus_in_para.append(paragraph.text)
            para_in_sent = tokenize.sent_tokenize(paragraph.text)
            for sent in para_in_sent:
                corpus_in_sent.append(sent)
                corpus_in_string += sent
            corpus_in_parasent.append(para_in_sent)
        corpus_in_docpara.append(doc_in_para)
        corpus_in_docstring.append('\n'.join(doc_in_para))

    title = ' '.join(sample(tokenize.word_tokenize(corpus_in_string), randint(2, 6)))
    f = docx.Document()
    print("number of paragraphs in source corpus: " + str(len(corpus_in_para)))
    while number_of_paragraphs <= 0:
        print('picking how many paragraphs to write')
        option_1 = max([len(x) for x in corpus_in_docpara])
        option_2 = sum([len(x) for x in corpus_in_docpara])/len(corpus_in_docpara)
        option_3 = randint(2,sum([len(x) for x in corpus_in_docpara]))
        number_of_paragraphs = int(choice([option_1,option_2,option_3]))
    print("let's write " + str(number_of_paragraphs) + " shuffled paragraphs")


    all_para_lengths = [len(x) for x in corpus_in_parasent]
    new_para_lengths = sample(all_para_lengths,number_of_paragraphs)
    new_corpus_in_sent = []
    new_corpus_in_str = ''
    new_corpus_in_para = []
    for length in new_para_lengths:
        p = f.add_paragraph()
        new_para = []
        for i in range(length):
            new_sent = choice(corpus_in_sent)
            if repetition == False:
                while new_sent in new_corpus_in_sent:
                    new_sent = choice(corpus_in_sent)
            new_para.append(new_sent)
            new_corpus_in_sent.append(new_sent)
            new_corpus_in_str += new_sent
            p.add_run(new_sent + '  ')
        new_corpus_in_para.append(new_para)
    print("number of sentences in new corpus: " + str(len(new_corpus_in_sent)))
    os.chdir(output_directory)
    filename = title + '.docx'
    f.save(filename)
    print("saved! view the new text (" + title + ") at " + output_directory)

# def shuffle_sentences(paths_to_use, output_path, set=True, number_of_paragraphs=0, scope = 'all', shuffle = 'globally'):
#     if not isinstance(paths_to_use, (list,)):
#         paths_to_use = [paths_to_use]
#     print("Shuffling " + str(len(paths_to_use)) + " document(s) by sentence")
#     list_of_sentences = []
#     total_paragraphs = 0
#     all_paragraphs = []
#     paragraph_lengths = []
#     para_dict = {}
#     for path in paths_to_use:
#         doc = None
#         try:
#             doc = docx.Document(path)
#         except:
#             print("couldn't scrape it")
#             all_runs = []
#         if doc:
#             for para in doc.paragraphs:
#                 txt = para.text.encode('ascii','ignore')
#                 para_in_sentences = tokenize.sent_tokenize(txt)
#                 all_paragraphs.append(para_in_sentences)
#                 for sentence in para_in_sentences:
#                     list_of_sentences.append(sentence)
#     if not number_of_paragraphs:
#         if scope == 'all':
#             number_of_paragraphs = len(all_paragraphs)
#         else:
#             number_of_paragraphs = randint(2,len(all_paragraphs))
#
#     sections_and_sources = {}
#     paragraphs = {}
#     f = docx.Document()
#     f.add_heading(str(number_of_paragraphs) + " Paragraphs shuffled" + shuffle)
#     print((str(number_of_paragraphs) + " Paragraphs"))
#     for i in range(number_of_paragraphs):
#         key = str(i)
#         print(str(i))
#         j = randint(0,)
#         number_of_sentences =
#         current_sentences = []
#         total_extracted = 0
#         remaining_sentences = number_of_sentences - total_extracted
#         while len(current_sentences) < number_of_sentences:
#             sentence = ""
#             while not sentence or not any(c.isalpha() for c in sentence):
#                 # try:
#                 if len(paths_to_use) > 1:
#                     random_path = sample(paths_to_use, 1)[0]
#                 else:
#                     random_path = paths_to_use[0]
#                 # print(random_path)
#                 # print(remaining_sentences)
#                 # try:
#                 sentence, page_number = extract_sentence(random_path)
#                 # except:
#                 # print(os.path.basename(random_path) + ' failed')
#                 continue
#
#             if sentence and len(sentence) > 1:
#                 # print(sentence)
#                 words_in_sentence = tokenize.word_tokenize(sentence)
#                 if len(words_in_sentence) > 25 or len(words_in_sentence) < 2:
#                     sentence = ' '.join(sample(partition_text_in_words(words_in_sentence, randint(7, 20)), 1)[0])
#
#                 end_punct = ['!', '?', '.', ';', '--', ':', '/']
#
#                 week_days = ['Sun ', 'Mon ', 'Tue ', 'Wed ', 'Thu ', 'Fri ']
#                 months = ['Jan ', 'Feb ', 'Mar ', 'Apr ', 'Jun ', 'Jul', 'Aug ', 'Sep ', 'Oct ', 'Nov ', 'Dec ']
#                 try:
#                     while not sentence[-1].isalpha():
#                         sentence = sentence[:len(sentence) - 1]
#                 except:
#                     continue
#                 # while not sentence[-1].isalpha():
#                 #     sentence = sentence[:-1]
#                 # if not sentence[-1] in end_punct:
#                 #     # sentence += sample(end_punct,1)[0]
#                 #     sentence += '.'
#                 outer_punct = [['"', '" '], ['[', '] '], ['{', '} '], ['(', ') '], ['...', '...']]
#                 chosen_punct = sample(outer_punct, 1)[0]
#                 pattern = r'\[.*?\]'
#                 sentence = re.sub(pattern, '', sentence)
#                 for c in [str(i) for i in range(0, 100)] + week_days + months:
#                     sentence = sentence.replace(c, '')
#                 try:
#                     sentence = time.strptime(sentence, '%H:%M:%S')
#                 except:
#                     # print('no time to strip')
#                     this = False
#                 try:
#                     while sentence[0] and not sentence[0].isalpha():
#                         sentence = sentence[1:]
#
#                     sentence = chosen_punct[0] + sentence + chosen_punct[1]
#                     current_sentences.append(sentence.capitalize())
#                     if not random_path in sections_and_sources.keys():
#                         sections_and_sources[(os.path.basename(random_path)[:len(random_path) - 4])] = page_number
#                         # print("We're adding this line")
#                         #     current_sentences.append(sentence)
#                         # store_source(sources, os.path.basename(random_path), page, sections_and_sources,
#                         #                                   key)
#                         # print("we've stored the sentence(s) and source(s)")
#                         # print("total sentence(s) extracted at this stage = " + str(total_extracted))
#                     remaining_sentences -= 1
#                 except:
#                     print('bad sentence')
#                 shuffle(current_sentences)
#                 # if not random_path + ', p.' + page_number in sources:
#
#         if not len(current_sentences) >= number_of_sentences:
#             continue
#         else:
#             print(str(len(current_sentences)) + " out of " + str(number_of_sentences) + " sentences being added")
#             if key not in vignettes:
#                 vignettes[key] = [" ".join(current_sentences)]
#             else:
#                 vignettes[key].append([" ".join(current_sentences)])
#     print('compiling vignette')
#     section_break_symbol = sample(end_punct, 1)[0] * randint(4, 30)
#     im_list = None
#     for key in vignettes:
#         print(key)
#         f.add_heading(section_break_symbol)
#         if im_list:
#             image_added = False
#             while not image_added:
#                 try:
#                     picture_path = sample(im_list, 1)[0]
#                     if isinstance(picture_path, str):
#                         f.add_picture(picture_path, width=Inches(4))
#                         image_added = True
#                 except:
#                     print("failed to insert " + picture_path)
#
#         f.add_paragraph(vignettes[key])
#         p = f.add_paragraph()
#         p.add_run().add_break()
#
#     f.add_page_break()
#     f.add_heading('Sources', 1)
#     # last_p = f.add_paragraph()
#     for x in sections_and_sources.keys():
#         p = f.add_paragraph(x)
#         if not sections_and_sources[x] == 0:
#             p.add_run(str(sections_and_sources[x]))
#
#             # for x in sources.keys():
#             #     f.add_paragraph(str(x) + " " + str(sources[x]))
#     # f.write(sources)
#     assemblage_name = 'prose_poem' + str(len(vignette_numbers)) + '.docx'
#     os.chdir(output_path)
#     if os.path.isfile(assemblage_name):
#         assemblage_name = 'prose_poem' + str(len(vignette_numbers)) + "_" + str(randint(1, 500)) + '.docx'
#
#     f.save(assemblage_name)
#
#     print("Complete! Check out: " + assemblage_name)

def scrape_facebook_html(directory,output_directory):
    html_paths = []
    message = ''
    all_messages = []
    for path in get_filepaths(directory):
        if path.endswith('html'):
            html_paths.append(path)

    for file in html_paths:
        HtmlFile = open(file, 'r', encoding='utf-8')
        source_code = HtmlFile.read()
        soup = BeautifulSoup(source_code)
        # print(soup.prettify())
        if soup.p.string:
            message = soup.p.string
            all_messages.append(message)
    f= docx.Document()
    for message in all_messages:
        f.add_paragraph(message)
        print('added message')
    f.save(output_directory + 'facebook_messages.docx')
    return all_messages

def partition_word_list(text_in_words,chunk_length=randint(2,6)):
    l = text_in_words
    n = chunk_length
    """Yield successive n-sized chunks from l."""
    chunk_list = []
    for i in range(0, len(l), n):
        chunk_list.append(l[i:i + n])
    return chunk_list

def extract_from_docx(filepath, output_type ='runs',output_size = randint(4,40)):
    doc = None
    try:
        doc = docx.Document(filepath)
        all_paragraphs = []
        all_runs = []
        length = len(doc.paragraphs)
        # print(str(length))
        if length >= 1:
            for paragraph in doc.paragraphs:
                all_paragraphs.append(paragraph)
                for run in paragraph.runs:
                    all_runs.append(run.text)
            doc_string = ' '.join(all_runs)
            if output_type == 'paragraphs':
                return all_paragraphs
            elif output_type == 'runs':
                return all_runs
            elif output_type == 'all':
                # print("got all " + str(length) + " run(s)")
                # print("sample: " + all_runs[0])
                return doc_string
            elif output_type == 'chunk':
                words = doc_string.split()
                number_of_words = len(words)
                # print(number_of_words)
                chunk_start = randint(0, number_of_words - output_size - 1)
                chunk_end = chunk_start + output_size
                chunk = ' '.join(words[chunk_start:chunk_end])
                return chunk
    except:
        return ''

def extract_from_document(filepath, output_type ='runs',output_size = randint(4,40)):
    # try:
    if filepath.endswith('.docx'):
        output = extract_from_docx(filepath, output_type,output_size)
        return output

    elif filepath.endswith('.pdf'):
        try:
            extract_from_pdf(filepath,type='page',subtype=output_type)
        except FileNotFoundError:
            print("couldn't scrape pdf")
            return

def get_pdf_paths(file_paths):
    paths_for_pdfs = []
    for path in file_paths:
        word = path
        regexp = re.compile(r'pdf')
        if regexp.search(word):
            paths_for_pdfs.append(path)
    return paths_for_pdfs

def fragment_writer(input_directory,number_of_fragments=10,output_directory=''):
    if not output_directory:
        output_directory = 'C:/Users/Vvayne/Documents/' + os.path.basename(input_directory)
    undesirable = ['[', ']', ':', ':', ')', '(']
    print("initializing directories")
    try:
        if not os.path.exists(output_directory):
            os.makedirs(output_directory)
    except:
        output_directory = input_directory + 'Outputs/' + os.path.basename(input_directory)
        if not os.path.exists(output_directory):
            os.makedirs(output_directory)
    os.chdir(output_directory)
    print("getting filepaths")
    files = get_document_paths(input_directory)
    files = [x if x.endswith('.docx') else '' for x in files]
    chunks_per_fragment = randint(3,10)
    # number_of_words = ''.join([int(x) if isinstance(x,int) else randint(10,40) for x in input("number of words?")])
    f = docx.Document()
    print("gathering " + str(number_of_fragments) + " fragments with " + str(chunks_per_fragment) + " chunks in each")
    fragments = []
    for i in range(number_of_fragments):
        # p = f.add_paragraph()
        chunk_list = []
        while len(chunk_list) < chunks_per_fragment:
            current_file = sample(files, 1)[0]
            # try:
                # print(current_file)
            chunk = extract_from_document(current_file,output_type='chunk',output_size=randint(4,14))
            if chunk:
                while len([c for c in undesirable if c in chunk]) > 0:
                    chunk = extract_from_document(current_file, output_type='chunk', output_size=randint(4, 14))
                # if not re.search(r'\[([^]]+)\]', chunk) or not len([c in undesirable for c in chunk]) > 0:
                tidy_chunk = tokenize.word_tokenize(chunk)
                for word in tokenize.word_tokenize(chunk):
                    try:
                        if not word or not word[0].isalpha() or not word.lower() in english_vocab:
                            tidy_chunk.remove(word)
                        if len(word) == 1 and not word.lower() in ['i','a']:
                            tidy_chunk.remove(word)
                    except:
                        continue

                chunk_list.append(' '.join(tidy_chunk))
            # except:
            #     print("couldn't get a chunk from " + current_file)
            # chunk_list = [str(chunk) if not isinstance(chunk,str) else chunk for chunk in chunk_list]
        if chunk_list and len(chunk_list) >= chunks_per_fragment:
            p = f.add_paragraph()
            while chunk_list:
                if len(chunk_list) > 2:
                    num_of_chunks = randint(1,int((len(chunk_list)-1)/2))
                else:
                    num_of_chunks = len(chunk_list)
                chunks = sample(chunk_list,num_of_chunks)
                p.add_run(' '.join(chunks))
                p.add_run().add_break()
                chunk_list = [chunk for chunk in chunk_list if not chunk in chunks]
            # fragment = ' '.join(chunk_list)
            # fragments.append(fragment)
            print("Fragment #" + str(i) + ' added.')
    title = ''
    while not title:
        try:
            title = extract_from_document(sample(get_filepaths(input_directory), 1)[0], output_type='chunk',
                                          output_size=randint(4, 14))
        except:
            continue
    clean_title = re.sub(r'[^\w\s]', '', title)

    try:
        f.save(clean_title + ' ' + str(number_of_fragments) + ' x ' + str(chunks_per_fragment) + '.docx')
    except:
        f.save(clean_title + str())
    print("Saved fragments")

def grab_words(path, max_words):
    fail_count = 0
    word_list = []
    print('parsing ' + os.path.basename(path) + ' for words')
    page_text = ""
    if path.endswith('.pdf'):
        while not page_text or len(page_text.split(' ')) < 2:
            # print("scraping a page of text (after " + str(fail_count) + " tries)")
            if fail_count > 20:
                print("couldn't get any text")
                return word_list
            try:
                page_text,page_number = extract_from_pdf(path)

            except:
                fail_count += 1
                # print('failed to get a page (after failing ' + str(fail_count) + ' times')
            fail_count +=1


        # page_as_sentences = split_text([page_text])
    elif path.endswith('.docx'):
        fail_count = 0
        page_text = ""
        print("trying to extract all text from this doc")
        page_text = extract_all_text_from_docx(path)
        if page_text:
            # print(page_text)
            # fail_count +=1
            # if fail_count > 20:
            #     print("can't parse this doc")
            #     return word_list
            if not isinstance(page_text[0],str):
                print('converting runs to strings')
                all_lines_as_list = []
                for run in page_text:
                    all_lines_as_list.append(run.text)
                page_text = " ".join(all_lines_as_list)

            if not isinstance(page_text, str):
                # print('this page text is in the wrong form')
                try:
                    page_text = " ".join(page_text)
                except:
                    print('no luck fixing it')
                    return word_list
        else:
            print('no text here')
            return word_list
    number_of_words_available = len(tokenize.word_tokenize(page_text))
    if not number_of_words_available:
        return word_list
    print("found " + str(number_of_words_available) + " words in this document")
    if number_of_words_available < 5:
        word_list = sample(tokenize.word_tokenize(page_text),len(tokenize.word_tokenize(page_text)))
        return word_list

    number_of_words = randint(2, max(2,min(max_words,number_of_words_available)))
    print("let's collect " + str(number_of_words) + " from this doc")
    page_in_words = tokenize.word_tokenize(page_text)
    number_of_available_words = len(page_in_words)
    start = randint(0, number_of_available_words - number_of_words)
    interval = range(start, start + number_of_words)
    for i in interval:

        word_list.append(page_in_words[i])
    # print(word_list)
    print('added ' + str(len(interval)) + ' words to list')
    return word_list

def assemble_phrases(input_paths,output_directory,number_of_words=randint(60,500)):
    all_words_so_far = []
    all_lines_so_far = []
    f = docx.Document() # initialize new document to write and save construction
    line_count = 0 # initialize line counter
    word_count = 0
    number_of_lines = randint(10,2*max(primes(number_of_words))+10) #define number of lines in new construction between 5 and 40 and ideally equal to the length of a random text in the materials folder
    print("let's get " + str(number_of_words) + " words into " + str(number_of_lines) + ' lines')
    sources = [] #initialize list to store the origin of each line
    while len(all_words_so_far) < number_of_words:
        remaining_words = number_of_words - len(all_words_so_far)
        print('still need ' + str(remaining_words) + ' words')
        if remaining_words <= 0:
            break
        current_path = sample(input_paths,1)[0]
        max_number_of_new_words = 0
        while max_number_of_new_words > .1*number_of_words or max_number_of_new_words == 0:
            max_number_of_new_words = min([randint(2,round(.1*number_of_words)),remaining_words])
        new_words_list = grab_words(current_path, max_number_of_new_words)
        if new_words_list:
            for word in new_words_list:
                if word and word[0].isalpha() and word in english_vocab:
                    all_words_so_far.append(word)
            print('added new word list to big word list')
            sources.append(os.path.basename(current_path))
            print('stored the source of the words too')
    print('[COLLECTED ' + str(len(all_words_so_far)) + ' WORDS]')
    print("let's make lines out of them")
    for line_number in range(1,number_of_lines):
        new_line_string = ""
        # refrain_boolean, refrain = refrain_or_no(
        #     all_lines_so_far)  # random integer to decide if a line should be repeated and returns it if yes
        # # print(refrain_boolean)
        # if refrain_boolean == 'yes':  # test for refrain_affirmation
        #     new_line = refrain  # stores old line as new line for refrain
        #     print('refrained')
        line_length = max(14, sample(primes(number_of_lines), 1)[0])
        words_in_chunks = partition_text_in_words(all_words_so_far,randint(3,7))
        while not new_line_string or len(new_line) < line_length:
            # start = randint(0,len(all_words_so_far) - line_length - 1)
            # end = start + line_length
            # new_line = " ".join(all_words_so_far[start:end])
            new_line = []
            while len(new_line) < line_length:
                chunk = sample(words_in_chunks,1)[0]
                while not isinstance(chunk,str):
                    chunk = " ".join(chunk)
                new_line.append(chunk + " ")
            new_line_string = " ".join(new_line)
        all_lines_so_far.append(new_line)
    shuffle(all_lines_so_far)
    for line in all_lines_so_far:
        f.add_paragraph(line)
        # print(line)

    f.add_page_break() #starts a new page in document being constructed
    f.add_heading('Sources', 1) #write header called Sources on new document
    last_p = f.add_paragraph() #start final block
    set_of_sources = [] #initialize unique set of sources
    for source in sources: #loop through list of all sources for lines in the new document
        if not source in set_of_sources: #check for uniqueness
            set_of_sources.append(source) #add if unique
    for source in set_of_sources: #loop through unique sources
        last_p.add_run(source) #write source
        last_p.add_run().add_break() #give space before the next one
    name = 'poem_by_phrases_' + str(len(set_of_sources)) + '_' + str(number_of_words) + '.docx' #generate a name using the number of different poems used and a random integer
    os.chdir(output_directory)
    f.save(name) # save new document
    print("Wrote a thing called: " + name)
    print("It's got " + str(len(all_lines_so_far)) + " lines from " + str(len(set_of_sources)) + " different sources")
    print("You can find it in " + output_directory)
    return all_lines_so_far

def generate_phrase(input_directory,n = randint(4,9)):
    phrase = ''
    while not phrase:
        source_path = choice(get_document_paths(input_directory))
        while not source_path or not source_path.endswith('.docx') or not source_path[0].isalnum():
            source_path = choice(get_document_paths(get_filepaths("C:/Users/Vvayne/Documents/FA 2018/Multigenre Workshop/")))
            # print(source_path)
        print('looking for words to make a title in ' + os.path.split(source_path)[-1])
        text_in_words = ''
        number_of_words = 0
        # try:
        all_text = extract_all_text_from_docx(source_path)
        if all_text:
            # print(type(all_text))
            # print(' '.join(all_text)[0:20])
            text_in_words = tokenize.word_tokenize(' '.join(all_text))
            # print(text_in_words[0:3])
            number_of_words = len(text_in_words)
            print(str(number_of_words))
        # except:
        #     print("couldn't get a phrase from " + source_path)
        #     continue

        if text_in_words and number_of_words > n:
            phrase = ' '.join(sample(text_in_words,n))

    clean_phrase = re.sub(r'[^\w\s]', '', phrase)
    clean_phrase = ' '.join(clean_phrase.split())
    print('phrase generated: ' + clean_phrase)
    return clean_phrase

def remove_repeated_words(string):
    clean_string = []
    string_in_words = string.split()
    for i in range(string_in_words):
        if not string_in_words[i] == string_in_words[i + 1]:
            clean_string.append(string_in_words[i])
    return clean_string

def strip_numbering_tags(string):
    # print("let's test for numbering/sectional tags so they don't get lumped in as lines")
    s = string
    if not isinstance(string, str):
        s = string.text
    if re.search('[iv]+\.', s):
        lst = re.findall('[iv]+\.', s)
        for tag in lst:
            s.replace(tag, "")
    elif re.search('\d+\.', s):
        lst = re.findall('\d+\.', s)
        for tag in lst:
            s.replace(tag, "")
            # s.replace([i for i in lst],"")
    return s

def generate_paragraphs(input_directory, number_of_paragraphs=randint(5, 10)):
    # get filepaths in input directory
    doc_paths = get_document_paths(input_directory)
    print('number of strophes: ' + str(number_of_paragraphs) + ' |  number of sourcetexts: ' + str(len(doc_paths)))
    paragraphs = []
    lines_added = []
    for i in range(number_of_paragraphs):
        print('populating strophe ' + str(i))
        sent_per_para = randint(4, 9)
        paragraph_sentences = []
        while len(paragraph_sentences) < sent_per_para:
            sentence_to_add = extract_sentence(choice(doc_paths))
            if sentence_to_add and sentence_to_add[:-1] not in lines_added:
                sent_no_numbers = []
                for string in sentence_to_add.split():
                    new_string = strip_numbering_tags(string)
                    sent_no_numbers.append(new_string)
                lines_added.append(' '.join(sent_no_numbers))
                ulist = []
                [ulist.append(x) for x in sentence_to_add.split() if x not in ulist]
                paragraph_sentences.append(' '.join(ulist))
                print('sentence ' + str(len(paragraph_sentences)) + ' of ' + str(sent_per_para) + ' added')
            else:
                print('either wrote this one already or fake sentence')

        paragraphs.append(paragraph_sentences)
        print('strophe ' + str(i) + ' added')
    print('finished strophe generation')
    return paragraphs

def write_list_of_text_units(text_list, output_directory, titles=False):
    print("now let's write these strophes to a document")
    f = docx.Document()
    print("first let's generate a title")
    title = generate_phrase()
    print("the title will be " + title)
    filename = title + '.docx'
    f.add_heading(title, 0)
    for i in range(len(text_list)):
        print('writing strophe ' + str(i) + ' of ' + str(len(text_list)))
        if titles:
            subtitle = text_list[0][0]
            f.add_heading(subtitle, 2)
            del text_list[0][0]
        p = f.add_paragraph()
        p.add_run(' '.join(text_list[i]))
        p.add_run().add_break()
        print('strophe ' + str(i) + ' added')
    print('all strophes written to the document')
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    os.chdir(output_directory)
    if os.path.isfile(filename):
        filename = title + "_" + str(randint(1, 500)) + '.docx'
    print('saving document')
    f.save(filename)
    print('all done! you can find ' + filename + ' @ ' + output_directory)

def generate_line_from_seed(corpus, cfdist, word, num_of_words=randint(9,16)):
    line_list = []
    fail_count = 0
    while len(line_list) < num_of_words:
        while word in line_list or word not in english_vocab:
            word = choice(list(cfdist[word].most_common()))[0]
            fail_count += 1
            if fail_count >= len(list(cfdist[word])):
                # print(pos_tag([word]))
                # print(word)
                token = tokenize.word_tokenize(word)
                # print(token)
                tagged_token = pos_tag(token)
                # print(tagged_token)
                tag = ' '.join([x[1] for x in tagged_token])
                # print(tag)
                # print(tag)

                word = choice(list_of_pos(corpus, tag))
        line_list.append(word)
        # if len(list(cfdist[word])) > 2:
        #     word = choice(cfdist[word].most_common(3))
        # else:

    line = ' '.join(line_list)
    return line

def generate_bigrams_cfd(text):
    bigrams = nltk.bigrams(text)
    cfd = nltk.ConditionalFreqDist(bigrams)
    return bigrams, cfd

def generate_trigrams_cfd(text):
    text_in_words = tokenize.word_tokenize(text)
    trigrams = []
    for i in range(text_in_words):
        trigrams.append((text_in_words[i - 1], text_in_words[i], text_in_words[i + 1]))
    cfd = nltk.ConditionalFreqDist(trigrams)
    return trigrams, cfd

def create_corpus(input_directory):
    f = docx.Document()
    list_of_fulldoc_strings = []
    doc_paths = get_document_paths(input_directory)
    intermediary_paths = [path for path in doc_paths if path.endswith('.docx')]
    paths = []
    for address in intermediary_paths:
        if not os.path.split(address)[-1] in ' '.join(paths):
            paths.append(address)
    for path in paths:
        try:
            all_text = extract_all_text_from_docx(path)
            if all_text:
                list_of_fulldoc_strings.append(' '.join(all_text))
        except:
            continue
    corpus_string = ' '.join(list_of_fulldoc_strings)
    corpus_list = re.findall(r"[\w']+|[.,!?;]", corpus_string)
    return corpus_list

def list_of_pos(corpus, tag):
    word_list = []
    words_and_tags = nltk.pos_tag(corpus)
    for a, b in words_and_tags:
        if tag in b:
            word_list.append(a)
    return word_list

def generate_lines_from_corpus(input_directory,number_of_lines = 12,seed_type='most common'):
    print("let's create a corpus from which to generate")
    current_corpus = create_corpus(input_directory)
    current_corpus_words = [x for x in tokenize.word_tokenize(' '.join(current_corpus)) if x in english_vocab]
    # stopwords = nltk.corpus.stopwords.words('english')
    # current_corpus_words = [word for word in current_corpus_words if word not in stopwords]
    if current_corpus_words:
        print(len(current_corpus_words))
        print("now let's produce a list of bigrams and conditional frequency distribution of them")
        bigrams, cfd = generate_bigrams_cfd(current_corpus_words)
        print("now let's initialize and find a seed word")
        used_seeds = []
        lines_generated = []
        for i in range(number_of_lines):
            seed_word = ''
            fail_counter = 0
            n = 10
            while not seed_word or seed_word in used_seeds:
                if seed_type == 'most common':
                    fdist1 = FreqDist(current_corpus_words)
                    top_words = [x for x, y in fdist1.most_common(n) if len(x) > 3]
                    seed_word = choice(top_words)
                    fail_counter += 1
                    if fail_counter > number_of_lines:
                        n += 5
                if seed_type == 'noun':
                    seed_word = choice(list_of_pos(current_corpus_words,'NN'))

            if seed_word:
                print("this seed word will be " + seed_word)
                used_seeds.append(seed_word)
                line_length = randint(3,20)
                print("now let's generate a line with it")
                new_line = generate_line_from_seed(current_corpus_words, cfd,seed_word, line_length)
                if new_line and not new_line in lines_generated:
                    print("it's a new line")
                    lines_generated.append(new_line)
        print("all lines generated")
        return lines_generated

def write_lines_in_strophes(lines,output_directory,number_of_strophes = 0,lines_per_strophe=0):
    f = docx.Document()
    title = choice(lines)
    filename = title + '.docx'
    f.add_heading(title,2)
    lines.remove(title)
    if number_of_strophes == 0:
        number_of_strophes = randint(1, len(lines))
    if lines_per_strophe == 0:
        lines_per_strophe = int(len(lines)/number_of_strophes)
    for i in range(number_of_strophes):
        p = f.add_paragraph()
        new_lines = sample(lines,lines_per_strophe)
        for line in new_lines:
            p.add_run(line)
            p.add_run().add_break()

        lines = [x for x in lines if not x in new_lines]
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    os.chdir(output_directory)
    if os.path.isfile(filename):
        filename = title + "_" + str(randint(1, 500)) + '.docx'
    print('saving document')
    f.save(filename)
    print('all done! you can find ' + filename + ' @ ' + output_directory)

def get_first_line(doc_paths,source=True):
    line = ''
    this_file = ''
    line_no = None
    while not line:
        this_path = sample(doc_paths,1)[0]
        this_file = os.path.basename(this_path)
        try:
            print("let's try to get an endword from " + this_file)
            line, line_no = extract_line_from_document(this_path)
        except:
            continue
    if source == True:
        source = (this_file[:-3],line_no)
    return line,source

def get_end_word(line):
    line_in_words = tokenize.word_tokenize(line)
    if len(line_in_words) > 1:
        end_word = line_in_words[-1]
        return end_word
    else:
        return ''

def line_from_word(word):
    line, source = "",()


def extract_line_from_document(file_path):
    if os.path.exists(file_path):
        filename = os.path.basename(file_path)
        if filename:
            line = ""
            section_num = 0

            if filename.endswith('.pdf'):
                line, section_num = extract_from_pdf(file_path, type='line')
                # print(line)
                return line, section_num

            elif filename.endswith('.doc') or filename.endswith('.docx'):
                # print("scraping " + filename + " at " + file_path + " for a line")
                # try:
                all_text = extract_all_text_from_docx(file_path)

                print('got all text from doc')
                if len(all_text) >= 1:
                    line_number = randint(0, len(all_text) - 1)
                    line = all_text[line_number]
                    # print('scraped a line')
                    return line, line_number

                else:
                    print('no text here')
                    line = ''
                    line_number = 0
                    return line, line_number

def end_word_poem_maker(input_directory, output_directory, number_to_make=1):
    f = docx.Document()
    all_lines = []
    sources = []
    all_end_words = []
    number_of_lines = randint(15, 20)

    print('initializing end word writer')
    all_paths = get_document_paths(input_directory)
    first_end_word = ''
    first_source = ''
    print('getting first end word')
    while not first_end_word or first_end_word not in english_vocab:
        first_line = ''
        while not first_line:
            # try:
            first_line, first_source = get_first_line(all_paths)
            # except:
            #     continue
        all_lines.append(first_line)
        sources.append(first_source)
        first_end_word = get_end_word(first_line)
    word = first_end_word
    all_end_words.append(first_end_word)
    print('first end word: ' + first_end_word)
    line, new_end_word = '', ''
    while len(all_lines) < number_of_lines:
        print('old word = ' + word)
        if new_end_word and not word == new_end_word:
            word = new_end_word
            new_end_word = ''
        print('current word to search for = ' + word)

        fail_count = 0
        # print('looking for a line with ' + word)
        line = ''
        while not line:
            if fail_count >= 50:
                nw_fail_count = 0
                while word in new_end_word or new_end_word not in english_vocab or len(new_end_word) < 3:
                    # print("let's try another word")
                    new_end_word = sample(tokenize.word_tokenize(' '.join(all_lines)), 1)[0]
                    nw_fail_count += 1
                    if nw_fail_count >= 24:
                        print('getting new word and line')
                        line = get_first_line(all_paths)
                        new_end_word = get_end_word(get_first_line(all_paths))
                # print('going with ' + new_end_word)
                fail_count = 0
                word = new_end_word
            c_path = sample(all_paths, 1)[0]
            # print('looking in ' + c_path)
            try:
                # print('extracting')
                text_list, page_no = extract_from_document(c_path, 'page')
            except:
                # print('nope')
                fail_count += 1
                continue
            text_in_words = [word for word in tokenize.word_tokenize(' '.join(text_list)) if word in english_vocab]

            if len(text_in_words) > 3:

                text_list = partition_text_in_words(tokenize.word_tokenize(' '.join(text_in_words)),
                                                    chunk_length=randint(3, min([8, len(text_in_words)])))
                for item in text_list:
                    fragment = ' '.join(item)
                    if re.search(word, fragment, re.I):
                        newfragmentwords = tokenize.word_tokenize(fragment)
                        newfragmentwords2 = [word if word in english_vocab else '' for word in newfragmentwords]
                        newfragment = ' '.join(newfragmentwords2)
                        source = ((os.path.basename(c_path)[:-5]), page_no)
                        line = newfragment
            fail_count += 1
            # print(fail_count)
        print(str(len(tokenize.word_tokenize(line))))
        if line and source and line not in all_lines and len(tokenize.word_tokenize(line)) > 1:
            print('adding line')
            if len(tokenize.word_tokenize(line)) > 24:
                line = sample(line, randint(6, 10))
            all_lines.append(line)
            print('grabbing new word to look for')
            new_end_word = get_end_word(line)
            nw_fail_count = 0
            while new_end_word == word or new_end_word not in english_vocab or len(new_end_word) < 3:
                print("let's try another word")
                new_end_word = sample(tokenize.word_tokenize(line), 1)[0]
                nw_fail_count += 1
                # print('new word will be ' + new_end_word)
                if nw_fail_count >= 24:
                    print('keeping old word')
                    new_end_word = word
                    break

            print('final end word to carry on: ' + new_end_word)
            sources.append(source)
        line = ""

    stanza_break = 0
    while stanza_break < 2:
        stanza_break = sample(primes(number_of_lines), 1)[0]
    line_counter = 1
    p = f.add_paragraph()
    for line in all_lines:
        if line_counter % stanza_break == 0 and line_counter > 0:
            p = f.add_paragraph()
            p.add_run().add_break()
        p.add_run(line)
        p.add_run().add_break()
        line_counter += 1

    f.add_page_break()  # start new page in document
    f.add_heading('Sources', 1)  # write 'Sources' at the top of the page
    last_p = f.add_paragraph()  # start final block
    set_of_sources = []  # initialize list of different sources (without duplicates)
    for source in sources:  # loop through each source that's been collected
        if not source in set_of_sources:  # test if current source is already in the list of unique sources
            set_of_sources.append(source)  # store source in in list of different sources if it isn't already in there
    for source in set_of_sources:  # loop through unique sources
        last_p.add_run(str(source))  # write source to page
        last_p.add_run().add_break()  # add space before next source
    name = first_end_word + str(len(set_of_sources)) + '_' + str(
        number_of_lines) + '.docx'  # define name of new document based on how many different texts contributed and a random integer
    print(name)
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
    os.chdir(output_directory)
    f.save(name)  # save the new document
    print("Wrote a poem called " + name)
    print(os.path.join(output_directory, name))
    print("It's got " + str(len(all_lines)) + " lines from " + str(len(set_of_sources)) + " different texts")
    print("You can find it in " + output_directory)

def construct_one(category):
    type = 'sentences'
    length_in_words = 800
    tense = 'present'
    common_words = ['but','and','like']
    sentence_endings = ['.','?','!']
    current_length = 0
    while current_length < length_in_words:
        line_length = randint(5,7)

def generate_ghost_tantra(text_path):
    if text_path.endswith('.docx'):
        text_as_string = ' '.join(extract_all_text_from_docx(text_path))
        text_as_words = text_as_string.split()
    num_of_lines = randint(11,18)
    for i in range(1,num_of_lines):
        line_length = randint(1,12)

def generate_ottavas(input_directory,output_directory,number_to_generate=3, repetition = False):
    doc_list = get_document_paths(input_directory)
    source_dict = {}
    corpus_in_docpara = []
    corpus_in_docstring = []
    corpus_in_para = []
    corpus_in_lines = []
    corpus_in_paralines = []
    corpus_in_string = ''
    for doc_path in doc_list:
        print('extracting text from ' + os.path.basename(doc_path))
        doc = docx.Document(doc_path)
        doc_in_para = []
        doc_in_string = ''
        for paragraph in doc.paragraphs:
            doc_in_para.append(paragraph.text)
            corpus_in_para.append(paragraph.text)
            para_in_lines = []
            for run in paragraph.runs:
                if len(tokenize.word_tokenize(run.text))>1:
                    para_in_lines.append(run.text)
                    corpus_in_lines.append(run.text)
                    doc_in_string += run.text
            corpus_in_paralines.append(para_in_lines)
        corpus_in_docpara.append(doc_in_para)
        corpus_in_string += doc_in_string
        corpus_in_docstring.append('\n'.join(doc_in_para))
        source_dict[os.path.basename(doc_path)] = doc_in_string

    title = ' '.join(sample(tokenize.word_tokenize(corpus_in_string), randint(2, 6)))
    f = docx.Document()
    print("number of ottavas in source corpus: " + str(len(corpus_in_para)))


    new_corpus_in_lines = []
    new_corpus_in_str = ''
    new_corpus_in_para = []
    for i in range(number_to_generate):
        f.add_heading(str(i))
        p = f.add_paragraph()
        new_para = []
        line_sources = []
        while len(new_para) < 8:
            new_line = choice(corpus_in_lines)
            if not repetition:
                while new_line in new_corpus_in_lines:
                    new_line = choice(corpus_in_lines)
            new_line = ''.join([c if not c.isdigit() else '' for c in new_line])
            new_para.append(new_line)
            new_corpus_in_lines.append(new_line)
            new_corpus_in_str += new_line
            new_line_source = ''
            for source in source_dict.keys():
                if new_line in source_dict[source]:
                    new_line_source = source
                    line_sources.append(new_line_source)
            p.add_run(new_line + ' ')
            p.add_run().add_break()
        new_corpus_in_para.append(new_para)
    f.add_page_break()
    f.add_heading('Sources', 1)
    last_p = f.add_paragraph()
    for j in range(len(line_sources)):
        for i in range(len(line_sources)):
            f.add_heading('Stanza ' + str(j))
            last_p.add_run(str(i+1) + '.' + line_sources[j])
            last_p = f.add_paragraph()
    print("number of lines in new corpus: " + str(len(new_corpus_in_lines)))
    os.chdir(output_directory)
    filename = title + '.docx'
    f.save(filename)
    print("saved! view the new text (" + title + ") at " + output_directory)

def generate_n_length_stanzas(input_directory,output_directory, number_to_generate = 3, repetition=True, lines_per_stanza = randint(3,9)):
    doc_list = get_document_paths(input_directory)
    source_dict = {}
    corpus_in_docpara = []
    corpus_in_docstring = []
    corpus_in_para = []
    corpus_in_lines = []
    corpus_in_paralines = []
    corpus_in_string = ''
    corpus_in_docchunks = []
    for doc_path in doc_list:
        print('extracting text from ' + os.path.basename(doc_path))
        doc = docx.Document(doc_path)
        doc_in_para = []
        doc_in_string = ''
        for paragraph in doc.paragraphs:
            doc_in_para.append(paragraph.text)
            corpus_in_para.append(paragraph.text)
            para_in_lines = []
            for run in paragraph.runs:
                if len(tokenize.word_tokenize(run.text)) > 1:
                    clean_string = ''.join([c for c in run.text])
                    para_in_lines.append(clean_string)
                    corpus_in_lines.append(clean_string)
                    doc_in_string += clean_string
            corpus_in_paralines.append(para_in_lines)
        doc_in_chunks = partition_text_in_words(doc_in_string.split(),randint(6,9))
        corpus_in_docchunks.append(doc_in_chunks)
        corpus_in_docpara.append(doc_in_para)
        corpus_in_string += doc_in_string
        corpus_in_docstring.append('\n'.join(doc_in_para))
        source_dict[os.path.basename(doc_path)] = doc_in_string


    title = ' '.join(sample(tokenize.word_tokenize(corpus_in_string), randint(2, 6)))
    title = ''.join([c if c.isalpha() else '' for c in title])
    f = docx.Document()
    print("number of stanzas in source corpus: " + str(len(corpus_in_para)))
    corpus_in_chunks = partition_text_in_words(tokenize.word_tokenize(corpus_in_string),randint(6,9))
    new_corpus_in_lines = []
    new_corpus_in_str = ''
    new_corpus_in_para = []
    for i in range(number_to_generate):
        f.add_heading(str(i))
        p = f.add_paragraph()
        new_para = []
        line_sources = []
        while len(new_para) < lines_per_stanza:
            new_line = ' '.join(choice(corpus_in_chunks))
            while "canto" in new_line.lower():
                new_line = ' '.join(choice(corpus_in_chunks))
            if not repetition:
                while new_line in new_corpus_in_lines:
                    new_line = choice(corpus_in_lines)
            new_line = ''.join([c if not c.isdigit() else '' for c in new_line]).capitalize()
            new_para.append(new_line)
            new_corpus_in_lines.append(new_line)
            new_corpus_in_str += new_line
            new_line_source = ''
            for source in source_dict.keys():
                if new_line in source_dict[source]:
                    new_line_source = source
                    line_sources.append(new_line_source)
            p.add_run(new_line + ' ')
            p.add_run().add_break()
        new_corpus_in_para.append(new_para)
    f.add_page_break()
    f.add_heading('Sources', 1)
    last_p = f.add_paragraph()
    for j in range(len(line_sources)):
        for i in range(len(line_sources)):
            f.add_heading('Stanza ' + str(j))
            last_p.add_run(str(i + 1) + '.' + line_sources[j])
            last_p = f.add_paragraph()
    print("number of lines in new corpus: " + str(len(new_corpus_in_lines)))
    os.chdir(output_directory)
    filename = title + '.docx'
    f.save(filename)
    print("saved! view the new text (" + title + ") at " + output_directory)


def generate_n_length_equalrep_stanzas(input_directory, output_directory, number_to_generate=3, repetition=True,
                              lines_per_stanza=randint(3, 9)):
    doc_list = get_document_paths(input_directory)
    source_dict = {}
    corpus_in_docpara = []
    corpus_in_docstring = []
    corpus_in_para = []
    corpus_in_lines = []
    corpus_in_paralines = []
    corpus_in_string = ''
    corpus_in_docchunks = []
    for doc_path in doc_list:
        print('extracting text from ' + os.path.basename(doc_path))
        doc = docx.Document(doc_path)
        doc_in_para = []
        doc_in_string = ''
        for paragraph in doc.paragraphs:
            doc_in_para.append(paragraph.text)
            corpus_in_para.append(paragraph.text)
            para_in_lines = []
            for run in paragraph.runs:
                if len(tokenize.word_tokenize(run.text)) > 1:
                    clean_string = ''.join([c for c in run.text])
                    para_in_lines.append(clean_string)
                    corpus_in_lines.append(clean_string)
                    doc_in_string += clean_string
            corpus_in_paralines.append(para_in_lines)
        doc_in_chunks = partition_text_in_words(doc_in_string.split(), randint(6, 9))
        corpus_in_docchunks.append(doc_in_chunks)
        corpus_in_docpara.append(doc_in_para)
        corpus_in_string += doc_in_string
        corpus_in_docstring.append('\n'.join(doc_in_para))
        source_dict[os.path.basename(doc_path)] = doc_in_string

    title = ' '.join(sample(tokenize.word_tokenize(corpus_in_string), randint(2, 6)))
    title = ''.join([c if c.isalpha() else '' for c in title])
    f = docx.Document()
    print("number of stanzas in source corpus: " + str(len(corpus_in_para)))
    # corpus_in_chunks = partition_text_in_words(tokenize.word_tokenize(corpus_in_string), randint(6, 9))
    new_corpus_in_lines = []
    new_corpus_in_str = ''
    new_corpus_in_para = []
    for i in range(number_to_generate):
        f.add_heading(str(i))
        p = f.add_paragraph()
        new_para = []
        line_sources = []
        while len(new_para) < lines_per_stanza:
            source_index = randint(0,len(corpus_in_docchunks)-1)
            # print(len(corpus_in_docchunks))
            new_line = ''
            while not new_line:
                try:
                    source_index = randint(0,len(corpus_in_docchunks)-1)
                    new_line = ' '.join(choice(corpus_in_docchunks[source_index]))
                    if 'canto' in new_line:
                        new_line = ''
                except:
                    print("couldn't get chunk from source whose index is: " + str(source_index))
                    new_line = ''
            print(new_line, source_index)
            # if not repetition:
            #     while new_line in new_corpus_in_lines:
            #         new_line = choice(corpus_in_lines)
            new_line = ''.join([c if not c.isdigit() else '' for c in new_line]).capitalize()
            new_para.append(new_line)
            new_corpus_in_lines.append(new_line)
            new_corpus_in_str += new_line
            new_line_source = ''
            for source in source_dict.keys():
                if new_line in source_dict[source]:
                    new_line_source = source
                    line_sources.append(new_line_source)
            p.add_run(new_line + ' ')
            p.add_run().add_break()
        new_corpus_in_para.append(new_para)
    f.add_page_break()
    f.add_heading('Sources', 1)
    last_p = f.add_paragraph()
    for j in range(len(line_sources)):
        for i in range(len(line_sources)):
            f.add_heading('Stanza ' + str(j))
            last_p.add_run(str(i + 1) + '.' + line_sources[j])
            last_p = f.add_paragraph()
    print("number of lines in new corpus: " + str(len(new_corpus_in_lines)))
    os.chdir(output_directory)
    filename = title + '.docx'
    f.save(filename)
    print("saved! view the new text (" + title + ") at " + output_directory)

# input_directory = 'C:/Users/Vvayne/Documents/FA 2018/Multigenre Workshop/workshop inputs/'
input_directory = 'C:/Users/Vvayne/Documents/Anne Waldman 0419/inputs/'
output_directory = "C:/Users/Vvayne/Documents/Anne Waldman 0419/outputs/"
# input_directory = 'C:/Users/Vvayne/Documents/inputs/'
# output_directory = 'C:/Users/Vvayne/Documents/he and she/'
# output_directory = 'C:/Users/Vvayne/Documents/grass outputs/'
if not os.path.exists(output_directory):
    os.makedirs(output_directory)

number_to_make = 4
path_to_use = get_document_paths(input_directory)[0]
for i in range(number_to_make):
    # fragment_writer(input_directory,output_directory=output_directory)
    # assemble_phrases(input_paths=get_document_paths(input_directory),output_directory=output_directory)
    # write_vignettes(paths_to_use=get_document_paths(input_directory),output_path=output_directory)
    # # write_paragraph_shuffle(path_to_use,output_directory)
    # # gen_lines = generate_lines_from_corpus(input_directory, number_of_lines=randint(11, 18), seed_type='most common')
    # # if gen_lines:
    # #     write_lines_in_strophes(gen_lines, output_directory)
    #
    # # end_word_poem_maker(input_directory,output_directory)
    # write_shuffle(input_directory,output_directory)
    generate_ottavas(input_directory,output_directory)
    generate_n_length_equalrep_stanzas(input_directory,output_directory, lines_per_stanza=8)
    generate_n_length_stanzas(input_directory,output_directory,lines_per_stanza=8)


