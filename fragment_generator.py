## fragment_generator.py will generate a fragment of word chunks from pdfs and docx in the input directory
import os
import docx
from random import randint,sample, shuffle
import PyPDF2
import nltk
import re
from nltk import tokenize
english_vocab = set(w.lower() for w in nltk.corpus.words.words())

def partition_word_list(text_in_words,chunk_length=randint(2,6)):
    l = text_in_words
    n = chunk_length
    """Yield successive n-sized chunks from l."""
    chunk_list = []
    for i in range(0, len(l), n):
        chunk_list.append(l[i:i + n])
    return chunk_list

def extract_from_pdf(filepath,type='page',output_type='chunk',output_size = randint(4,40)):
    # filename = os.path.basename(filepath)
    with open(filepath, 'rb') as pdf_file_obj:  # open current pdf
        try:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file_obj)  # store pdf as PyPDF2 reader object
        except:
            print('unable to parse pdf')
            return
        if pdf_reader and not pdf_reader.isEncrypted:
            # print("I'm looking for a line in " + filename)
            try:
                number_of_pages = pdf_reader.getNumPages()  # count number of pages in pdf
            except:
                print("unparseable pdf")
                return

            if number_of_pages > 1:
                if type == 'all' :
                    text_in_lines = []
                    text_in_pages = []
                    for i in range(0,number_of_pages - 1):
                            page_obj = pdf_reader.getPage(i)
                            page_as_list = page_obj.extractText().splitlines()
                            text_in_pages.append(page_as_list)
                            for line in page_as_list:
                                text_in_lines.append(line)
                    if output_type == 'runs':
                        return text_in_lines
                    elif output_type == 'pages':
                        return text_in_pages
                    elif output_type == 'all':
                        return ' '.join(text_in_lines)
                    elif output_type == 'chunk':
                        words = nltk.tokenize.word_tokenize(' '.join(text_in_lines))
                        number_of_words = len(words)
                        # print(number_of_words)
                        chunk_start = randint(0, number_of_words - output_size - 1)
                        chunk_end = chunk_start + output_size
                        chunk = ' '.join(words[chunk_start:chunk_end])
                        if chunk:
                            print('chunk from pdf')
                        return chunk
                elif type == 'page':
                    for j in sample(range(number_of_pages),min(round(number_of_pages/2),number_of_pages)):
                        page_number = j
                        page_obj = pdf_reader.getPage(page_number)
                        if page_obj:
                            try:
                                page_as_string = page_obj.extractText()
                            except:
                                print('text extraction fail')
                                return
                            words = nltk.tokenize.word_tokenize(page_as_string)
                            if words and len(words) > 4:
                                # print(page_as_string)
                                if output_type == 'runs':
                                    page_in_lines = page_obj.extractText().splitlines()
                                    # text_in_pages = [text_in_lines]
                                    return page_in_lines
                                elif output_type == 'all':
                                    return page_as_string
                                elif output_type == 'chunk':

                                    number_of_words = len(words)
                                    if number_of_words >= 4:
                                        if number_of_words > output_size + 1:
                                            chunk_start = randint(0, number_of_words - output_size - 1)
                                            chunk_end = chunk_start + output_size
                                            chunk = ''.join(words[chunk_start:chunk_end])
                                        else:
                                            chunk = words
                                        return chunk
                            # except:
                            #     print('bad page')



                # elif type == 'sentence':
                #     fail_count = 0
                #     sentence = ""
                #     page_num = 0
                #     while not sentence and fail_count < number_of_pages:
                #         page_num = randint(0, number_of_pages-1)
                #         page_obj = pdf_reader.getPage(page_num)
                #         try:
                #             page_as_string = page_obj.extractText()
                #             page_in_sentences = nltk.tokenize.sent_tokenize(page_as_string)
                #             sentence = sample(page_in_sentences,1)[0]
                #         except:
                #             print('failed to extract a sentence from pdf page ' + str(page_num))
                #             fail_count +=1
                #             continue
                #     return sentence, page_num

def extract_from_docx(filepath, output_type ='runs',output_size = randint(4,40)):
    doc = None
    try:
        doc = docx.Document(filepath)
        all_paragraphs = []
        all_runs = []
        length = len(doc.paragraphs)
        # print(str(length))
        if length >= 1:
            for paragraph in doc.paragraphs:
                all_paragraphs.append(paragraph)
                for run in paragraph.runs:
                    all_runs.append(run.text)
            doc_string = ' '.join(all_runs)
            if output_type == 'paragraphs':
                return all_paragraphs
            elif output_type == 'runs':
                return all_runs
            elif output_type == 'all':
                # print("got all " + str(length) + " run(s)")
                # print("sample: " + all_runs[0])
                return doc_string
            elif output_type == 'chunk':
                words = doc_string.split()
                number_of_words = len(words)
                # print(number_of_words)
                chunk_start = randint(0, number_of_words - output_size - 1)
                chunk_end = chunk_start + output_size
                chunk = ' '.join(words[chunk_start:chunk_end])
                return chunk
    except:
        return ''

def extract_from_document(filepath, output_type ='runs',output_size = randint(4,40)):
    # try:
    if filepath.endswith('.docx'):
        output = extract_from_docx(filepath, output_type,output_size)
        return output

    elif filepath.endswith('.pdf'):
        try:
            extract_from_pdf(filepath,type='page',output_type=output_type,output_size = output_size)
        except FileNotFoundError:
            print("couldn't scrape pdf")
            return

def get_filepaths(directory, type = ''):

    """
    This function will generate the file names in a directory
    tree by walking the tree either top-down or bottom-up. For each
    directory in the tree rooted at directory top (including top itself),
    it yields a 3-tuple (dirpath, dirnames, filenames).
    """
    file_paths = []  # List which will store all of the full filepaths.

    # Walk the tree.
    if isinstance(directory,list):
        for folder in directory:
            for root, directories, files in os.walk(folder):
                for filename in files:
                    # Join the two strings in order to form the full filepath.
                    filepath = os.path.join(root, filename)
                    if type == 'pdf':
                        if filepath.endswith('.pdf'):
                            file_paths.append(filepath)
                    elif type == 'document':
                        if filepath.endswith('.pdf') or filepath.endswith('.docx'):
                            file_paths.append(filepath)  # Add it to the list.
                    else:
                        file_paths.append(filepath)  # Add it to the list.
    else:
        for root, directories, files in os.walk(directory):
            for filename in files:
                # Join the two strings in order to form the full filepath.
                filepath = os.path.join(root, filename)
                if type == 'pdf':
                    if filepath.endswith('.pdf'):
                        file_paths.append(filepath)
                elif type == 'document':
                    if filepath.endswith('.pdf') or filepath.endswith('.docx'):
                        file_paths.append(filepath)  # Add it to the list.
                else:
                    file_paths.append(filepath)  # Add it to the list.

    return file_paths  # Self-explanatory.

def get_pdf_paths(file_paths):
    paths_for_pdfs = []
    for path in file_paths:
        word = path
        regexp = re.compile(r'pdf')
        if regexp.search(word):
            paths_for_pdfs.append(path)
    return paths_for_pdfs

def fragment_writer(input_directory,number_of_fragments=10):
    output_directory = 'C:/Users/Vvayne/Documents/FA 2018/Outputs/' + input_directory.split('/')[-2] + '/'
    undesirable = ['[', ']', ':', ':', ')', '(']

    print("initializing directories")
    try:
        if not os.path.exists(output_directory):
            os.makedirs(output_directory)
    except:
        output_directory = input_directory + 'Outputs/' + os.path.basename(input_directory)
        if not os.path.exists(output_directory):
            os.makedirs(output_directory)
    os.chdir(output_directory)
    print("getting filepaths")
    files = get_filepaths(input_directory,type='document')
    files = [x if x.endswith('.docx') else '' for x in files]
    chunks_per_fragment = randint(3,10)
    # number_of_words = ''.join([int(x) if isinstance(x,int) else randint(10,40) for x in input("number of words?")])
    f = docx.Document()
    print("gathering " + str(number_of_fragments) + " fragments with " + str(chunks_per_fragment) + " chunks in each")
    for i in range(number_of_fragments):
        # p = f.add_paragraph()
        chunk_list = []
        while len(chunk_list) < chunks_per_fragment:
            current_file = sample(files, 1)[0]
            # try:
                # print(current_file)
            chunk = extract_from_document(current_file,output_type='chunk',output_size=randint(4,14))
            if chunk:
                while len([c for c in undesirable if c in chunk]) > 0:
                    chunk = extract_from_document(current_file, output_type='chunk', output_size=randint(4, 14))
                # if not re.search(r'\[([^]]+)\]', chunk) or not len([c in undesirable for c in chunk]) > 0:
                tidy_chunk = tokenize.word_tokenize(chunk)
                for word in tokenize.word_tokenize(chunk):
                    try:
                        if not word or not word[0].isalpha() or not word.lower() in english_vocab:
                            tidy_chunk.remove(word)
                        if len(word) == 1 and not word.lower() in ['i','a']:
                            tidy_chunk.remove(word)
                    except:
                        continue

                chunk_list.append(' '.join(tidy_chunk))
            # except:
            #     print("couldn't get a chunk from " + current_file)
            # chunk_list = [str(chunk) if not isinstance(chunk,str) else chunk for chunk in chunk_list]
        if chunk_list and len(chunk_list) >= chunks_per_fragment:
            p = f.add_paragraph()
            while chunk_list:
                if len(chunk_list) > 2:
                    num_of_chunks = randint(1,int((len(chunk_list)-1)/2))
                else:
                    num_of_chunks = len(chunk_list)
                chunks = sample(chunk_list,num_of_chunks)
                p.add_run(' '.join(chunks))
                p.add_run().add_break()
                chunk_list = [chunk for chunk in chunk_list if not chunk in chunks]
            # fragment = ' '.join(chunk_list)
            # fragments.append(fragment)
            print("Fragment #" + str(i) + ' added.')
    title = ''
    while not title:
        try:
            title = extract_from_document(sample(get_filepaths(input_directory), 1)[0], output_type='chunk',
                                          output_size=randint(4, 14))
        except:
            continue
    clean_title = re.sub(r'[^\w\s]', '', title)

    try:
        f.save(clean_title + ' ' + str(number_of_fragments) + ' x ' + str(chunks_per_fragment) + '.docx')
    except:
        f.save(clean_title + str())
    print("Saved fragments")


# input_directory = 'C:/Users/Vvayne/Documents/FA 2018/'
# input_directory = 'C:/Users/Vvayne/Documents/FA 2018/Critical Thesis/'
alt_directory = 'C:/Users/Vvayne/Documents/FA 2018/Multigenre Workshop/Workshop/'

# print(poem_as_list)
# print(poem_as_list)
# keywords = ['mechanics']
# relevant_directories,keywords,output_directory = initialize_keyword_directories(input_directory,keywords)
# print(keywords)
# print(relevant_directories)
#
# relevant_paths = []
# for directory in relevant_directories:
#     paths = get_document_paths(get_filepaths(directory))
#     for path in paths:
#         if not os.path.basename(path) in relevant_paths:
#             relevant_paths.append(path)
#
# poem_as_list = text_by_keywords(keywords)




# for fragment in fragments:
#     p = f.add_paragraph()
#     fragment_in_words = nltk.tokenize.word_tokenize(fragment)
#     chunks = partition_word_list(fragment_in_words)
#     title = ' '.join(chunks[0])
#     for chunk in chunks:
#         p.add_run(' '.join(chunk) + ' ')
#     p.add_run().add_break()
# alt_directory = 'C:/Users/Vvayne/Documents/FA 2018/Multigenre Workshop/current inputs/'
# output_directory = "C:/Users/Vvayne/Documents/FA 2018/Multigenre Workshop/current outputs/"
fragment_writer(alt_directory, 4)

