##program that takes one or more documents, parses out
import time
import docx
from docx.shared import Inches
import os
import re
from random import sample, randint, shuffle
from nltk import tokenize, pos_tag, FreqDist
import nltk
import string
import shutil
def get_filepaths(directory):
    """
    This function will generate the file names in a directory
    tree by walking the tree either top-down or bottom-up. For each
    directory in the tree rooted at directory top (including top itself),
    it yields a 3-tuple (dirpath, dirnames, filenames).
    """
    file_paths = []  # List which will store all of the full filepaths.

    # Walk the tree.
    for root, directories, files in os.walk(directory):
        for filename in files:
            # Join the two strings in order to form the full filepath.
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)  # Add it to the list.

    return file_paths  # Self-explanatory.
def get_document_paths(file_paths):
    paths_for_potentials = []
    for path in file_paths:
        word = path
        if word.endswith('.docx') or word.endswith('.pdf') or word.endswith('.doc'):
            if not os.path.basename(path).startswith('~$'):
                paths_for_potentials.append(path)
    return paths_for_potentials
def extract_all_text_from_docx(filepath, paragraph_or_runs='runs'):
    try:
        doc = docx.Document(filepath)
    except:
        print("couldn't scrape it")
        all_runs = []
        return all_runs
    if doc:
        # print("attempting to scrape " + os.path.basename(filepath))
        file_fail_count = 0
        all_paragraphs = []
        all_runs = []
        length = len(doc.paragraphs)
        # print(str(length))
        if length >= 1:
            for paragraph in doc.paragraphs:
                all_paragraphs.append(paragraph)
                for run in paragraph.runs:
                    all_runs.append(run.text)
            if paragraph_or_runs == 'paragraphs':
                return all_paragraphs
            else:
                # print("got all " + str(length) + " run(s)")
                # print("sample: " + all_runs[0])
                return all_runs

#initialize variables
input_directory = 'C:/Users/Vvayne/Documents/SUM 2018/SWP/'
# input_path =
list_of_lines = []
list_of_sentences = []
list_of_quotes = []
end_punctuation = ['.',';',':','—']
number_of_paragraphs = randint(28-32)
list_of_cohesive_ties = ['also', 'moreover', 'and', 'in addition', 'besides', 'as well', 'furthermore', 'not only',
                         'additionally', 'whereas', 'besides','but','however','in contrast','instead','conversely',
                         'it may be the case that', 'certainly', 'also', 'likewise', 'naturally', 'nevertheless',
                         'of course', 'on the contrary', 'on the other hand', 'regardless', 'granted', 'like',
                         'alternatively', 'still', 'whereas', 'while', 'yet' , 'although', 'despite', 'it is true that',
                         'notwithstanding', 'it may appear', 'regardless', 'certainly', 'granted that', 'naturally',
                         'I admit that', 'it may be the case that', 'admittedly', 'all things considered',
                         'as a general rule', 'as far as we know', 'astonishingly', 'broadly', 'by and large',
                         'characteristically', 'clearly', 'coincidentally', 'conveniently', 'curiously',
                         'disappointingly', 'equally', 'essentially', 'explicitly', 'even so', 'eventually',
                         'fortunately', 'fundamentally', 'generally speaking', 'interestingly', 'ironically',
                         'in essence', 'in general', 'in particular', 'in practice', 'in reality',
                         'in retrospect/hindsight', 'in theory', 'in view of this', 'more interestingly',
                         'more seriously', 'more specifically', 'naturally', 'on balance', 'obviously', 'on reflection',
                         'overall', 'paradoxically', 'potentially', 'predictably', 'presumably', 'primarily', 'probably',
                         'remarkably', 'seemingly', 'significantly', 'surprisingly', 'theoretically',
                         'to all intents and purposes', 'typically', 'ultimately', 'understandably', 'undoubtedly',
                         'unfortunately', 'with hindsight']

f = docx.Document()

#get filepaths in input directory
doc_paths = get_document_paths(get_filepaths(input_directory))

tokenizer = nltk.tokenize.RegexpTokenizer(r'\w+|[^\w\s]+')

content_text = ' '.join(article.content for article in articles)
tokenized_content = tokenizer.tokenize(content_text)
content_model = NgramModel(3, tokenized_content)

starting_words = content_model.generate(100)[-2:]
content = content_model.generate(words_to_generate, starting_words)
print (' '.join(content))

#extract lines from text
# for input_file in doc_paths:
#     try:
#         list_of_runs = extract_all_text_from_docx(input_file)
#
#     #convert to sentences if not a sentence and store in list
#     for run in list_of_runs:
#         detected_sentences = tokenize.sent_tokenize(run)
#         if len(detected_sentences) > 1:
#             for sentence in detected_sentences:
#                 quote = re.search('”(.+?)”', sentence)
#                 if quote:
#                     list_of_quotes.append(quote.group(1))
#                 else:
#                     cohesive_tie_bool = randint(0,1)
#                     if cohesive_tie_bool == 1:
#                         cohesive_tie = sample(list_of_cohesive_ties,1)[0]
#                         line = cohesive_tie + ', ' + sentence[0].lower() + sentence[1:-1]
#                     else:
#                         line = sentence
#                     if not line[0].isupper():
#                         #capitalize first word
#                         line[0] = line[0].upper()
#
#                     if not line[-1] in end_punctuation or not line[-1] == '?':
#                         #append with punctuation
#                         line = line + str(sample(end_punctuation),1)[0]
#                     if not line in list_of_sentences:
#                         list_of_sentences.append(line)
#         else:
#             quote = re.search('”(.+?)”', run)
#             if quote:
#                 list_of_quotes.append(quote.group(1))
#             else:
#                 cohesive_tie_bool = randint(0, 1)
#                 if cohesive_tie_bool == 1:
#                     cohesive_tie = sample(list_of_cohesive_ties, 1)[0]
#                     line = cohesive_tie + ', ' + run[0].lower() + sentence[1:-1]
#                 else:
#                     line = run
#                 if not line[0].isupper():
#                     # capitalize first word
#                     line[0] = line[0].upper()
#
#                 if not line[-1] in end_punctuation or not line[-1] == '?':
#                     # append with punctuation
#                     line = line + str(sample(end_punctuation), 1)[0]
#                 list_of_sentences.append(line)
#
# # populate paragraphs by sampling without repetition
# list_of_available_sentences = list_of_sentences
# for i in range(number_of_paragraphs):
# 	sent_per_para = randint(1,6)
# 	p = f.add_paragraph()
# 	for j in range(sent_per_para):
# 		current_sentence = sample(list_of_available_sentences,1)[0]
# 		p.add_run(current_sentence)
# 		list_of_available_sentences.remove(current_sentence)
#
# output_directory = input_directory

